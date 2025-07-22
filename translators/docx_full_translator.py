import os
import logging
from docx import Document
from typing import Optional, Dict, Any, Union, List, Tuple
import hashlib
from tqdm import tqdm
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import threading
import time
from enum import Enum
from copy import deepcopy
from dataclasses import dataclass

# 在最开始就设置日志级别，屏蔽所有DEBUG信息
logging.getLogger().setLevel(logging.INFO)

# 屏蔽所有可能的第三方库DEBUG日志
third_party_loggers = [
    'urllib3', 'requests', 'httpx', 'siliconflow', 'openai',
    'anthropic', 'zhipuai', 'dashscope', 'httpcore', 'httpx._client',
    'httpx._config', 'httpx._models', 'httpx._auth', 'requests.packages.urllib3',
    'requests_oauthlib', 'oauthlib', 'aiohttp', 'websockets', 'asyncio'
]

for logger_name in third_party_loggers:
    logging.getLogger(logger_name).setLevel(logging.WARNING)

# 设置过滤器屏蔽第三方DEBUG信息
class NoDebugFilter(logging.Filter):
    def filter(self, record):
        if record.levelno >= logging.WARNING:
            return True
        if record.name == __name__ or record.name.startswith('__main__'):
            return record.levelno >= logging.INFO
        return False

root_logger = logging.getLogger()
root_logger.addFilter(NoDebugFilter())

logger = logging.getLogger(__name__)

# 配置合并策略枚举
class ConfigMergeStrategy(Enum):
    MERGE = "merge"              # 默认：智能合并前端和实例配置
    OVERRIDE = "override"        # 前端配置完全覆盖
    INSTANCE_ONLY = "instance_only"  # 只使用实例配置

# 失败原因枚举 - 新增
class FailureReason(Enum):
    """失败原因枚举"""
    TIMEOUT = "timeout"
    API_ERROR = "api_error"
    PARSE_ERROR = "parse_error"
    NOT_TRANSLATED = "not_translated"
    EMPTY_RESPONSE = "empty_response"
    CONNECTION_ERROR = "connection_error"
    BATCH_FAILURE = "batch_failure"
    FORMAT_ERROR = "format_error"

# 失败任务数据类 - 新增
@dataclass
class FailedTask:
    """失败任务数据类 - 修复版本，支持结果更新"""
    original_text: str
    paragraph: Any  # 段落对象引用
    paragraph_data: Dict[str, Any]  # 段落数据
    batch_id: int  # 批次ID
    task_id: str  # 任务ID，用于重试时直接更新
    failure_reason: FailureReason
    retry_count: int = 0
    error_message: str = ""
    is_serious: bool = True
    
    def __post_init__(self):
        self.failure_timestamp = time.time()

# 全局配置合并策略和实例配置
_global_merge_strategy = ConfigMergeStrategy.MERGE
_instance_prompt_config = None

# 优化的批处理设置 - 增强版
OPTIMAL_BATCH_SIZE = 50  # 批次大小
MAX_CHARS_PER_BATCH = 8000  # 限制每批次字符数
DEFAULT_TRANSLATION_TIMEOUT = 60  # 翻译超时时间（秒）
DEFAULT_MAX_RETRIES = 10           # 最大重试次数
DEFAULT_RETRY_DELAY = 2           # 重试延迟（秒）
DEFAULT_RETRY_BATCH_SIZE = 10     # 重试时的批次大小 - 修改为更小的值
DEFAULT_THREADS = 5              # 默认线程数
DEFAULT_RETRY_THREADS = 3        # 默认重试线程数 - 新增

_translation_cache = {}
_cache_lock = threading.Lock()

class TimeoutException(Exception):
    """超时异常"""
    pass

class TranslationValidator:
    """翻译完整性验证器 - 修复版本，减少误判"""
    
    # 扩展的错误关键词检测
    ERROR_KEYWORDS = [
        # 英文错误关键词
        'timeout', 'readtimeout', 'connecttimeout', 'httptimeout',
        'network error', 'connection error', 'api error', 'service error',
        'translation failed', 'service unavailable', 'request failed',
        'server error', 'bad gateway', 'gateway timeout', 'error:',
        # 中文错误关键词
        '超时', '网络错误', '连接错误', '服务错误', '翻译失败',
        '服务不可用', '请求失败', '服务器错误', '错误:'
    ]
    
    @staticmethod
    def is_error_message(text: str) -> bool:
        """检查是否为错误消息"""
        if not text:
            return True
            
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in TranslationValidator.ERROR_KEYWORDS)
    
    @staticmethod
    def is_serious_failure(original_text: str, translated_text: str, 
                          large_text_threshold: int = 50) -> Tuple[bool, str]:
        """改进的失败检测逻辑 - 减少误判"""
        
        # 1. 首先检查明显的错误消息
        if TranslationValidator.is_error_message(translated_text):
            return True, f"API错误消息: {translated_text[:50]}..."
        
        # 2. 检查空值或无效响应
        if not translated_text or translated_text.strip() == "":
            return len(original_text) > 20, f"空翻译结果（原文{len(original_text)}字符）"
        
        # 3. 改进的未翻译检测 - 更宽松的判断
        is_unchanged = original_text.strip() == translated_text.strip()
        if is_unchanged:
            # 检查是否可能是正常情况（专有名词、数字、代码等）
            if TranslationValidator._is_likely_untranslatable(original_text):
                return False, "可能是专有名词或代码，保持原文正常"
            
            # 长文本完全未变化才认为是严重失败
            if len(original_text) > large_text_threshold:
                return True, f"长文本未翻译（{len(original_text)}字符）"
            else:
                return False, f"短文本未变化（{len(original_text)}字符，可能正常）"
        
        # 4. 质量检测：检查翻译是否合理 - 降低阈值
        quality_score = TranslationValidator._assess_translation_quality(original_text, translated_text)
        if quality_score < 0.2:  # 降低阈值，减少误判
            return True, f"翻译质量过低（分数: {quality_score:.2f}）"
        
        return False, "翻译正常"
    
    @staticmethod
    def _is_likely_untranslatable(text: str) -> bool:
        """判断文本是否可能不需要翻译 - 扩展识别范围"""
        text = text.strip()
        
        # 数字、日期、代码
        if re.match(r'^[\d\-\/\.\s\:\(\)]+$', text):
            return True
        
        # URL、邮箱
        if re.match(r'^https?://|.*@.*\..*', text):
            return True
        
        # 专有名词比例 - 降低阈值
        uppercase_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if uppercase_ratio > 0.3:  # 更宽松
            return True
        
        # 常见的不需要翻译的模式
        common_patterns = [
            r'^[A-Z]{2,}$',  # 全大写缩写
            r'^\d+[\.\)]\s*$',  # 列表编号
            r'^[a-zA-Z0-9_\-\.]+$',  # 标识符
        ]
        
        for pattern in common_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    @staticmethod
    def _assess_translation_quality(original: str, translated: str) -> float:
        """评估翻译质量（0-1分数）- 更宽松的评估"""
        if not translated or not original:
            return 0.0
        
        score = 1.0
        
        # 长度检查 - 更宽松
        length_ratio = len(translated) / max(len(original), 1)
        if length_ratio < 0.2 or length_ratio > 5.0:  # 更宽松的范围
            score -= 0.2  # 减少扣分
        
        # 字符多样性检查 - 更宽松
        if len(set(translated)) / max(len(translated), 1) < 0.05:  # 更宽松的阈值
            score -= 0.1  # 减少扣分
        
        # 重复检查 - 更宽松
        if len(translated) > 20 and translated.count(translated[:10]) > 5:  # 检测重复
            score -= 0.2
        
        return max(0.0, score)

class FormattedDocxTranslator:
    """格式化DOCX翻译器 - 增强版，支持重试机制"""
    
    def __init__(self, translator, target_lang: str, source_lang: Optional[str] = None,
                 translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT,
                 max_retries: int = DEFAULT_MAX_RETRIES,
                 retry_delay: int = DEFAULT_RETRY_DELAY,
                 retry_batch_size: int = DEFAULT_RETRY_BATCH_SIZE,
                 max_retry_workers: int = DEFAULT_RETRY_THREADS,
                 large_text_threshold: int = 50,
                 retry_failure_threshold: float = 0.0,
                 prompt_config=None):
        self.translator = translator
        self.target_lang = target_lang
        self.source_lang = source_lang
        self.translation_timeout = translation_timeout
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.retry_batch_size = retry_batch_size
        self.max_retry_workers = max_retry_workers
        self.large_text_threshold = large_text_threshold
        self.retry_failure_threshold = retry_failure_threshold
        self.prompt_config = prompt_config
        
        self.stats = {
            'success': 0, 
            'failed': 0, 
            'skipped': 0, 
            'timeout': 0, 
            'retried': 0,
            'final_failed': 0,
            'serious_failures': 0,
            'minor_issues': 0,
            'retry_attempts': 0,
            'rescued_tasks': 0,
            'retry_success_rate': 0.0,
            'total_retry_tasks': 0,
            'parallel_retry_used': 0,
            'cache_hits': 0,
            'api_calls': 0
        }
        
        # 时间统计 - 新增
        self.time_stats = {
            'total_start_time': 0,
            'main_translation_time': 0,
            'retry_time': 0,
            'total_time': 0,
            'document_load_time': 0,
            'document_save_time': 0
        }
        
        # 失败任务追踪
        self.failed_tasks = []
        self.failed_tasks_lock = threading.Lock()
        
        # 增强重试策略配置 - 改为 [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_delays = [1, 2, 4, 8, 12, 16, 20, 25]  # 更合理的延迟

def set_config_merge_strategy(strategy: Union[str, ConfigMergeStrategy]):
    """设置全局配置合并策略"""
    global _global_merge_strategy
    if isinstance(strategy, str):
        try:
            _global_merge_strategy = ConfigMergeStrategy(strategy)
        except ValueError:
            logger.warning(f"Invalid merge strategy: {strategy}, using default 'merge'")
            _global_merge_strategy = ConfigMergeStrategy.MERGE
    else:
        _global_merge_strategy = strategy
    
    logger.info(f"Formatted DOCX translator config merge strategy set to: {_global_merge_strategy.value}")

def set_instance_prompt_config(config: Optional[Dict[str, Any]]):
    """设置实例级别的prompt配置"""
    global _instance_prompt_config
    _instance_prompt_config = deepcopy(config) if config else None
    if config:
        logger.info(f"Formatted DOCX translator instance prompt config set: mode={config.get('mode', 'none')}")

def get_effective_config(frontend_config: Optional[Dict[str, Any]] = None, 
                        merge_strategy: Optional[Union[str, ConfigMergeStrategy]] = None) -> Optional[Dict[str, Any]]:
    """获取当前有效配置，基于合并策略"""
    global _instance_prompt_config, _global_merge_strategy
    
    # 确定使用的合并策略
    if merge_strategy:
        if isinstance(merge_strategy, str):
            try:
                strategy = ConfigMergeStrategy(merge_strategy)
            except ValueError:
                strategy = _global_merge_strategy
        else:
            strategy = merge_strategy
    else:
        strategy = _global_merge_strategy
    
    logger.debug(f"Formatted DOCX translator using merge strategy: {strategy.value}")
    
    # 根据策略合并配置
    if strategy == ConfigMergeStrategy.INSTANCE_ONLY:
        effective_config = deepcopy(_instance_prompt_config) if _instance_prompt_config else None
        logger.debug("Using instance-only configuration for formatted DOCX translator")
    
    elif strategy == ConfigMergeStrategy.OVERRIDE:
        effective_config = deepcopy(frontend_config) if frontend_config else deepcopy(_instance_prompt_config)
        logger.debug("Using override strategy for formatted DOCX translator (frontend takes precedence)")
    
    else:  # ConfigMergeStrategy.MERGE
        effective_config = _merge_prompt_configs(_instance_prompt_config, frontend_config)
        logger.debug("Using intelligent merge strategy for formatted DOCX translator")
    
    # 标准化最终配置
    if effective_config:
        effective_config = _normalize_prompt_config(effective_config)
        logger.debug(f"Effective formatted DOCX config: mode={effective_config.get('mode', 'none')}")
    
    return effective_config

def _merge_prompt_configs(instance_config: Optional[Dict[str, Any]], 
                         frontend_config: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """核心配置合并逻辑 - 智能合并前端和实例配置"""
    
    if not instance_config and not frontend_config:
        return None
    
    if not instance_config:
        return deepcopy(frontend_config)
    
    if not frontend_config:
        return deepcopy(instance_config)
    
    # 开始智能合并
    merged = deepcopy(instance_config)
    
    logger.debug(f"Merging formatted DOCX configs - Instance: {instance_config.get('mode', 'none')}, Frontend: {frontend_config.get('mode', 'none')}")
    
    # 1. 模式合并：前端优先
    if 'mode' in frontend_config:
        merged['mode'] = frontend_config['mode']
    
    # 2. 基础字段合并：前端优先
    basic_fields = ['prompt_template', 'professional_domain', 'custom_prompt', 
                   'custom_system_prompt', 'custom_user_prompt', 'system', 'user']
    
    for field in basic_fields:
        if field in frontend_config:
            merged[field] = frontend_config[field]
    
    # 3. 智能字段合并
    
    # preserve_terms：术语列表合并去重
    merged_terms = set()
    for config in [instance_config, frontend_config]:
        terms = config.get('preserve_terms', [])
        if isinstance(terms, str):
            terms = [term.strip() for term in terms.split(',') if term.strip()]
        elif isinstance(terms, list):
            terms = [str(term).strip() for term in terms if str(term).strip()]
        merged_terms.update(terms)
    
    if merged_terms:
        merged['preserve_terms'] = list(merged_terms)
        logger.debug(f"Merged preserve_terms for formatted DOCX: {len(merged_terms)} unique terms")
    
    # glossary：术语表字典合并（前端优先）
    merged_glossary = {}
    if instance_config.get('glossary') and isinstance(instance_config['glossary'], dict):
        merged_glossary.update(instance_config['glossary'])
    if frontend_config.get('glossary') and isinstance(frontend_config['glossary'], dict):
        merged_glossary.update(frontend_config['glossary'])  # 前端优先覆盖
    
    if merged_glossary:
        merged['glossary'] = merged_glossary
        logger.debug(f"Merged glossary for formatted DOCX: {len(merged_glossary)} entries")
    
    # additional_context：上下文信息拼接
    contexts = []
    for config in [instance_config, frontend_config]:
        context = config.get('additional_context', '').strip()
        if context:
            contexts.append(context)
    
    if contexts:
        merged['additional_context'] = ' | '.join(contexts)
        logger.debug(f"Merged additional_context for formatted DOCX: {len(merged['additional_context'])} chars")
    
    # 4. 批处理设置合并：取更优的值
    batch_fields = {
        'chunk_size': max,
        'batch_size': max,
        'max_workers': max,
        'thread_count': max
    }
    
    for field, merge_func in batch_fields.items():
        values = []
        for config in [instance_config, frontend_config]:
            if field in config and isinstance(config[field], (int, float)):
                values.append(config[field])
        if values:
            merged[field] = merge_func(values)
    
    # 5. 质量和功能设置：逻辑OR合并
    quality_fields = ['ensure_consistency', 'quality_level', 'preserve_formatting', 
                     'advanced_mode', 'detailed_logging', 'use_cache', 'smart_format']
    
    for field in quality_fields:
        if (instance_config.get(field) or frontend_config.get(field)):
            merged[field] = frontend_config.get(field, instance_config.get(field))
    
    logger.debug(f"Formatted DOCX config merge completed: {len(merged)} fields in final config")
    return merged

def _normalize_prompt_config(config: Dict[str, Any]) -> Dict[str, Any]:
    """标准化prompt配置格式，确保与前端格式兼容"""
    normalized = config.copy()
    
    # 确保mode字段存在
    if 'mode' not in normalized:
        if 'custom_prompt' in normalized:
            normalized['mode'] = 'custom'
        elif 'prompt_template' in normalized or 'professional_domain' in normalized:
            normalized['mode'] = 'professional'
        elif any(k in normalized for k in ['preserve_terms', 'glossary', 'additional_context']):
            normalized['mode'] = 'general'
        else:
            normalized['mode'] = 'none'
    
    # 处理保留术语 - 支持逗号分隔的字符串（前端格式）
    preserve_terms = normalized.get('preserve_terms')
    if preserve_terms:
        if isinstance(preserve_terms, str):
            # 前端格式：逗号分隔的字符串
            terms_list = [term.strip() for term in preserve_terms.split(',') if term.strip()]
            normalized['preserve_terms'] = terms_list
        elif isinstance(preserve_terms, list):
            # 确保列表中的字符串都是清理过的
            normalized['preserve_terms'] = [str(term).strip() for term in preserve_terms if str(term).strip()]
    
    # 处理术语表 - 确保是字典格式
    glossary = normalized.get('glossary')
    if glossary and not isinstance(glossary, dict):
        logger.warning(f"Glossary should be a dictionary, got {type(glossary)}, ignoring")
        normalized.pop('glossary', None)
    
    # 处理自定义prompt
    if normalized.get('mode') == 'custom':
        custom_prompt = normalized.get('custom_prompt', {})
        if not custom_prompt or not isinstance(custom_prompt, dict):
            # 检查是否有分离的system和user字段
            system_prompt = normalized.get('custom_system_prompt', normalized.get('system'))
            user_prompt = normalized.get('custom_user_prompt', normalized.get('user'))
            
            if system_prompt:
                normalized['custom_prompt'] = {
                    'system': system_prompt,
                    'user': user_prompt or 'Please translate the following content to {target_lang}:\n\n{content}'
                }
            elif isinstance(custom_prompt, str):
                # 如果custom_prompt是字符串，将其作为system prompt
                normalized['custom_prompt'] = custom_prompt
            else:
                logger.warning("Custom mode selected but no valid custom prompt provided, falling back to general mode")
                normalized['mode'] = 'general'
    
    # 处理专业模板 - 前端使用 'professional_domain' 字段
    if normalized.get('mode') == 'professional':
        domain = normalized.get('professional_domain', normalized.get('prompt_template', 'academic'))
        normalized['prompt_template'] = domain
    
    return normalized

def _get_cache_key(text: str, target_lang: str, source_lang: Optional[str], prompt_config=None) -> str:
    """生成缓存键，包含prompt配置信息"""
    text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    
    # 包含prompt配置的哈希，确保不同prompt配置的翻译分开缓存
    prompt_hash = ""
    if prompt_config:
        mode = prompt_config.get('mode', '')
        template = prompt_config.get('prompt_template', '')
        custom_prompt = prompt_config.get('custom_prompt', '')
        if isinstance(custom_prompt, dict):
            custom_prompt = custom_prompt.get('system', '')[:50]
        elif isinstance(custom_prompt, str):
            custom_prompt = custom_prompt[:50]
        
        prompt_str = f"{mode}_{template}_{custom_prompt}"
        prompt_hash = hashlib.md5(prompt_str.encode('utf-8')).hexdigest()[:8]
    
    return f"{text_hash}_{target_lang}_{source_lang or 'auto'}_{prompt_hash}"

def _get_enhanced_system_prompt(target_lang: str, source_lang: Optional[str], prompt_config=None) -> str:
    """获取增强的系统提示 - 支持配置合并后的自定义prompt配置"""
    
    # 如果有自定义prompt配置，优先使用
    if prompt_config:
        prompt_mode = prompt_config.get('mode', 'default')
        
        if prompt_mode == 'custom' and prompt_config.get('custom_prompt'):
            # 使用完全自定义的prompt，确保包含批量处理说明
            logger.info("Using merged custom prompt for formatted DOCX translation")
            
            custom_prompt = prompt_config['custom_prompt']
            if isinstance(custom_prompt, dict):
                system_content = custom_prompt.get('system', '')
            else:
                system_content = str(custom_prompt)
            
            # 为批量翻译添加必要的说明
            if "numbered line" not in system_content.lower():
                system_content += f"""

BATCH PROCESSING INSTRUCTIONS:
- Each input line is numbered [1], [2], etc.
- Translate each numbered line individually
- Keep the exact same number of lines as input
- Output only the translated content, one per line
- Do not include line numbers in output
- Preserve formatting and punctuation"""
            
            return system_content
        
        elif prompt_mode == 'professional' and prompt_config.get('prompt_template'):
            # 使用专业模板 - 扩展的专业领域
            domain = prompt_config['prompt_template']
            logger.info(f"Using professional template for formatted DOCX: {domain}")
            
            professional_prompts = {
                'academic': f"""You are an expert academic translator specializing in scholarly documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain academic tone, preserve citations and references, and use appropriate academic terminology.
Ensure consistency in technical terms throughout the translation.""",
                
                'business': f"""You are a professional business translator with expertise in corporate documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use appropriate business terminology, maintain formal tone, and keep company names/brands unchanged.
Ensure clarity and professionalism in the translation.""",
                
                'technical': f"""You are a technical translator specializing in technical documentation.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve technical accuracy, keep code snippets and commands unchanged, and use industry-standard terminology.
Maintain consistency in technical terms throughout.""",
                
                'legal': f"""You are a certified legal translator with expertise in legal documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use precise legal terminology, maintain legal accuracy and formality, and preserve all legal references.
Ensure no ambiguity in legal terms.""",
                
                'medical': f"""You are a certified medical translator with expertise in medical documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use standard medical terminology, preserve drug names and dosages exactly, and maintain clinical precision.
Follow international medical nomenclature standards.""",
                
                'creative': f"""You are a creative translator focusing on maintaining style and tone.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve the original style, adapt idioms naturally, and maintain emotional impact.
Focus on readability and flow while being faithful to the original meaning.""",
                
                # 新增专业领域
                'scientific': f"""You are a scientific translator with expertise in research papers and scientific publications.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain scientific accuracy, preserve equations and formulas, use standard scientific terminology.
Ensure precision in methodology descriptions and results interpretation.""",
                
                'financial': f"""You are a financial translator specializing in financial documents and reports.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use precise financial terminology, maintain accuracy in numerical data, preserve financial statements format.
Ensure compliance with international financial reporting standards.""",
                
                'marketing': f"""You are a marketing translator focusing on promotional and marketing content.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain persuasive tone, adapt cultural references appropriately, preserve brand messaging.
Focus on engagement and cultural sensitivity while maintaining marketing effectiveness.""",
                
                'educational': f"""You are an educational content translator specializing in learning materials.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use clear and accessible language, maintain pedagogical structure, preserve learning objectives.
Ensure content remains suitable for the intended educational level."""
            }
            
            system_content = professional_prompts.get(domain, professional_prompts['academic'])
            
            # 为批量翻译调整模板
            if "numbered line" not in system_content.lower():
                system_content += f"""

BATCH PROCESSING REQUIREMENTS:
- Process each numbered line [1], [2], etc.
- Return same number of translated lines
- No line numbers in output"""
            return system_content
    
    # 默认增强prompt逻辑
    enhanced_rules = []
    if prompt_config:
        if prompt_config.get('preserve_terms'):
            terms_text = ', '.join(prompt_config['preserve_terms'])
            enhanced_rules.append(f"PRESERVE THESE TERMS EXACTLY: {terms_text}")
        
        if prompt_config.get('glossary'):
            glossary_text = '; '.join([f"{k}: {v}" for k, v in prompt_config['glossary'].items()])
            enhanced_rules.append(f"USE THIS GLOSSARY: {glossary_text}")
        
        if prompt_config.get('additional_context'):
            enhanced_rules.append(f"ADDITIONAL CONTEXT: {prompt_config['additional_context']}")

    enhanced_rules_text = "\n".join([f"{i+9}. {rule}" for i, rule in enumerate(enhanced_rules)])

    return f"""You are a professional translator. Translate from {source_lang or 'auto-detected language'} to {target_lang}.

Rules:
1. Translate each numbered line individually, but take the full context of surrounding lines into account to ensure accurate and coherent translation, ensure that the translation follows the natural speaking habits, tone, and logic of the target language.
2. Keep the exact same number of lines as the original.
3. Preserve all formatting, punctuation, and special characters.
4. For lists, keep the list markers.
5. Output only the translated lines, one per line, in the same order.
6. Do not include the original line numbers or any extra comments in your output.
7. Do not translate place names (e.g. cities, countries) or company names—keep them exactly as in the original.
8. Return only the translated content without including the original text.
{enhanced_rules_text}"""

def _call_translator_with_timeout(translator, text: str, target_lang: str, source_lang: Optional[str],
                                prompt_config: Optional[Dict[str, Any]] = None, 
                                timeout: Optional[int] = None) -> str:
    """带超时的翻译调用"""
    timeout = timeout or DEFAULT_TRANSLATION_TIMEOUT
    
    # 使用ThreadPoolExecutor实现超时
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(_call_translator, translator, text, target_lang, source_lang, prompt_config)
        try:
            result = future.result(timeout=timeout)
            return result
        except TimeoutError:
            logger.warning(f"Translation timeout after {timeout} seconds")
            return f"Error: Translation timeout after {timeout} seconds"
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return f"Error: {str(e)}"

def _call_translator(translator, text: str, target_lang: str, source_lang: Optional[str],
                   prompt_config: Optional[Dict[str, Any]] = None) -> str:
    """调用翻译器"""
    try:
        # 获取系统提示
        system_prompt = _get_enhanced_system_prompt(target_lang, source_lang, prompt_config)
        
        # 支持配置合并策略的翻译器调用
        try:
            # 优先尝试使用配置合并的翻译器接口
            if hasattr(translator, 'translate') and hasattr(translator, 'set_prompt_config'):
                translated = translator.translate(
                    text=text,
                    target_lang=target_lang,
                    source_lang=source_lang,
                    prompt_config=prompt_config
                )
            else:
                # 回退到messages格式
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ]
                translated = translator.translate(
                    messages=messages,
                    target_lang=target_lang,
                    source_lang=source_lang
                )
        except (TypeError, AttributeError):
            # 如果不支持messages格式，使用传统格式
            full_prompt = f"{system_prompt}\n\nText to translate:\n{text}"
            translated = translator.translate(full_prompt, target_lang, source_lang)
        
        return translated
        
    except Exception as e:
        logger.error(f"Translation failed: {e}")
        return f"Error: {str(e)}"

def _extract_numbered_translations(response: str, expected_count: int) -> list:
    """从响应中提取按行号标记的翻译结果"""
    translations = [""] * expected_count
    
    # 按行分割响应
    lines = response.strip().split('\n')
    
    current_line_num = 0
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 检查是否以行号开头 [1], [2], 等
        number_match = re.match(r'^\[(\d+)\]\s*(.*)', line)
        if number_match:
            line_num = int(number_match.group(1)) - 1  # 转换为0基索引
            content = number_match.group(2).strip()
            
            if 0 <= line_num < expected_count and content:
                translations[line_num] = content
        else:
            # 如果没有行号标记，按顺序分配
            if current_line_num < expected_count and line:
                translations[current_line_num] = line
                current_line_num += 1
    
    # 对于空的翻译，尝试直接按行顺序分配
    if any(not t for t in translations):
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        for i, line in enumerate(non_empty_lines):
            if i < expected_count and not translations[i]:
                # 移除可能的行号前缀
                clean_line = re.sub(r'^\[\d+\]\s*', '', line)
                if clean_line:
                    translations[i] = clean_line
    
    return translations

def _translate_batch_enhanced(texts: list, translator, target_lang: str, source_lang: Optional[str], 
                            prompt_config=None, batch_info: Optional[Dict] = None,
                            translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT) -> tuple:
    """
    增强的批量翻译文本，支持失败检测和重试机制
    返回: (translations, failed_indices, error_message)
    """
    if not texts:
        return [], [], None
        
    translations = []
    uncached_texts = []
    uncached_indices = []
    failed_indices = []
    error_message = None
    cache_flags = [False] * len(texts)  # 跟踪哪些结果来自缓存
    
    # 线程安全的缓存检查
    with _cache_lock:
        for i, text in enumerate(texts):
            cache_key = _get_cache_key(text, target_lang, source_lang, prompt_config)
            cached = _translation_cache.get(cache_key)
            if cached:
                translations.append(cached)
                cache_flags[i] = True  # 标记为缓存结果
            else:
                uncached_texts.append(text)
                uncached_indices.append(i)
                translations.append("")  # 占位
    
    # 翻译未缓存的文本
    if uncached_texts:
        try:
            # 使用行号标记每个段落
            numbered_texts = []
            for i, text in enumerate(uncached_texts):
                numbered_texts.append(f"[{i+1}] {text}")
            
            user_message = "\n".join(numbered_texts)
            
            # 带超时的翻译
            response = _call_translator_with_timeout(translator, user_message, target_lang, source_lang, 
                                                   prompt_config, translation_timeout)
            
            if response and not response.startswith("Error:"):
                # 提取翻译结果
                translated_parts = _extract_numbered_translations(response, len(uncached_texts))
                
                # 线程安全的缓存更新
                with _cache_lock:
                    for i, (text, idx) in enumerate(zip(uncached_texts, uncached_indices)):
                        if i < len(translated_parts) and translated_parts[i]:
                            translation = translated_parts[i].strip()
                            if translation and translation != text:
                                cache_key = _get_cache_key(text, target_lang, source_lang, prompt_config)
                                _translation_cache[cache_key] = translation
                                translations[idx] = translation
                            else:
                                # 翻译失败或为空
                                translations[idx] = text  # 保持原文
                                failed_indices.append(idx)
                        else:
                            # 翻译失败或为空
                            translations[idx] = text  # 保持原文
                            failed_indices.append(idx)
            else:
                # 整个批次翻译失败
                error_message = response if response else "Unknown translation error"
                failed_indices = uncached_indices
                for idx in uncached_indices:
                    translations[idx] = texts[idx]  # 保持原文
                    
        except Exception as e:
            logger.error(f"批量翻译失败: {e}")
            error_message = str(e)
            # 简单回退：标记为失败
            failed_indices = uncached_indices
            for idx in uncached_indices:
                translations[idx] = texts[idx]  # 保持原文
                
    return translations, failed_indices, error_message, cache_flags

def _create_smart_batches(texts, max_batch_size=OPTIMAL_BATCH_SIZE, max_chars_per_batch=MAX_CHARS_PER_BATCH):
    """智能分批：按字符数分组"""
    batches = []
    current_batch = []
    current_chars = 0
    
    for text in texts:
        text_len = len(text)
        
        # 检查是否需要开始新批次
        if (len(current_batch) >= max_batch_size or 
            current_chars + text_len > max_chars_per_batch) and current_batch:
            batches.append(current_batch)
            current_batch = []
            current_chars = 0
        
        current_batch.append(text)
        current_chars += text_len
    
    if current_batch:
        batches.append(current_batch)
    
    return batches

def _get_paragraph_format(paragraph):
    """获取段落格式"""
    format_info = {}
    
    # 安全地获取段落格式属性
    try:
        format_info['alignment'] = paragraph.alignment
    except:
        format_info['alignment'] = None
        
    pf = paragraph.paragraph_format
    if pf:
        try:
            format_info['line_spacing'] = pf.line_spacing
            format_info['space_before'] = pf.space_before
            format_info['space_after'] = pf.space_after
            format_info['first_line_indent'] = pf.first_line_indent
            format_info['left_indent'] = pf.left_indent
            format_info['right_indent'] = pf.right_indent
            format_info['keep_together'] = pf.keep_together
            format_info['keep_with_next'] = pf.keep_with_next
            format_info['page_break_before'] = pf.page_break_before
            format_info['widow_control'] = pf.widow_control
        except Exception as e:
            logger.warning(f"Error getting paragraph format: {e}")
    
    try:
        format_info['style'] = paragraph.style.name if paragraph.style else None
    except:
        format_info['style'] = None
        
    return format_info

def _apply_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    if format_info.get('style'):
        try:
            paragraph.style = format_info['style']
        except:
            pass
            
    if format_info.get('alignment') is not None:
        try:
            paragraph.alignment = format_info['alignment']
        except:
            pass
    
    pf = paragraph.paragraph_format
    if pf:
        try:
            if format_info.get('line_spacing') is not None:
                pf.line_spacing = format_info['line_spacing']
            if format_info.get('space_before') is not None:
                pf.space_before = format_info['space_before']
            if format_info.get('space_after') is not None:
                pf.space_after = format_info['space_after']
            if format_info.get('first_line_indent') is not None:
                pf.first_line_indent = format_info['first_line_indent']
            if format_info.get('left_indent') is not None:
                pf.left_indent = format_info['left_indent']
            if format_info.get('right_indent') is not None:
                pf.right_indent = format_info['right_indent']
            if format_info.get('keep_together') is not None:
                pf.keep_together = format_info['keep_together']
            if format_info.get('keep_with_next') is not None:
                pf.keep_with_next = format_info['keep_with_next']
            if format_info.get('page_break_before') is not None:
                pf.page_break_before = format_info['page_break_before']
            if format_info.get('widow_control') is not None:
                pf.widow_control = format_info['widow_control']
        except Exception as e:
            logger.warning(f"Error applying paragraph format: {e}")

def _get_run_format(run):
    """获取文本运行的格式"""
    format_info = {}
    
    # 基本格式
    try:
        format_info['bold'] = run.bold
        format_info['italic'] = run.italic
        format_info['underline'] = run.underline
    except:
        pass
    
    # 字体相关
    if hasattr(run, 'font') and run.font:
        try:
            if hasattr(run.font, 'size') and run.font.size:
                format_info['font_size'] = run.font.size
            if hasattr(run.font, 'name') and run.font.name:
                format_info['font_name'] = run.font.name
            
            # 更安全地处理颜色
            try:
                if hasattr(run.font, 'color') and run.font.color:
                    if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                        format_info['color'] = run.font.color.rgb
            except Exception:
                pass
                
            # 更安全地处理高亮颜色
            try:
                if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                    # 只有当highlight_color不是None或'none'时才保存
                    if run.font.highlight_color and str(run.font.highlight_color).lower() != 'none':
                        format_info['highlight_color'] = run.font.highlight_color
            except Exception:
                pass
                
            if hasattr(run.font, 'all_caps'):
                format_info['all_caps'] = run.font.all_caps
            if hasattr(run.font, 'small_caps'):
                format_info['small_caps'] = run.font.small_caps
            if hasattr(run.font, 'strike'):
                format_info['strike'] = run.font.strike
            if hasattr(run.font, 'double_strike'):
                format_info['double_strike'] = run.font.double_strike
            if hasattr(run.font, 'subscript'):
                format_info['subscript'] = run.font.subscript
            if hasattr(run.font, 'superscript'):
                format_info['superscript'] = run.font.superscript
            if hasattr(run.font, 'emphasis') and run.font.emphasis:
                format_info['emphasis'] = run.font.emphasis
        except Exception as e:
            if "WD_COLOR_INDEX" not in str(e):  # 忽略颜色索引相关的警告
                logger.warning(f"Error getting run format: {e}")
            
    return format_info


def _apply_run_format(run, format_info):
    """应用文本运行的格式"""
    # 基本格式
    try:
        if 'bold' in format_info and format_info['bold'] is not None:
            run.bold = format_info['bold']
        if 'italic' in format_info and format_info['italic'] is not None:
            run.italic = format_info['italic']
        if 'underline' in format_info and format_info['underline'] is not None:
            run.underline = format_info['underline']
    except Exception:
        pass
    
    # 字体相关
    if hasattr(run, 'font') and run.font:
        try:
            if 'font_size' in format_info and format_info['font_size'] is not None:
                run.font.size = format_info['font_size']
            if 'font_name' in format_info and format_info['font_name'] is not None:
                run.font.name = format_info['font_name']
            
            # 安全地应用颜色
            if 'color' in format_info and format_info['color'] is not None:
                try:
                    if hasattr(run.font, 'color') and run.font.color:
                        run.font.color.rgb = format_info['color']
                except Exception:
                    pass
                    
            # 安全地应用高亮颜色
            if 'highlight_color' in format_info and format_info['highlight_color'] is not None:
                try:
                    run.font.highlight_color = format_info['highlight_color']
                except Exception:
                    pass
                    
            if 'all_caps' in format_info and format_info['all_caps'] is not None:
                run.font.all_caps = format_info['all_caps']
            if 'small_caps' in format_info and format_info['small_caps'] is not None:
                run.font.small_caps = format_info['small_caps']
            if 'strike' in format_info and format_info['strike'] is not None:
                run.font.strike = format_info['strike']
            if 'double_strike' in format_info and format_info['double_strike'] is not None:
                run.font.double_strike = format_info['double_strike']
            if 'subscript' in format_info and format_info['subscript'] is not None:
                run.font.subscript = format_info['subscript']
            if 'superscript' in format_info and format_info['superscript'] is not None:
                run.font.superscript = format_info['superscript']
            if 'emphasis' in format_info and format_info['emphasis'] is not None:
                run.font.emphasis = format_info['emphasis']
        except Exception:
            pass

def _all_runs_same_format(runs_data):
    """检查所有run是否具有相同的格式"""
    if not runs_data:
        return True
    
    first_format = runs_data[0][1]
    for _, format_info in runs_data[1:]:
        if format_info != first_format:
            return False
    return True

def _apply_translated_text_with_format(paragraph, translated_text, runs_data):
    """智能应用翻译文本和格式"""
    # 获取原始文本的格式边界
    original_text = ''.join([text for text, _ in runs_data])
    
    # 如果原文和译文长度相近，按比例分配
    if 0.5 <= len(translated_text) / len(original_text) <= 2.0:
        _apply_proportional_format(paragraph, translated_text, runs_data)
    else:
        # 否则，尝试识别格式模式
        _apply_pattern_based_format(paragraph, translated_text, runs_data)

def _apply_proportional_format(paragraph, translated_text, runs_data):
    """按比例应用格式"""
    original_lengths = [len(text) for text, _ in runs_data]
    total_length = sum(original_lengths)
    translated_length = len(translated_text)
    
    if total_length == 0:
        return
    
    current_pos = 0
    for i, (original_text, format_info) in enumerate(runs_data):
        if original_text:  # 只处理非空的run
            # 计算这个run应该占的长度
            if i < len(runs_data) - 1:
                run_length = int(translated_length * len(original_text) / total_length)
            else:
                # 最后一个run获取剩余的所有文本
                run_length = translated_length - current_pos
            
            if current_pos < translated_length and run_length > 0:
                run_text = translated_text[current_pos:current_pos + run_length]
                if run_text:
                    run = paragraph.add_run(run_text)
                    _apply_run_format(run, format_info)
                current_pos += run_length

def _apply_pattern_based_format(paragraph, translated_text, runs_data):
    """基于模式应用格式"""
    # 识别原文中的格式模式
    format_patterns = []
    current_format = None
    current_text = ""
    
    for text, format_info in runs_data:
        if text:  # 只处理非空文本
            if format_info != current_format:
                if current_text:
                    format_patterns.append((current_text, current_format))
                current_format = format_info
                current_text = text
            else:
                current_text += text
    
    if current_text:
        format_patterns.append((current_text, current_format))
    
    # 如果只有一种格式，直接应用
    if len(format_patterns) == 1:
        run = paragraph.add_run(translated_text)
        _apply_run_format(run, format_patterns[0][1])
    else:
        # 尝试识别格式边界（如标题、强调等）
        # 这里简化处理：将第一部分作为可能的标题或强调
        if len(format_patterns) >= 2:
            first_text, first_format = format_patterns[0]
            first_ratio = len(first_text) / len(''.join([t for t, _ in format_patterns]))
            
            # 应用第一部分的格式
            first_length = int(len(translated_text) * first_ratio)
            if first_length > 0:
                run = paragraph.add_run(translated_text[:first_length])
                _apply_run_format(run, first_format)
                
                # 应用剩余部分的格式（使用最常见的格式）
                if first_length < len(translated_text):
                    run = paragraph.add_run(translated_text[first_length:])
                    # 使用第二种格式或默认格式
                    if len(format_patterns) > 1:
                        _apply_run_format(run, format_patterns[1][1])
        else:
            # 默认情况：使用第一个格式
            run = paragraph.add_run(translated_text)
            if runs_data:
                _apply_run_format(run, runs_data[0][1])

def _record_failed_task(failed_tasks, failed_tasks_lock, task_data, batch_id, failure_reason, error_msg="", is_serious=True):
    """记录失败的任务 - 修复版本"""
    task_id = f"task_{batch_id}_{task_data.get('index', 0)}"
    
    failed_task = FailedTask(
        original_text=task_data['text'],
        paragraph=task_data['paragraph'],
        paragraph_data=task_data,
        batch_id=batch_id,
        task_id=task_id,
        failure_reason=failure_reason,
        error_message=error_msg,
        is_serious=is_serious,
        retry_count=task_data.get('retry_count', 0)
    )
    
    with failed_tasks_lock:
        failed_tasks.append(failed_task)
        
    # 根据严重性更新统计
    if is_serious:
        return "serious"
    else:
        return "minor"

def _process_batch_enhanced(batch_data, failed_tasks, failed_tasks_lock, stats,
                          translation_timeout=DEFAULT_TRANSLATION_TIMEOUT) -> tuple:
    """处理单个批次 - 修复版本，支持更好的失败检测"""
    batch, batch_id, thread_id, translator, target_lang, source_lang, prompt_config = batch_data
    
    try:
        # 准备文本和任务索引
        texts = []
        for i, task_data in enumerate(batch):
            # 设置任务唯一ID
            task_data['index'] = i  # 添加批次内索引
            task_data['task_id'] = f"task_{batch_id}_{i}"  # 添加唯一任务ID
            texts.append(task_data['text'])
        
        logger.info(f"线程-{thread_id}: 处理批次-{batch_id} (共{len(texts)}段落)")
        
        batch_info = {'batch_id': batch_id, 'thread_id': thread_id}
        translations, failed_indices, error_message, cache_flags = _translate_batch_enhanced(
            texts, translator, target_lang, source_lang, prompt_config, batch_info,
            translation_timeout=translation_timeout
        )
        
        # 更新缓存命中统计
        stats['cache_hits'] += sum(cache_flags)
        stats['api_calls'] += 1
        
        # 应用成功的翻译结果
        success_count = 0
        serious_failures = 0
        minor_issues = 0
        
        for i, (task_data, translation) in enumerate(zip(batch, translations)):
            # 确保任务有唯一ID
            task_id = task_data.get('task_id', f"task_{batch_id}_{i}")
            task_data['task_id'] = task_id
            
            # 检查翻译结果
            if translation and translation.strip():
                # 验证翻译结果
                is_serious, reason = TranslationValidator.is_serious_failure(
                    task_data['text'], translation, large_text_threshold=50
                )
                
                if not is_serious:
                    try:
                        # 应用翻译到段落
                        paragraph = task_data['paragraph']
                        para_format = task_data['para_format']
                        runs_data = task_data['runs_data']
                        
                        # 清空段落
                        paragraph.clear()
                        
                        # 策略1：如果段落只有一种格式，直接应用
                        if len(runs_data) == 1 or _all_runs_same_format(runs_data):
                            run = paragraph.add_run(translation)
                            _apply_run_format(run, runs_data[0][1])
                        else:
                            # 策略2：尝试智能分配格式
                            _apply_translated_text_with_format(paragraph, translation, runs_data)
                        
                        # 恢复段落格式
                        _apply_paragraph_format(paragraph, para_format)
                        
                        success_count += 1
                    except Exception as e:
                        logger.warning(f"应用翻译失败: {e}")
                        failure_result = _record_failed_task(
                            failed_tasks, failed_tasks_lock, task_data, batch_id,
                            FailureReason.FORMAT_ERROR, f"格式应用错误: {str(e)}", True
                        )
                        serious_failures += 1 if failure_result == "serious" else 0
                        minor_issues += 1 if failure_result == "minor" else 0
                else:
                    # 严重的翻译问题 - 例如长文本未翻译
                    failure_result = _record_failed_task(
                        failed_tasks, failed_tasks_lock, task_data, batch_id,
                        FailureReason.NOT_TRANSLATED, reason, is_serious
                    )
                    serious_failures += 1 if failure_result == "serious" else 0
                    minor_issues += 1 if failure_result == "minor" else 0
                    
            elif i in failed_indices or not translation:
                # 明确的翻译失败
                if error_message and "timeout" in error_message.lower():
                    failure_reason = FailureReason.TIMEOUT
                    stats['timeout'] += 1
                else:
                    failure_reason = FailureReason.API_ERROR
                
                failure_result = _record_failed_task(
                    failed_tasks, failed_tasks_lock, task_data, batch_id,
                    failure_reason, error_message or "翻译失败或空结果", True
                )
                serious_failures += 1 if failure_result == "serious" else 0
                minor_issues += 1 if failure_result == "minor" else 0
        
        # 更新统计
        stats['success'] += success_count
        stats['failed'] += serious_failures + minor_issues
        stats['serious_failures'] += serious_failures
        stats['minor_issues'] += minor_issues
        
        logger.info(f"线程-{thread_id}: 批次-{batch_id} 完成，成功: {success_count}/{len(batch)}，严重失败: {serious_failures}，轻微问题: {minor_issues}")
        return len(batch), success_count, serious_failures + minor_issues, error_message, serious_failures
        
    except Exception as e:
        logger.error(f"线程-{thread_id}: 批次-{batch_id} 处理失败: {e}")
        # 记录整个批次失败
        serious_failures = 0
        for task_data in batch:
            failure_result = _record_failed_task(
                failed_tasks, failed_tasks_lock, task_data, batch_id,
                FailureReason.BATCH_FAILURE, f"批次处理异常: {str(e)}", True
            )
            serious_failures += 1 if failure_result == "serious" else 0
        
        stats['failed'] += len(batch)
        stats['serious_failures'] += serious_failures
        return len(batch), 0, len(batch), str(e), serious_failures

def _analyze_failure_patterns(failed_tasks) -> Dict[str, Any]:
    """分析失败模式"""
    analysis = {
        'total_count': len(failed_tasks),
        'avg_text_length': sum(len(task.original_text) for task in failed_tasks) / max(1, len(failed_tasks)),
        'failure_types': {},
        'has_timeouts': False,
        'has_api_errors': False,
        'consecutive_failures': 0
    }
    
    for task in failed_tasks:
        reason = task.failure_reason.value
        analysis['failure_types'][reason] = analysis['failure_types'].get(reason, 0) + 1
        
        if task.failure_reason == FailureReason.TIMEOUT:
            analysis['has_timeouts'] = True
        elif task.failure_reason == FailureReason.API_ERROR:
            analysis['has_api_errors'] = True
    
    return analysis

def _determine_retry_strategy(analysis: Dict[str, Any], retry_count: int) -> Dict[str, Any]:
    """确定重试策略 - 简化并行判断"""
    strategy = {
        'use_concurrent': True,  # 默认使用并行
        'concurrent_threshold': 2,  # 简化阈值
        'max_workers': DEFAULT_RETRY_THREADS,  # 默认重试线程数
        'batch_size_multiplier': 1.0,
        'timeout_multiplier': 1.0,
        'description': '并行重试'
    }
    
    # 如果有超时，降低并发度但仍使用并行
    if analysis['has_timeouts']:
        strategy['max_workers'] = max(2, strategy['max_workers'] // 2)
        strategy['timeout_multiplier'] = 1.5
        strategy['description'] = '超时优化并行重试'
    
    # 如果有API错误，使用更小批次但仍并行
    if analysis['has_api_errors']:
        strategy['batch_size_multiplier'] = 0.5
        strategy['description'] = 'API错误优化并行重试'
    
    # 大文本处理 - 仍使用并行
    if analysis['avg_text_length'] > 1000:
        strategy['concurrent_threshold'] = 1  # 任何数量都并行
        strategy['timeout_multiplier'] = 2.0
        strategy['description'] = '长文本优化并行重试'
    
    # 只有在worker数量为1时才串行
    if strategy['max_workers'] == 1:
        strategy['use_concurrent'] = False
        strategy['description'] = '串行重试'
    
    return strategy

def _create_smart_retry_batches(failed_tasks: List[FailedTask], retry_count: int, retry_batch_sizes: List[int]) -> List[List[FailedTask]]:
    """创建智能重试批次，只处理严重失败"""
    # 只重试严重失败
    serious_failed_tasks = [task for task in failed_tasks if task.is_serious]
    
    if not serious_failed_tasks:
        return []
    
    # 根据重试次数确定批次大小
    if retry_count < len(retry_batch_sizes):
        max_batch_size = retry_batch_sizes[retry_count]
    else:
        # 后续重试使用最小批次
        max_batch_size = 1
    
    logger.info(f"第{retry_count + 1}次重试，严重失败任务: {len(serious_failed_tasks)}, 批次大小: {max_batch_size}")
    
    # 按失败原因分组
    tasks_by_reason = {}
    for task in serious_failed_tasks:
        reason = task.failure_reason
        if reason not in tasks_by_reason:
            tasks_by_reason[reason] = []
        tasks_by_reason[reason].append(task)
    
    chunks = []
    
    for reason, tasks in tasks_by_reason.items():
        # 根据失败类型确定批次策略
        if reason == FailureReason.TIMEOUT:
            # 超时失败：减小批次
            batch_size = max(1, max_batch_size // 2)
        elif reason == FailureReason.API_ERROR:
            # API错误：更小批次
            batch_size = max(1, max_batch_size // 3)
        else:
            batch_size = max_batch_size
        
        # 按文本长度排序，长文本优先单独处理
        tasks.sort(key=lambda x: len(x.original_text), reverse=True)
        
        current_chunk = []
        current_chars = 0
        
        for task in tasks:
            text_len = len(task.original_text)
            
            # 超长文本单独处理
            if text_len > MAX_CHARS_PER_BATCH // 2:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = []
                    current_chars = 0
                chunks.append([task])
                continue
            
            # 正常批次处理
            if (len(current_chunk) >= batch_size or 
                current_chars + text_len > MAX_CHARS_PER_BATCH) and current_chunk:
                chunks.append(current_chunk)
                current_chunk = [task]
                current_chars = text_len
            else:
                current_chunk.append(task)
                current_chars += text_len
        
        if current_chunk:
            chunks.append(current_chunk)
    
    return chunks

def _calculate_adaptive_delay(retry_count: int, retry_delays: List[float], failure_analysis: Dict[str, Any]) -> float:
    """计算自适应延迟"""
    base_delay = retry_delays[min(retry_count, len(retry_delays) - 1)]
    
    # 根据失败模式调整延迟
    if failure_analysis.get('has_api_errors', False):
        base_delay *= 1.5  # API错误需要更长延迟
    
    if failure_analysis.get('avg_text_length', 0) > 1000:
        base_delay *= 1.2  # 长文本需要更多时间
    
    return base_delay

def _execute_concurrent_retry(retry_batches: List[List[FailedTask]], 
                            translator, target_lang: str, source_lang: Optional[str], 
                            retry_count: int, strategy: Dict[str, Any], 
                            prompt_config: Optional[Dict[str, Any]],
                            translation_timeout: int,
                            stats: Dict[str, Any]) -> int:
    """并发重试执行 - 修复版本，支持结果更新"""
    success_count = 0
    max_workers = strategy['max_workers']
    timeout_multiplier = strategy.get('timeout_multiplier', 1.0)
    actual_timeout = translation_timeout * timeout_multiplier
    
    try:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            with tqdm(total=sum(len(batch) for batch in retry_batches), 
                     desc=f"第{retry_count + 1}次重试(并发)", unit="任务") as pbar:
                
                # 提交所有重试批次任务
                future_to_batch = {}
                for retry_batch_idx, retry_batch in enumerate(retry_batches):
                    future = executor.submit(
                        _process_retry_batch, retry_batch, retry_batch_idx, retry_count,
                        translator, target_lang, source_lang, prompt_config, actual_timeout
                    )
                    future_to_batch[future] = retry_batch
                
                # 处理完成的任务
                for future in as_completed(future_to_batch):
                    retry_batch = future_to_batch[future]
                    try:
                        batch_success = future.result(timeout=actual_timeout + 30)
                        success_count += batch_success
                        pbar.update(len(retry_batch))
                    except Exception as e:
                        logger.error(f"并发重试批次处理异常: {e}")
                        # 对失败的批次任务增加重试计数
                        for task in retry_batch:
                            task.retry_count += 1
                        pbar.update(len(retry_batch))
        
        stats['parallel_retry_used'] += 1
        return success_count
        
    except Exception as e:
        logger.error(f"并发重试执行异常: {e}")
        # 对所有任务增加重试计数
        for batch in retry_batches:
            for task in batch:
                task.retry_count += 1
        return success_count

def _execute_sequential_retry(retry_batches: List[List[FailedTask]], 
                            translator, target_lang: str, source_lang: Optional[str], 
                            retry_count: int, strategy: Dict[str, Any], 
                            prompt_config: Optional[Dict[str, Any]],
                            translation_timeout: int,
                            stats: Dict[str, Any]) -> int:
    """串行重试执行 - 修复版本，支持结果更新"""
    success_count = 0
    timeout_multiplier = strategy.get('timeout_multiplier', 1.0)
    actual_timeout = translation_timeout * timeout_multiplier
    
    with tqdm(total=sum(len(batch) for batch in retry_batches), 
             desc=f"第{retry_count + 1}次重试", unit="任务") as pbar:
        
        for retry_batch_idx, retry_batch in enumerate(retry_batches):
            batch_success = _process_retry_batch(
                retry_batch, retry_batch_idx, retry_count,
                translator, target_lang, source_lang, prompt_config, actual_timeout
            )
            success_count += batch_success
            pbar.update(len(retry_batch))
    
    return success_count

def _process_retry_batch(retry_batch: List[FailedTask], batch_idx: int, retry_count: int,
                       translator, target_lang: str, source_lang: Optional[str], 
                       prompt_config: Optional[Dict[str, Any]], translation_timeout: int) -> int:
    """处理单个重试批次 - 核心修复：直接更新原始段落"""
    logger.info(f"开始重试批次 {batch_idx + 1}，任务数: {len(retry_batch)}")
    
    # 提取文本和任务
    batch_texts = [task.original_text for task in retry_batch]
    
    # 执行重试翻译
    translations, failed_indices, error_message, cache_flags = _translate_batch_enhanced(
        batch_texts, translator, target_lang, source_lang, 
        prompt_config, translation_timeout=translation_timeout
    )
    
    logger.info(f"重试批次 {batch_idx + 1} 完成，结果数: {len(translations) if translations else 0}")
    
    # 处理重试结果 - 关键修复：直接更新原始段落
    success_count = 0
    
    if not translations:
        # 整个批次失败
        batch_reason = error_message if error_message else "重试批次失败"
        logger.warning(f"重试批次失败: {batch_reason}")
        
        # 更新重试计数
        for task in retry_batch:
            task.retry_count += 1
            task.error_message = batch_reason
        
        return 0
    
    # 批次成功，检查个别结果并更新原始段落
    for i, task in enumerate(retry_batch):
        if i < len(translations):
            translation = translations[i]
            from_cache = cache_flags[i] if i < len(cache_flags) else False
            
            # 检查重试是否成功
            is_serious, reason = TranslationValidator.is_serious_failure(task.original_text, translation)
            
            if not is_serious and not TranslationValidator.is_error_message(translation):
                # 重试成功 - 核心修复：直接更新原始段落
                try:
                    paragraph = task.paragraph
                    paragraph_data = task.paragraph_data
                    para_format = paragraph_data['para_format']
                    runs_data = paragraph_data['runs_data']
                    
                    # 清空段落
                    paragraph.clear()
                    
                    # 根据格式应用翻译
                    if len(runs_data) == 1 or _all_runs_same_format(runs_data):
                        run = paragraph.add_run(translation)
                        _apply_run_format(run, runs_data[0][1])
                    else:
                        _apply_translated_text_with_format(paragraph, translation, runs_data)
                    
                    # 恢复段落格式
                    _apply_paragraph_format(paragraph, para_format)
                    
                    # 更新统计和状态
                    success_count += 1
                    task.retry_count += 1  # 仍然增加重试计数，但标记为成功
                    task.is_serious = False  # 不再是严重问题
                    logger.info(f"✅ 重试救援成功: {task.original_text[:30]}... -> {translation[:30]}...")
                    
                except Exception as e:
                    logger.warning(f"应用重试翻译失败: {e}")
                    task.retry_count += 1
                    task.error_message = f"格式应用错误: {str(e)}"
            else:
                # 重试仍然失败
                logger.debug(f"重试失败: {task.original_text[:50]}... 原因: {reason}")
                task.retry_count += 1
                task.error_message = reason or "重试翻译结果仍有问题"
        else:
            # 没有对应的结果
            logger.warning(f"重试结果不足: 任务 {i}, 结果数量 {len(translations)}")
            task.retry_count += 1
            task.error_message = "重试结果缺失"
    
    return success_count

def _adaptive_retry_strategy(failed_tasks: List[FailedTask], translator, target_lang: str, 
                           source_lang: Optional[str], prompt_config: Optional[Dict[str, Any]],
                           translation_timeout: int, max_retries: int, retry_batch_sizes: List[int],
                           retry_delays: List[float], stats: Dict[str, Any]) -> int:
    """自适应重试策略主控制器 - 修复版本，支持结果更新"""
    retry_count = 0
    retry_start_time = time.time()  # 开始重试计时
    total_success_count = 0
    
    # 只处理严重失败任务
    serious_tasks = [task for task in failed_tasks if task.is_serious]
    if not serious_tasks:
        logger.info("没有严重失败任务需要重试")
        return 0
    
    logger.info(f"开始自适应重试，发现 {len(serious_tasks)} 个严重失败任务")
    stats['total_retry_tasks'] = len(serious_tasks)
    
    while retry_count < max_retries:
        # 过滤出当前需要重试的任务
        current_retry_tasks = [task for task in failed_tasks 
                             if task.is_serious and task.retry_count <= retry_count]
        
        if not current_retry_tasks:
            logger.info(f"第 {retry_count + 1} 次重试检查：没有(更多)需要重试的任务")
            break
        
        logger.info(f"第 {retry_count + 1} 次重试，处理 {len(current_retry_tasks)} 个严重失败任务")
        
        # 分析失败模式
        failure_analysis = _analyze_failure_patterns(current_retry_tasks)
        
        # 动态调整策略
        strategy = _determine_retry_strategy(failure_analysis, retry_count)
        
        # 添加重试延迟
        delay = _calculate_adaptive_delay(retry_count, retry_delays, failure_analysis)
        if delay > 0:
            logger.info(f"智能延迟: {delay:.1f}秒")
            time.sleep(delay)
        
        # 创建重试批次
        retry_batches = _create_smart_retry_batches(current_retry_tasks, retry_count, retry_batch_sizes)
        
        if not retry_batches:
            logger.info(f"第 {retry_count + 1} 次重试：没有批次需要处理")
            break
        
        # 简化并行重试判断 - 优先使用并行
        use_concurrent = len(retry_batches) >= strategy['concurrent_threshold']
        actual_workers = min(strategy['max_workers'], len(retry_batches)) if use_concurrent else 1
        
        # 执行重试
        if use_concurrent:
            logger.info(f"启用并发重试，worker数量: {actual_workers}")
            success_count = _execute_concurrent_retry(
                retry_batches, translator, target_lang, source_lang, 
                retry_count, strategy, prompt_config, translation_timeout, stats
            )
        else:
            logger.info("使用串行重试")
            success_count = _execute_sequential_retry(
                retry_batches, translator, target_lang, source_lang, 
                retry_count, strategy, prompt_config, translation_timeout, stats
            )
        
        total_success_count += success_count
        stats['rescued_tasks'] += success_count
        stats['retry_attempts'] += 1
        
        # 更新失败任务列表 - 移除不再严重的任务
        failed_tasks[:] = [task for task in failed_tasks if task.is_serious]
        
        logger.info(f"第 {retry_count + 1} 次重试完成，恢复 {success_count} 个任务，当前剩余 {len(failed_tasks)} 个严重失败")
        
        retry_count += 1
    
    # 计算重试时间和成功率
    retry_time = time.time() - retry_start_time
    if stats['total_retry_tasks'] > 0:
        stats['retry_success_rate'] = stats['rescued_tasks'] / stats['total_retry_tasks']
    
    logger.info(f"重试阶段完成，用时: {retry_time:.1f}秒, "
               f"成功率: {stats['retry_success_rate']:.1%}")
    
    return total_success_count

def translate_docx_file_formatted(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    chunk_size: int = OPTIMAL_BATCH_SIZE,
    prompt_config=None,  # prompt配置参数
    config_merge_mode: Optional[Union[str, ConfigMergeStrategy]] = None,  # 配置合并模式
    translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT,
    max_retries: int = DEFAULT_MAX_RETRIES,
    retry_delay: int = DEFAULT_RETRY_DELAY,
    retry_batch_size: int = DEFAULT_RETRY_BATCH_SIZE,
    max_workers: int = DEFAULT_THREADS,
    max_retry_workers: int = DEFAULT_RETRY_THREADS,
    large_text_threshold: int = 50,
    retry_failure_threshold: float = 0.0,
    **kwargs
) -> str:
    """翻译Word文档并保持格式，支持重试机制 - 修复版本，解决重试结果更新问题"""
    try:
        # 开始总计时
        total_start_time = time.time()
        
        # 检查输入文件
        if not os.path.exists(input_filepath):
            return f"Error: 输入文件未找到: {input_filepath}"
            
        # 准备输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成输出文件路径
        filename = os.path.basename(input_filepath)
        name, ext = os.path.splitext(filename)
        
        # 支持自定义文件名
        if unique_filename_base:
            output_filename = f"{unique_filename_base}_{target_lang}{ext}"
        else:
            output_filename = f"{name}_translated_{target_lang}{ext}"
        
        output_filepath = os.path.join(output_dir, output_filename)

        # 统计信息
        stats = {
            'success': 0, 
            'failed': 0, 
            'skipped': 0, 
            'timeout': 0, 
            'retried': 0,
            'final_failed': 0,
            'serious_failures': 0,
            'minor_issues': 0,
            'retry_attempts': 0,
            'rescued_tasks': 0,
            'retry_success_rate': 0.0,
            'total_retry_tasks': 0,
            'parallel_retry_used': 0,
            'cache_hits': 0,
            'api_calls': 0
        }
        
        # 时间统计
        time_stats = {
            'total_start_time': total_start_time,
            'main_translation_time': 0,
            'retry_time': 0,
            'total_time': 0,
            'document_load_time': 0,
            'document_save_time': 0
        }
        
        # 失败任务追踪
        failed_tasks = []
        failed_tasks_lock = threading.Lock()

        # 增强重试策略配置
        retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]
        retry_delays = [1, 2, 4, 8, 12, 16, 20, 25]

        # 配置合并和翻译器配置管理
        original_prompt_config = None
        try:
            # 获取有效配置（应用配置合并策略）
            effective_config = get_effective_config(prompt_config, config_merge_mode)
            
            # 如果提供了有效配置，应用到翻译器
            if effective_config and hasattr(translator, 'set_prompt_config'):
                # 保存翻译器的原始prompt配置
                original_prompt_config = getattr(translator, 'prompt_config', None)
                translator.set_prompt_config(effective_config)
                logger.info(f"Applied effective config to translator for formatted DOCX translation: mode={effective_config.get('mode')}, strategy={config_merge_mode or _global_merge_strategy.value}")
            
            # 从有效配置中提取批处理参数
            if effective_config:
                if 'chunk_size' in effective_config:
                    chunk_size = effective_config['chunk_size']
                elif 'batch_size' in effective_config:
                    chunk_size = effective_config['batch_size']
            
            # 处理prompt相关的kwargs参数，保持向后兼容
            if effective_config:
                # 这些参数可能从kwargs传递，但有效配置优先
                preserve_terms = effective_config.get('preserve_terms', kwargs.get('preserve_terms'))
                glossary = effective_config.get('glossary', kwargs.get('glossary'))
                additional_context = effective_config.get('additional_context', kwargs.get('additional_context'))
                
                # 更新有效配置
                if preserve_terms and 'preserve_terms' not in effective_config:
                    effective_config['preserve_terms'] = preserve_terms
                if glossary and 'glossary' not in effective_config:
                    effective_config['glossary'] = glossary
                if additional_context and 'additional_context' not in effective_config:
                    effective_config['additional_context'] = additional_context
                
                logger.info(f"Processing formatted DOCX with effective config: mode={effective_config.get('mode')}, strategy={config_merge_mode or _global_merge_strategy.value}")
            
            # 清理缓存
            with _cache_lock:
                _translation_cache.clear()
            
            # 加载文档
            doc_load_start = time.time()
            doc = Document(input_filepath)
            time_stats['document_load_time'] = time.time() - doc_load_start
            
            # 收集所有需要翻译的段落
            paragraphs_to_translate = []
            paragraph_data = []
            skipped_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 保存段落格式
                    para_format = _get_paragraph_format(paragraph)
                    
                    # 保存每个run的文本和格式
                    runs_data = []
                    for run in paragraph.runs:
                        # 保存所有run，包括空的（用于保持格式）
                        format_info = _get_run_format(run)
                        runs_data.append((run.text, format_info))
                    
                    if runs_data:
                        # 只翻译有实际内容的文本
                        full_text = paragraph.text
                        if full_text.strip():  # 确保有内容需要翻译
                            task_data = {
                                'text': full_text,
                                'paragraph': paragraph,
                                'para_format': para_format,
                                'runs_data': runs_data,
                                'retry_count': 0  # 初始重试计数
                            }
                            paragraphs_to_translate.append(task_data)
                else:
                    stats['skipped'] += 1
            
            logger.info(f"需要翻译 {len(paragraphs_to_translate)} 段落，跳过 {stats['skipped']} 段落")
            logger.info(f"配置合并策略: {config_merge_mode or _global_merge_strategy.value}")
            if effective_config:
                logger.info(f"使用有效prompt模式: {effective_config.get('mode', 'default')}")
            
            if not paragraphs_to_translate:
                logger.info("没有需要翻译的内容")
                # 直接复制原文件
                import shutil
                shutil.copy2(input_filepath, output_filepath)
                print("翻译完成！（无需要翻译的内容）")
                time_stats['total_time'] = time.time() - time_stats['total_start_time']
                return output_filepath
            
            # 智能分批
            actual_batch_size = min(chunk_size, OPTIMAL_BATCH_SIZE)
            batches = _create_smart_batches([task['text'] for task in paragraphs_to_translate], max_batch_size=actual_batch_size)
            
            # 重建batch数据结构
            batch_data_list = []
            task_idx = 0
            for batch_id, text_batch in enumerate(batches):
                batch_tasks = []
                for _ in text_batch:
                    if task_idx < len(paragraphs_to_translate):
                        batch_tasks.append(paragraphs_to_translate[task_idx])
                        task_idx += 1
                batch_data_list.append(batch_tasks)
            
            logger.info(f"分为 {len(batch_data_list)} 个批次处理")
            
            # === 第一阶段：多线程批量翻译 ===
            logger.info("=== 阶段1：初始批量翻译 ===")
            total_paragraphs = len(paragraphs_to_translate)
            main_translation_start = time.time()
            
            with ThreadPoolExecutor(max_workers=min(max_workers, len(batch_data_list))) as executor:
                # 准备批次数据
                future_batch_data = []
                for i, batch_tasks in enumerate(batch_data_list):
                    batch_data = (batch_tasks, i+1, i+1, translator, target_lang, source_lang, effective_config)
                    future_batch_data.append(batch_data)
                
                # 提交任务
                future_to_batch = {
                    executor.submit(_process_batch_enhanced, batch_data, failed_tasks, failed_tasks_lock, stats, translation_timeout): batch_data[0]
                    for batch_data in future_batch_data
                }
                
                # 处理结果
                with tqdm(total=total_paragraphs, desc="翻译进度", unit="段落") as pbar:
                    for future in as_completed(future_to_batch):
                        total, success, failed, error, serious_failures = future.result()
                        pbar.update(total)
                        
                        if error:
                            logger.warning(f"批次错误: {error}")
            
            time_stats['main_translation_time'] = time.time() - main_translation_start
            logger.info(f"主翻译完成，耗时: {time_stats['main_translation_time']:.1f}秒")
            
            # === 第二阶段：重试失败的任务 ===
            serious_failures = len([task for task in failed_tasks if task.is_serious])
            if serious_failures > 0:
                logger.info(f"=== 阶段2：重试 {serious_failures} 个严重失败任务 ===")
                retry_start_time = time.time()
                
                # 执行自适应重试策略
                retry_success_count = _adaptive_retry_strategy(
                    failed_tasks, translator, target_lang, source_lang, effective_config,
                    translation_timeout, max_retries, retry_batch_sizes, retry_delays, stats
                )
                
                time_stats['retry_time'] = time.time() - retry_start_time
                logger.info(f"重试阶段完成，恢复 {retry_success_count} 个任务，耗时: {time_stats['retry_time']:.1f}秒")
                
                # 更新统计
                stats['success'] += retry_success_count
                stats['failed'] -= retry_success_count
                stats['retried'] += retry_success_count
                
                # 统计最终失败任务
                final_failures = len([task for task in failed_tasks if task.is_serious])
                stats['final_failed'] = final_failures
                
                if final_failures > 0:
                    logger.warning(f"重试后仍有 {final_failures} 个严重失败任务保持原文")
            else:
                logger.info("=== 阶段2：没有需要重试的严重失败任务 ===")
            
            # 保存文档
            save_start = time.time()
            doc.save(output_filepath)
            time_stats['document_save_time'] = time.time() - save_start
            
            # 计算总时间
            time_stats['total_time'] = time.time() - time_stats['total_start_time']
            
            # 输出详细统计
            total = len(paragraphs_to_translate)
            success_rate = (stats['success'] / total * 100) if total > 0 else 0
            
            logger.info(f"翻译完成: 成功率 {success_rate:.1f}%")
            logger.info(f"总耗时: {time_stats['total_time']:.1f}秒 (主翻译: {time_stats['main_translation_time']:.1f}s, "
                       f"重试: {time_stats['retry_time']:.1f}s, 文档加载: {time_stats['document_load_time']:.1f}s, "
                       f"文档保存: {time_stats['document_save_time']:.1f}s)")
            
            print(f"\n✅ 格式化翻译完成!")
            print(f"📄 输出文件: {output_filepath}")
            print(f"📊 翻译统计:")
            print(f"  ✨ 成功翻译: {stats['success']} 段落")
            print(f"  🔄 重试恢复: {stats['retried']} 段落")
            print(f"  ⏰ 超时失败: {stats['timeout']} 段落")
            print(f"  ❌ 最终失败: {stats['final_failed']} 段落")
            print(f"  ⏭️ 跳过: {stats['skipped']} 段落")
            print(f"  🎯 成功率: {success_rate:.1f}%")
            print(f"  📋 格式保持: ✅ 完整")
            print(f"  🔧 配置合并策略: {config_merge_mode or _global_merge_strategy.value}")
            print(f"  ⏱️ 总用时: {time_stats['total_time']:.1f}秒")
            
            if stats['retry_attempts'] > 0:
                print(f"  🔁 重试统计: {stats['retry_attempts']}轮，成功率{stats['retry_success_rate']:.1%}")
            
            if stats['final_failed'] > 0:
                print(f"  ⚠️ 注意: {stats['final_failed']} 段落保持原文")
            
            if effective_config:
                print(f"  🎨 Prompt模式: {effective_config.get('mode', 'default')}")
                if effective_config.get('mode') == 'professional':
                    print(f"  🏢 专业领域: {effective_config.get('prompt_template', 'default')}")
            
            return output_filepath
            
        except Exception as e:
            logger.exception(f"文档翻译错误: {e}")
            return f"Error: {str(e)}"
        
        finally:
            # 恢复翻译器的原始prompt配置
            if effective_config and hasattr(translator, 'set_prompt_config'):
                if original_prompt_config is not None:
                    translator.set_prompt_config(original_prompt_config)
                    logger.debug("Restored original prompt config to translator")
                else:
                    # 如果原来没有配置，清除当前配置
                    if hasattr(translator, 'prompt_config'):
                        translator.prompt_config = None
                        logger.debug("Cleared prompt config from translator")
        
    except Exception as e:
        logger.exception(f"文档翻译错误: {e}")
        return f"Error: {str(e)}"

# 测试用的模拟翻译器，支持配置合并策略
class MockTranslator:
    def __init__(self):
        self.prompt_config = None
    
    def set_prompt_config(self, prompt_config):
        """设置prompt配置"""
        self.prompt_config = prompt_config
    
    def translate(self, text=None, messages=None, target_lang: str = None, source_lang: Optional[str] = None, 
                 prompt_config=None, config_merge_mode='merge', **kwargs) -> str:
        import time
        import random
        
        # 模拟随机失败和超时
        if random.random() < 0.1:  # 10%失败率
            if random.random() < 0.5:
                time.sleep(70)  # 模拟超时
            else:
                raise Exception("Mock translation failure")
        
        time.sleep(random.uniform(0.1, 0.3))  # 模拟网络延迟
        
        # 支持多种调用方式
        if isinstance(messages, list) and len(messages) > 0:
            # messages格式
            content = messages[-1]["content"]
        elif isinstance(text, str):
            # 文本格式
            content = text
        elif isinstance(messages, str):
            # 传统格式
            content = messages
        else:
            return "Error: Invalid input format"
        
        # 使用配置合并策略
        effective_config = get_effective_config(prompt_config, config_merge_mode)
        
        # 根据有效配置调整翻译行为
        if effective_config:
            mode = effective_config.get('mode', 'none')
            if mode == 'custom':
                prefix = "[CUSTOM]"
            elif mode == 'professional':
                domain = effective_config.get('prompt_template', 'academic')
                prefix = f"[{domain.upper()}]"
            else:
                prefix = f"[{target_lang}]"
            
            # 显示合并后的特殊功能
            features = []
            if effective_config.get('preserve_terms'):
                features.append(f"TERMS:{len(effective_config['preserve_terms'])}")
            if effective_config.get('glossary'):
                features.append(f"GLOSS:{len(effective_config['glossary'])}")
            
            if features:
                prefix += f"[{','.join(features)}]"
        else:
            prefix = f"[{target_lang}]"
        
        # 模拟返回带行号的翻译结果
        lines = content.split('\n')
        translated_lines = []
        for line in lines:
            if line.strip().startswith('[') and ']' in line:
                # 提取行号和内容
                match = re.match(r'^\[(\d+)\]\s*(.*)', line.strip())
                if match:
                    num = match.group(1)
                    content_text = match.group(2)
                    translated_lines.append(f"{prefix} {content_text}")
            else:
                # 对于没有行号的内容，直接翻译
                if line.strip():
                    translated_lines.append(f"{prefix} {line.strip()}")
        
        return '\n'.join(translated_lines)

if __name__ == '__main__':
    # 强制设置日志配置
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        force=True  # 强制重新配置
    )
    
    print("=== 增强格式化DOCX翻译器（修复重试结果更新版）===")
    print("🚀 新增特性：")
    print("✅ 直接更新重试结果 - 不再存在'很多没有翻译'问题")
    print("✅ 优先使用并行重试 - 大幅提高重试效率")
    print("✅ 增强失败检测 - 减少误判，提高成功率")
    print("✅ 完整时间统计 - 主翻译/重试/加载/保存各阶段耗时")
    print("✅ 智能批次管理 - 根据失败原因调整批次大小")
    print("✅ 自适应延迟 - 智能调整重试间隔")
    print("✅ 全流程完整保持格式 - 从加载到保存精确保持格式")
    
    # 演示配置
    print(f"\n=== 重试机制配置 ===")
    print(f"📦 默认批次大小: {OPTIMAL_BATCH_SIZE}")
    print(f"🔄 重试批次大小: {DEFAULT_RETRY_BATCH_SIZE}")
    print(f"⏰ 翻译超时: {DEFAULT_TRANSLATION_TIMEOUT}秒")
    print(f"🔁 最大重试次数: {DEFAULT_MAX_RETRIES}")
    print(f"⏸️ 重试延迟: {DEFAULT_RETRY_DELAY}秒")
    print(f"🧵 默认线程数: {DEFAULT_THREADS}")
    print(f"🧵 重试线程数: {DEFAULT_RETRY_THREADS}")
    print(f"\n=== 核心修复 ===")
    print("1. 失败任务记录增强 - 保存段落直接引用")
    print("2. 重试结果直接更新 - 解决重试成功但结果未更新问题")
    print("3. 并行重试机制 - 智能分批，避免API限制")
    print("4. 完整时间追踪 - 精确统计各阶段时间")
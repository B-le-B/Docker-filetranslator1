# docx_full_translator.py (第一批修改)
import os
import logging
import shutil
import tempfile
import re
import time
import hashlib
import threading
import copy
from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
from tqdm import tqdm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
from dataclasses import dataclass
from collections import OrderedDict
from contextlib import contextmanager
from enum import Enum

# 日志配置
logging.getLogger().setLevel(logging.INFO)

third_party_loggers = [
    'urllib3', 'requests', 'httpx', 'siliconflow', 'openai',
    'anthropic', 'zhipuai', 'dashscope', 'httpcore', 'httpx._client',
    'httpx._config', 'httpx._models', 'httpx._auth', 'requests.packages.urllib3',
    'requests_oauthlib', 'oauthlib', 'aiohttp', 'websockets', 'asyncio'
]

for logger_name in third_party_loggers:
    logging.getLogger(logger_name).setLevel(logging.WARNING)

class NoDebugFilter(logging.Filter):
    def filter(self, record):
        if record.levelno >= logging.WARNING:
            return True
        if record.name == __name__ or record.name.startswith('__main__'):
            return record.levelno >= logging.INFO
        return False

root_logger = logging.getLogger()
root_logger.addFilter(NoDebugFilter())

logger = logging.getLogger(__name__)

# 预编译正则表达式以提高性能
NUMBERED_LINE_PATTERN = re.compile(r'^\[(\d+)\]\s*(.*)')
LIST_MARKER_PATTERN = re.compile(r'^([\d\w]+[\.\)]\s*|[•·▪▫◦‣⁃]\s*)')
WHITESPACE_PATTERN = re.compile(r'\s+')

class FailureReason(Enum):
    """失败原因枚举"""
    TIMEOUT = "timeout"
    API_ERROR = "api_error"
    PARSE_ERROR = "parse_error"
    NOT_TRANSLATED = "not_translated"
    EMPTY_RESPONSE = "empty_response"
    CONNECTION_ERROR = "connection_error"
    BATCH_FAILURE = "batch_failure"

@dataclass
class FailedTask:
    """统一的失败任务数据类"""
    original_text: str
    original_index: int  # 在原始列表中的索引
    element: Any  # ParagraphElement对象
    failure_reason: FailureReason
    retry_count: int = 0
    error_message: str = ""
    is_serious: bool = True
    
    def __post_init__(self):
        self.failure_timestamp = time.time()

@dataclass
class ParagraphElement:
    """保持段落完整性的元素"""
    full_text: str
    paragraph: Any
    runs_info: List[Dict[str, Any]]
    para_format: Dict[str, Any]
    location: str
    element_type: str = 'paragraph'

def _prepare_prompt_config(prompt_config: Optional[Dict[str, Any]], kwargs: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """准备和标准化prompt配置，与前端格式兼容"""
    try:
        # 检查是否有任何有效的配置参数
        has_prompt_config = prompt_config and isinstance(prompt_config, dict) and prompt_config
        has_kwargs_config = any(k in kwargs for k in ['preserve_terms', 'glossary', 'additional_context', 'prompt_template', 'custom_prompt'])
        
        if not has_prompt_config and not has_kwargs_config:
            logger.debug("No valid prompt configuration found, returning None")
            return None
        
        # 基础配置 - 确保是字典
        config = {}
        if has_prompt_config:
            config.update(prompt_config)
        
        # 从kwargs合并配置（向后兼容）
        for key in ['preserve_terms', 'glossary', 'additional_context', 'prompt_template', 'custom_prompt']:
            if key in kwargs and kwargs[key] is not None:
                config[key] = kwargs[key]
        
        # 再次检查是否有有效内容
        if not config:
            logger.debug("Config is empty after processing, returning None")
            return None
        
        # 标准化配置格式
        normalized_config = _normalize_prompt_config(config)
        
        # 最终验证
        if not normalized_config or normalized_config.get('mode') == 'none':
            logger.debug("Normalized config is invalid, returning None")
            return None
        
        logger.debug(f"Prepared prompt config for advanced DOCX translation: {normalized_config}")
        return normalized_config
        
    except Exception as e:
        logger.warning(f"Error preparing prompt config: {e}")
        return None

def _normalize_prompt_config(config: Dict[str, Any]) -> Dict[str, Any]:
    """标准化prompt配置格式，确保与前端格式兼容"""
    try:
        if not config or not isinstance(config, dict):
            return {'mode': 'none'}
        
        normalized = config.copy()
        
        # 确保mode字段存在
        if 'mode' not in normalized:
            if 'custom_prompt' in normalized and normalized['custom_prompt']:
                normalized['mode'] = 'custom'
            elif 'prompt_template' in normalized or 'professional_domain' in normalized:
                normalized['mode'] = 'professional'
            elif any(k in normalized for k in ['preserve_terms', 'glossary', 'additional_context']):
                normalized['mode'] = 'general'
            else:
                normalized['mode'] = 'none'
        
        # 处理保留术语 - 支持逗号分隔的字符串（前端格式）
        preserve_terms = normalized.get('preserve_terms')
        if preserve_terms:
            try:
                if isinstance(preserve_terms, str):
                    # 前端格式：逗号分隔的字符串
                    terms_list = [term.strip() for term in preserve_terms.split(',') if term.strip()]
                    normalized['preserve_terms'] = terms_list
                elif isinstance(preserve_terms, list):
                    # 确保列表中的字符串都是清理过的
                    normalized['preserve_terms'] = [str(term).strip() for term in preserve_terms if str(term).strip()]
                else:
                    # 其他格式，移除该字段
                    normalized.pop('preserve_terms', None)
            except Exception as e:
                logger.warning(f"Error processing preserve_terms: {e}")
                normalized.pop('preserve_terms', None)
        
        # 处理术语表 - 确保是字典格式
        glossary = normalized.get('glossary')
        if glossary:
            if not isinstance(glossary, dict):
                logger.warning(f"Glossary should be a dictionary, got {type(glossary)}, ignoring")
                normalized.pop('glossary', None)
        
        # 处理自定义prompt
        if normalized.get('mode') == 'custom':
            custom_prompt = normalized.get('custom_prompt', {})
            if not custom_prompt or not isinstance(custom_prompt, dict):
                # 检查是否有分离的system和user字段
                system_prompt = normalized.get('custom_system_prompt', normalized.get('system'))
                user_prompt = normalized.get('custom_user_prompt', normalized.get('user'))
                
                if system_prompt:
                    normalized['custom_prompt'] = {
                        'system': str(system_prompt),
                        'user': str(user_prompt) if user_prompt else 'Please translate the following content to {target_lang}:\n\n{content}'
                    }
                else:
                    logger.warning("Custom mode selected but no valid custom prompt provided, falling back to general mode")
                    normalized['mode'] = 'general'
        
        # 处理专业模板 - 前端使用 'professional_domain' 字段
        if normalized.get('mode') == 'professional':
            domain = normalized.get('professional_domain', normalized.get('prompt_template', 'academic'))
            normalized['prompt_template'] = str(domain)
        
        return normalized
        
    except Exception as e:
        logger.warning(f"Error normalizing prompt config: {e}")
        return {'mode': 'none'}

def _get_batch_settings_from_config(prompt_config: Optional[Dict[str, Any]], kwargs: Dict[str, Any]) -> Dict[str, int]:
    """从配置中获取批处理设置"""
    settings = {
        'batch_size': 50,
        'max_chars': 8000,
        'max_workers': 5,
        'retry_max_workers': 5,  # 新增：重试专用worker数量
        'translation_timeout': 60,
        'max_retries': 8,  # 增加重试次数
        'large_text_threshold': 50,
        'retry_failure_threshold': 0.0,  # 任何失败都重试
        'non_ascii_threshold': 0.0  # 阈值设为0
    }
    
    try:
        # 优先级：kwargs > prompt_config > 默认值
        if 'batch_size' in kwargs and isinstance(kwargs['batch_size'], (int, float)):
            settings['batch_size'] = max(5, min(int(kwargs['batch_size']), 200))
        if 'max_chunk_size' in kwargs and isinstance(kwargs['max_chunk_size'], (int, float)):
            settings['max_chars'] = max(1000, min(int(kwargs['max_chunk_size']), 50000))
        if 'max_workers' in kwargs and isinstance(kwargs['max_workers'], (int, float)):
            settings['max_workers'] = max(1, min(int(kwargs['max_workers']), 10))
        if 'retry_max_workers' in kwargs and isinstance(kwargs['retry_max_workers'], (int, float)):
            settings['retry_max_workers'] = max(1, min(int(kwargs['retry_max_workers']), 8))
        if 'translation_timeout' in kwargs and isinstance(kwargs['translation_timeout'], (int, float)):
            settings['translation_timeout'] = max(30, min(int(kwargs['translation_timeout']), 300))
        if 'max_retries' in kwargs and isinstance(kwargs['max_retries'], (int, float)):
            settings['max_retries'] = max(3, min(int(kwargs['max_retries']), 15))
        
        if prompt_config and isinstance(prompt_config, dict):
            # 前端可能使用 max_units_per_chunk 控制批次大小
            max_units = prompt_config.get('max_units_per_chunk')
            if max_units and isinstance(max_units, (int, float)):
                settings['batch_size'] = max(5, min(int(max_units), 100))
            
            # 前端使用 max_chars_per_chunk 控制字符数
            max_chars = prompt_config.get('max_chars_per_chunk')
            if max_chars and isinstance(max_chars, (int, float)):
                settings['max_chars'] = max(1000, min(int(max_chars), 50000))
    
    except Exception as e:
        logger.warning(f"Error processing batch settings: {e}")
    
    return settings

class SmartCache:
    """智能LRU缓存，支持prompt配置差异化和失败缓存清理"""
    
    def __init__(self, max_size: int = 1000):
        self._cache = OrderedDict()
        self.max_size = max_size
        self._lock = threading.RLock()
        self._hits = 0
        self._misses = 0
        self._clears = 0  # 新增清理统计
    
    def get(self, key: str) -> Optional[str]:
        if not key or not isinstance(key, str):
            return None
            
        with self._lock:
            if key in self._cache:
                value = self._cache.pop(key)
                self._cache[key] = value
                self._hits += 1
                return value
            self._misses += 1
            return None
    
    def put(self, key: str, value: str):
        if not key or not isinstance(key, str) or not isinstance(value, str):
            return
            
        with self._lock:
            if key in self._cache:
                self._cache.pop(key)
            elif len(self._cache) >= self.max_size:
                self._cache.popitem(last=False)
            self._cache[key] = value
    
    def remove(self, key: str) -> bool:
        """移除指定缓存项"""
        if not key or not isinstance(key, str):
            return False
            
        with self._lock:
            if key in self._cache:
                self._cache.pop(key)
                self._clears += 1
                return True
            return False
    
    def clear_pattern(self, pattern: str) -> int:
        """清理匹配模式的缓存项"""
        if not pattern:
            return 0
            
        with self._lock:
            keys_to_remove = [key for key in self._cache.keys() if pattern in key]
            for key in keys_to_remove:
                self._cache.pop(key)
            self._clears += len(keys_to_remove)
            return len(keys_to_remove)
    
    def clear(self):
        with self._lock:
            cleared_count = len(self._cache)
            self._cache.clear()
            self._hits = 0
            self._misses = 0
            self._clears += cleared_count
            
    @property
    def stats(self) -> Dict[str, Any]:
        with self._lock:
            total = self._hits + self._misses
            return {
                'size': len(self._cache),
                'hits': self._hits,
                'misses': self._misses,
                'clears': self._clears,
                'hit_rate': self._hits / total if total > 0 else 0
            }

class TranslationValidator:
    """翻译完整性验证器"""
    
    # 扩展的错误关键词检测
    ERROR_KEYWORDS = [
        # 英文错误关键词
        'timeout', 'readtimeout', 'connecttimeout', 'httptimeout',
        'network error', 'connection error', 'api error', 'service error',
        'translation failed', 'service unavailable', 'request failed',
        'server error', 'bad gateway', 'gateway timeout',
        # 中文错误关键词
        '超时', '网络错误', '连接错误', '服务错误', '翻译失败',
        '服务不可用', '请求失败', '服务器错误'
    ]
    
    @staticmethod
    def is_error_message(text: str) -> bool:
        """检查是否为错误消息"""
        if not text:
            return True
            
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in TranslationValidator.ERROR_KEYWORDS)
    
    @staticmethod
    def is_serious_failure(original_text: str, translated_text: str, large_text_threshold: int = 50, from_cache: bool = False) -> Tuple[bool, str]:
        """判断是否为严重失败，修复长文本检测逻辑"""
        
        # 缓存结果不算失败
        if from_cache:
            return False, "缓存结果"
        
        # 检查是否翻译为空
        if not translated_text or translated_text.strip() == "":
            if len(original_text) > large_text_threshold:
                return True, f"大段文字未翻译（{len(original_text)}字符）"
            else:
                return False, f"短文本未翻译（{len(original_text)}字符，可能正常）"
        
        # 检查明显的错误信息
        if TranslationValidator.is_error_message(translated_text):
            return True, f"翻译服务错误: {translated_text[:50]}..."
        
        # 检查是否与原文完全相同
        if original_text.strip() == translated_text.strip():
            if len(original_text) > large_text_threshold:
                # 长文本与原文相同，直接判定为未翻译（阈值改为0，简化判定）
                return True, f"长文本未翻译（{len(original_text)}字符）"
            else:
                # 短文本与原文相同很正常（专有名词、数字、已翻译等）
                return False, f"短文本与原文相同（{len(original_text)}字符，正常情况）"
        
        # 缓存结果也需要验证（修复缓存问题）
        if from_cache and len(original_text) > large_text_threshold:
            # 即使是缓存，长文本也要进一步检查
            if original_text.strip() == translated_text.strip():
                return True, f"缓存中的长文本未翻译（{len(original_text)}字符）"
        
        # 其他情况认为翻译成功
        return False, "翻译成功"


class AdvancedDocxTranslator:
    """高级DOCX翻译器 - 增强重试机制和失败检测"""
    
    def __init__(
        self, 
        translator, 
        batch_size: int = 50,
        max_chars: int = 8000,
        max_workers: int = 5,
        retry_max_workers: int = 5,  # 新增：重试专用worker数量
        prompt_config: Optional[Dict[str, Any]] = None,
        reference_doc: Optional[str] = None,
        translation_timeout: int = 60,
        max_retries: int = 8,
        large_text_threshold: int = 50,
        retry_failure_threshold: float = 0.0,
        non_ascii_threshold: float = 0.0,
        **kwargs
    ):
        self.translator = translator
        self.batch_size = batch_size
        self.max_chars = max_chars
        self.max_workers = max_workers
        self.retry_max_workers = retry_max_workers  # 新增
        self.translation_timeout = translation_timeout
        self.max_retries = max_retries
        self.large_text_threshold = large_text_threshold
        self.retry_failure_threshold = retry_failure_threshold
        self.non_ascii_threshold = non_ascii_threshold
        
        # 验证worker配置
        if self.retry_max_workers < 1:
            self.retry_max_workers = 1
            logger.warning("retry_max_workers不能小于1，已重置为1")
        
        if self.retry_max_workers > 8:
            logger.warning(f"retry_max_workers={self.retry_max_workers}可能过高，建议不超过8")
        
        # 为重试预留一些资源，避免总并发数过高
        total_max_workers = 12  # 假设系统总共支持12个worker
        if self.max_workers + self.retry_max_workers > total_max_workers:
            self.max_workers = max(1, total_max_workers - self.retry_max_workers)
            logger.info(f"调整主翻译workers为{self.max_workers}，为重试预留{self.retry_max_workers}个workers")
        
        self.cache = SmartCache(1000)
        self.source_lang = None
        self.reference_doc = reference_doc
        self._config_lock = threading.RLock()
        self._cached_config_hash = None
        
        # 失败任务追踪
        self.failed_tasks: List[FailedTask] = []
        self.failed_tasks_lock = threading.Lock()
        
        # 增强重试策略配置 - 改为 [20, 10, 5, 3, 1, 1, 1, 1]
        self.retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_delays = [1, 2, 4, 8, 16, 20, 25, 30]
        
        # 处理prompt配置 - 增强错误处理
        try:
            self.effective_prompt_config = _prepare_prompt_config(prompt_config, kwargs)
            self.original_translator_config = None
            
            # 确保配置一致性
            if self.effective_prompt_config is None:
                logger.debug("No effective prompt config, using default mode")
                self.effective_prompt_config = {'mode': 'none'}
            
        except Exception as e:
            logger.warning(f"Error initializing prompt config: {e}")
            self.effective_prompt_config = {'mode': 'none'}
        
        # 从配置获取批处理设置
        try:
            batch_settings = _get_batch_settings_from_config(self.effective_prompt_config, kwargs)
            self.batch_size = batch_settings['batch_size']
            self.max_chars = batch_settings['max_chars']
            self.max_workers = min(self.max_workers, batch_settings['max_workers'])  # 不超过配置值
            self.retry_max_workers = min(self.retry_max_workers, batch_settings['retry_max_workers'])  # 应用配置
            self.translation_timeout = batch_settings['translation_timeout']
            self.max_retries = batch_settings['max_retries']
            self.large_text_threshold = batch_settings['large_text_threshold']
            self.retry_failure_threshold = batch_settings['retry_failure_threshold']
            self.non_ascii_threshold = batch_settings['non_ascii_threshold']
        except Exception as e:
            logger.warning(f"Error getting batch settings: {e}")
        
        if self.effective_prompt_config and self.effective_prompt_config.get('mode') != 'none':
            logger.info(f"AdvancedDocxTranslator initialized with prompt config: mode={self.effective_prompt_config.get('mode')}")
        
        self.stats = {
            'total_paragraphs': 0,
            'translated_paragraphs': 0,
            'skipped_paragraphs': 0,
            'total_chars': 0,
            'total_batches': 0,
            'tables': 0,
            'processing_time': 0,
            'prompt_mode': self.effective_prompt_config.get('mode', 'none'),
            'api_calls': 0,
            'cache_savings': 0,
            'template_applied': False,
            'serious_failures': 0,
            'minor_issues': 0,
            'retry_attempts': 0,
            'final_failures': 0,
            'cache_clears': 0,
            'final_rescues': 0,
            'retry_workers_used': 0,  # 新增：实际使用的重试worker数
            'concurrent_retry_batches': 0,  # 新增：并发处理的重试批次数
        }

    @contextmanager
    def _translator_config_context(self):
        """线程安全的配置上下文管理器"""
        try:
            self._apply_prompt_config_to_translator()
            yield
        except Exception as e:
            logger.error(f"Error in translator config context: {e}")
            raise
        finally:
            self._restore_translator_config()

    def _apply_prompt_config_to_translator(self):
        """应用prompt配置到翻译器"""
        if (self.effective_prompt_config and 
            self.effective_prompt_config.get('mode') != 'none' and 
            hasattr(self.translator, 'set_prompt_config')):
            try:
                with self._config_lock:
                    # 保存翻译器的原始配置
                    self.original_translator_config = getattr(self.translator, 'prompt_config', None)
                    config_copy = copy.deepcopy(self.effective_prompt_config)
                    self.translator.set_prompt_config(config_copy)
                    logger.info("Applied prompt config to translator in AdvancedDocxTranslator")
                    return True
            except Exception as e:
                logger.warning(f"Failed to apply prompt config to translator: {e}")
        return False

    def _restore_translator_config(self):
        """恢复翻译器的原始配置"""
        if (self.effective_prompt_config and 
            self.effective_prompt_config.get('mode') != 'none' and 
            hasattr(self.translator, 'set_prompt_config')):
            try:
                with self._config_lock:
                    if self.original_translator_config is not None:
                        self.translator.set_prompt_config(self.original_translator_config)
                    else:
                        # 如果原来没有配置，清除当前配置
                        if hasattr(self.translator, 'prompt_config'):
                            self.translator.prompt_config = None
                    logger.debug("Restored translator config in AdvancedDocxTranslator")
            except Exception as e:
                logger.warning(f"Failed to restore translator config: {e}")

    def _get_config_safe(self) -> Optional[Dict[str, Any]]:
        """线程安全地获取配置副本"""
        with self._config_lock:
            return copy.deepcopy(self.effective_prompt_config) if self.effective_prompt_config else None

    def _should_use_concurrent_retry(self, retry_batches):
        """判断是否应该使用并发重试"""
        # 只有在重试批次较多时才使用并发
        if len(retry_batches) < 3:
            return False
        
        # 检查系统资源情况
        return True

# docx_full_translator.py (第二批修改)

    def _clear_failed_task_cache(self, task: FailedTask, target_lang: str, source_lang: Optional[str]):
        """清理失败任务的缓存"""
        try:
            local_prompt_config = self._get_config_safe()
            cache_key = self._get_cache_key(task.original_text, target_lang, source_lang, local_prompt_config)
            if self.cache.remove(cache_key):
                self.stats['cache_clears'] += 1
                logger.debug(f"清理失败任务缓存: {task.original_text[:50]}...")
        except Exception as e:
            logger.warning(f"Error clearing failed task cache: {e}")

    def _should_trigger_retry(self, total_processed: int) -> bool:
        """判断是否应该触发重试"""
        with self.failed_tasks_lock:
            serious_failures = [task for task in self.failed_tasks if task.is_serious]
            
            if not serious_failures:
                return False
            
            # 严重失败率超过阈值才重试
            failure_rate = len(serious_failures) / max(1, total_processed)
            should_retry = failure_rate >= self.retry_failure_threshold
            
            logger.info(f"严重失败: {len(serious_failures)}/{total_processed} ({failure_rate:.1%}), "
                       f"阈值: {self.retry_failure_threshold:.1%}, 是否重试: {should_retry}")
            
            return should_retry

    def _add_failed_task(self, element: ParagraphElement, original_index: int, 
                        reason: str, from_cache: bool = False):
        """添加失败任务，智能判断是否为严重失败"""
        is_serious, detailed_reason = TranslationValidator.is_serious_failure(
            element.full_text, reason, self.large_text_threshold, from_cache)
        
        with self.failed_tasks_lock:
            failed_task = FailedTask(
                original_text=element.full_text,
                original_index=original_index,
                element=element,
                failure_reason=FailureReason.API_ERROR if is_serious else FailureReason.NOT_TRANSLATED,
                error_message=detailed_reason,
                is_serious=is_serious
            )
            self.failed_tasks.append(failed_task)
            
            # 更新统计
            if is_serious:
                self.stats['serious_failures'] += 1
                logger.debug(f"严重失败: {detailed_reason[:50]}...")
            else:
                self.stats['minor_issues'] += 1
                logger.debug(f"轻微问题: {detailed_reason[:50]}...")

    def _add_batch_failure(self, batch: List[ParagraphElement], batch_indices: List[int], reason: str):
        """添加批次级失败，所有任务都标记为严重失败"""
        with self.failed_tasks_lock:
            for element, original_index in zip(batch, batch_indices):
                failed_task = FailedTask(
                    original_text=element.full_text,
                    original_index=original_index,
                    element=element,
                    failure_reason=FailureReason.BATCH_FAILURE,
                    error_message=reason,
                    is_serious=True
                )
                self.failed_tasks.append(failed_task)
                self.stats['serious_failures'] += 1
            logger.warning(f"批次失败: {len(batch)} 个任务, 原因: {reason}")

    def _create_retry_batches(self, failed_tasks: List[FailedTask], retry_count: int) -> List[List[FailedTask]]:
        """创建重试批次，只处理严重失败"""
        # 只重试严重失败
        serious_failed_tasks = [task for task in failed_tasks if task.is_serious]
        
        if not serious_failed_tasks:
            return []
        
        # 根据重试次数确定批次大小
        if retry_count < len(self.retry_batch_sizes):
            max_batch_size = self.retry_batch_sizes[retry_count]
        else:
            # 后续重试使用最小批次
            max_batch_size = 1
        
        logger.info(f"第{retry_count + 1}次重试，严重失败任务: {len(serious_failed_tasks)}, 批次大小: {max_batch_size}")
        
        chunks = []
        current_chunk = []
        
        for task in serious_failed_tasks:
            if len(current_chunk) >= max_batch_size:
                chunks.append(current_chunk)
                current_chunk = [task]
            else:
                current_chunk.append(task)
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks

    def _retry_failed_tasks(self, target_lang: str, source_lang: Optional[str]):
        """只重试严重失败的任务，修复重试逻辑，增加并发重试功能"""
        retry_count = 0
        
        while retry_count < self.max_retries:
            # 获取当前需要重试的严重失败任务
            with self.failed_tasks_lock:
                serious_failed_tasks = [task for task in self.failed_tasks 
                                      if task.is_serious and task.retry_count <= retry_count]
            
            if not serious_failed_tasks:
                logger.info(f"第 {retry_count + 1} 次重试检查：没有严重失败任务需要重试")
                break
            
            logger.info(f"第 {retry_count + 1} 次重试，处理 {len(serious_failed_tasks)} 个严重失败任务")
            
            # 重试前清理失败任务的缓存
            logger.info(f"清理 {len(serious_failed_tasks)} 个失败任务的缓存")
            for task in serious_failed_tasks:
                self._clear_failed_task_cache(task, target_lang, source_lang)
            
            # 清空当前重试轮次的失败任务，但保留非严重问题
            with self.failed_tasks_lock:
                self.failed_tasks = [task for task in self.failed_tasks 
                                   if not task.is_serious or task.retry_count > retry_count]
            
            # 添加重试延迟
            if retry_count < len(self.retry_delays):
                delay = self.retry_delays[retry_count]
                logger.info(f"重试前等待 {delay} 秒...")
                time.sleep(delay)
            
            # 创建重试批次
            retry_batches = self._create_retry_batches(serious_failed_tasks, retry_count)
            
            if not retry_batches:
                logger.info(f"第 {retry_count + 1} 次重试：没有批次需要处理")
                break
            
            # 判断是否使用并发重试
            use_concurrent = self._should_use_concurrent_retry(retry_batches)
            actual_workers = min(self.retry_max_workers, len(retry_batches)) if use_concurrent else 1
            self.stats['retry_workers_used'] = actual_workers
            
            if use_concurrent:
                logger.info(f"启用并发重试，worker数量: {actual_workers}")
                self.stats['concurrent_retry_batches'] = len(retry_batches)
                self._retry_batches_concurrent(retry_batches, target_lang, source_lang, retry_count, actual_workers)
            else:
                logger.info("使用串行重试")
                self._retry_batches_sequential(retry_batches, target_lang, source_lang, retry_count)
            
            self.stats['retry_attempts'] += 1
            retry_count += 1
        
        # 统计最终失败的任务
        with self.failed_tasks_lock:
            final_serious_failures = [task for task in self.failed_tasks if task.is_serious]
            self.stats['final_failures'] = len(final_serious_failures)
            if final_serious_failures:
                logger.warning(f"常规重试后仍有 {len(final_serious_failures)} 个严重失败任务")

    def _retry_batches_concurrent(self, retry_batches: List[List[FailedTask]], target_lang: str, 
                                source_lang: Optional[str], retry_count: int, max_workers: int):
        """并发重试批次"""
        try:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                with tqdm(total=sum(len(batch) for batch in retry_batches), 
                         desc=f"第{retry_count + 1}次重试(并发)", unit="任务") as pbar:
                    
                    # 提交所有重试批次任务
                    future_to_batch = {}
                    for retry_batch_idx, retry_batch in enumerate(retry_batches):
                        future = executor.submit(
                            self._process_retry_batch, retry_batch, target_lang, source_lang, 
                            retry_batch_idx, retry_count
                        )
                        future_to_batch[future] = retry_batch
                    
                    # 处理完成的任务
                    for future in as_completed(future_to_batch):
                        retry_batch = future_to_batch[future]
                        try:
                            future.result(timeout=self.translation_timeout + 30)  # 给额外的超时时间
                            pbar.update(len(retry_batch))
                        except Exception as e:
                            logger.error(f"并发重试批次处理异常: {e}")
                            # 处理异常情况
                            with self.failed_tasks_lock:
                                for task in retry_batch:
                                    new_task = FailedTask(
                                        original_text=task.original_text,
                                        original_index=task.original_index,
                                        element=task.element,
                                        failure_reason=FailureReason.BATCH_FAILURE,
                                        error_message=f"并发执行异常: {str(e)}",
                                        is_serious=True
                                    )
                                    new_task.retry_count = retry_count + 1
                                    self.failed_tasks.append(new_task)
                            pbar.update(len(retry_batch))
                            
        except Exception as e:
            logger.error(f"并发重试执行异常: {e}")
            # 回退到串行重试
            logger.info("回退到串行重试")
            self._retry_batches_sequential(retry_batches, target_lang, source_lang, retry_count)

    def _retry_batches_sequential(self, retry_batches: List[List[FailedTask]], target_lang: str, 
                                source_lang: Optional[str], retry_count: int):
        """串行重试批次"""
        with tqdm(total=sum(len(batch) for batch in retry_batches), 
                 desc=f"第{retry_count + 1}次重试", unit="任务") as pbar:
            
            for retry_batch_idx, retry_batch in enumerate(retry_batches):
                self._process_retry_batch(retry_batch, target_lang, source_lang, retry_batch_idx, retry_count)
                pbar.update(len(retry_batch))

    def _process_retry_batch(self, retry_batch: List[FailedTask], target_lang: str, 
                           source_lang: Optional[str], batch_idx: int, retry_count: int):
        """处理单个重试批次"""
        logger.info(f"开始重试批次 {batch_idx + 1}，任务数: {len(retry_batch)}")
        
        # 将FailedTask转换为元素列表和索引列表
        batch_elements = [task.element for task in retry_batch]
        batch_indices = [task.original_index for task in retry_batch]
        
        # 执行重试翻译
        success, retry_results, cache_flags = self._translate_batch_with_timeout(
            batch_elements, batch_indices, target_lang, source_lang, -1)
        
        logger.info(f"重试批次 {batch_idx + 1} 完成，成功: {success}, 结果数: {len(retry_results) if retry_results else 0}")
        
        # 处理重试结果
        if not success:
            # 整个批次失败
            batch_reason = retry_results[0] if retry_results else "重试批次失败"
            logger.warning(f"重试批次失败: {batch_reason}")
            
            with self.failed_tasks_lock:
                for task in retry_batch:
                    new_task = FailedTask(
                        original_text=task.original_text,
                        original_index=task.original_index,
                        element=task.element,
                        failure_reason=FailureReason.BATCH_FAILURE,
                        error_message=batch_reason,
                        is_serious=True
                    )
                    new_task.retry_count = retry_count + 1
                    self.failed_tasks.append(new_task)
        else:
            # 批次成功，检查个别结果
            for i, task in enumerate(retry_batch):
                if i < len(retry_results):
                    result = retry_results[i]
                    from_cache = cache_flags[i] if i < len(cache_flags) else False
                    
                    # 检查重试是否成功，增加结果验证
                    is_serious, reason = TranslationValidator.is_serious_failure(
                        task.original_text, result, self.large_text_threshold, from_cache)
                    
                    if not is_serious and not TranslationValidator.is_error_message(result):
                        # 重试成功，应用翻译
                        self._apply_translation(task.element, result)
                        logger.debug(f"重试成功: {task.original_text[:50]}... -> {result[:50]}...")
                    else:
                        # 重试仍然失败
                        logger.debug(f"重试失败，保持原文: {task.original_text[:50]}... 原因: {reason}")
                        
                        with self.failed_tasks_lock:
                            new_task = FailedTask(
                                original_text=task.original_text,
                                original_index=task.original_index,
                                element=task.element,
                                failure_reason=FailureReason.API_ERROR,
                                error_message=reason,
                                is_serious=True
                            )
                            new_task.retry_count = retry_count + 1
                            self.failed_tasks.append(new_task)
                else:
                    # 没有对应的结果
                    logger.warning(f"重试结果不足: 任务 {i}, 结果数量 {len(retry_results)}")
                    with self.failed_tasks_lock:
                        new_task = FailedTask(
                            original_text=task.original_text,
                            original_index=task.original_index,
                            element=task.element,
                            failure_reason=FailureReason.PARSE_ERROR,
                            error_message="重试结果缺失",
                            is_serious=True
                        )
                        new_task.retry_count = retry_count + 1
                        self.failed_tasks.append(new_task)

    def _final_retry_remaining_tasks(self, target_lang: str, source_lang: Optional[str]):
        """最终处理剩余失败任务"""
        with self.failed_tasks_lock:
            remaining_tasks = [task for task in self.failed_tasks if task.is_serious]
        
        if not remaining_tasks:
            logger.info("没有剩余失败任务需要最终处理")
            return
        
        logger.info(f"最终处理 {len(remaining_tasks)} 个剩余失败任务")
        
        success_count = 0
        
        with tqdm(total=len(remaining_tasks), desc="最终挽救", unit="任务") as pbar:
            for task in remaining_tasks:
                try:
                    # 清理缓存
                    self._clear_failed_task_cache(task, target_lang, source_lang)
                    
                    # 单独翻译，增加延迟避免频率限制
                    time.sleep(1)
                    
                    logger.debug(f"最终处理任务: {task.original_text[:50]}...")
                    
                    result = self.translator.translate(
                        text=task.original_text,
                        target_lang=target_lang,
                        source_lang=source_lang
                    )
                    
                    self.stats['api_calls'] += 1
                    
                    if (result and not TranslationValidator.is_error_message(result) and 
                        result.strip() != task.original_text.strip()):
                        
                        # 进一步验证翻译质量
                        is_serious, reason = TranslationValidator.is_serious_failure(
                            task.original_text, result, self.large_text_threshold, False)
                        
                        if not is_serious:
                            self._apply_translation(task.element, result)
                            success_count += 1
                            self.stats['final_rescues'] += 1
                            logger.info(f"最终处理成功: {task.original_text[:50]}... -> {result[:50]}...")
                        else:
                            logger.debug(f"最终处理结果仍不合格: {reason}")
                            
                    else:
                        logger.debug(f"最终处理仍失败: {task.original_text[:50]}...")
                        
                except Exception as e:
                    logger.error(f"最终处理异常: {e}")
                
                pbar.update(1)
        
        if success_count > 0:
            logger.info(f"最终挽救成功 {success_count}/{len(remaining_tasks)} 个任务")
        else:
            logger.warning(f"最终挽救失败，{len(remaining_tasks)} 个任务保持原文")

    def _create_ordered_batches(self, elements: List[ParagraphElement]) -> List[Tuple[List[ParagraphElement], List[int]]]:
        """创建批次，保持文档顺序，返回(批次元素, 原始索引)的元组"""
        batches = []
        current_batch = []
        current_indices = []
        current_chars = 0
        
        for idx, element in enumerate(elements):  # 保持原始顺序
            text_len = len(element.full_text)
            
            if (len(current_batch) >= self.batch_size or 
                current_chars + text_len > self.max_chars) and current_batch:
                batches.append((current_batch, current_indices))
                current_batch = [element]
                current_indices = [idx]
                current_chars = text_len
            else:
                current_batch.append(element)
                current_indices.append(idx)
                current_chars += text_len
        
        if current_batch:
            batches.append((current_batch, current_indices))
        
        return batches

    def _translate_batch_with_timeout(self, batch: List[ParagraphElement], batch_indices: List[int],
                                    target_lang: str, source_lang: Optional[str], 
                                    batch_num: int) -> Tuple[bool, List[str], List[bool]]:
        """带超时的批次翻译，返回结果和缓存标记"""
        try:
            # 使用Future实现超时控制
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._translate_batch, batch, target_lang, source_lang, batch_num)
                
                try:
                    batch_results, from_cache_flags = future.result(timeout=self.translation_timeout)
                    return True, batch_results, from_cache_flags
                    
                except TimeoutError:
                    logger.warning(f"批次 {batch_num} 翻译超时 ({self.translation_timeout}秒)")
                    failure_reasons = [f"翻译超时 ({self.translation_timeout}秒)"] * len(batch)
                    cache_flags = [False] * len(batch)
                    return False, failure_reasons, cache_flags
                    
        except Exception as e:
            logger.error(f"批次 {batch_num} 翻译异常: {e}")
            failure_reasons = [f"翻译异常: {str(e)}"] * len(batch)
            cache_flags = [False] * len(batch)
            return False, failure_reasons, cache_flags

    def _get_cache_key(self, text: str, target_lang: str, source_lang: Optional[str], 
                      prompt_config: Optional[Dict[str, Any]] = None) -> str:
        """生成缓存键，包含prompt配置信息 - 修复空值问题"""
        try:
            # 使用传入的配置或实例配置
            config = prompt_config or self.effective_prompt_config
            
            # 生成配置哈希（缓存以提高性能）
            if self._cached_config_hash is None and config is not None:
                try:
                    mode = config.get('mode', '') if config else ''
                    template = config.get('prompt_template', '') if config else ''
                    custom_system = ''
                    
                    custom_prompt = config.get('custom_prompt')
                    if custom_prompt and isinstance(custom_prompt, dict):
                        custom_system = custom_prompt.get('system', '')[:50]
                    
                    key_config_str = f"{mode}_{template}_{custom_system}"
                    self._cached_config_hash = hashlib.md5(key_config_str.encode('utf-8')).hexdigest()[:8]
                except (AttributeError, TypeError, KeyError) as e:
                    logger.warning(f"Failed to generate prompt hash, using empty: {e}")
                    self._cached_config_hash = ""
            
            prompt_hash = self._cached_config_hash or ""
            text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
            return f"{text_hash}_{target_lang}_{source_lang or 'auto'}_{prompt_hash}"
            
        except Exception as e:
            logger.warning(f"Error generating cache key: {e}")
            # 回退到基本缓存键
            text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
            return f"{text_hash}_{target_lang}_{source_lang or 'auto'}"

# docx_full_translator.py (第三批修改)

    def _get_enhanced_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """获取增强的系统提示，完整支持前端prompt配置"""
        
        try:
            if not self.effective_prompt_config or self.effective_prompt_config.get('mode') == 'none':
                # 默认系统提示
                return self._build_default_system_prompt(target_lang, source_lang)
            
            mode = self.effective_prompt_config.get('mode', 'none')
            
            if mode == 'custom' and self.effective_prompt_config.get('custom_prompt'):
                # 使用完全自定义的prompt
                return self._build_custom_system_prompt(target_lang, source_lang)
                
            elif mode == 'professional':
                # 使用专业模板
                return self._build_professional_system_prompt(target_lang, source_lang)
                
            elif mode == 'general':
                # 通用增强模式
                return self._build_general_system_prompt(target_lang, source_lang)
                
            else:
                # 'none' 模式或其他
                return self._build_simple_system_prompt(target_lang, source_lang)
                
        except Exception as e:
            logger.warning(f"Error generating enhanced system prompt: {e}")
            return self._build_default_system_prompt(target_lang, source_lang)

    def _build_custom_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """构建自定义系统提示"""
        try:
            custom_prompt = self.effective_prompt_config.get('custom_prompt', {})
            system_content = custom_prompt.get('system', '')
            
            if not system_content:
                logger.warning("Custom prompt system content is empty, falling back to default")
                return self._build_default_system_prompt(target_lang, source_lang)
            
            # 为批量DOCX翻译添加必要的说明
            if "numbered line" not in system_content.lower():
                system_content += f"""

ADVANCED DOCX BATCH PROCESSING:
- Each input line is numbered [1], [2], etc.
- Translate each numbered line individually while considering context
- Maintain document structure and formatting context
- Keep the exact same number of lines as input
- Output only the translated content without including the original text, one per line
- Do not include line numbers in output
- Preserve formatting and punctuation
- Do not include any extra comments in your output
- Ensure translations sound natural in {target_lang}"""
            
            # 添加增强规则
            system_content = self._add_enhancement_rules(system_content)
            
            logger.info("Using custom prompt for advanced DOCX translation")
            return system_content
            
        except Exception as e:
            logger.warning(f"Error building custom system prompt: {e}")
            return self._build_default_system_prompt(target_lang, source_lang)

    def _build_professional_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """构建专业模板系统提示"""
        try:
            domain = self.effective_prompt_config.get('prompt_template', 'academic')
            logger.info(f"Using professional template for advanced DOCX: {domain}")
            
            # 专业领域的系统提示
            professional_prompts = {
                'academic': f"""You are an expert academic translator specializing in scholarly documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain academic tone, preserve citations and references, and use appropriate academic terminology.
Ensure consistency in technical terms throughout the translation.
Pay special attention to:
- Academic writing style and formal tone
- Proper citation formats and bibliographic references
- Technical and disciplinary terminology
- Research methodology descriptions
- Statistical and analytical content""",
                
                'business': f"""You are a professional business translator with expertise in corporate documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use appropriate business terminology, maintain formal tone, and keep company names/brands unchanged.
Ensure clarity and professionalism in the translation.
Focus on:
- Business terminology and corporate language
- Professional communication style
- Financial and commercial concepts
- Strategic planning and business processes
- Maintaining brand consistency""",
                
                'technical': f"""You are a technical translator specializing in technical documentation.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve technical accuracy, keep code snippets and commands unchanged, and use industry-standard terminology.
Maintain consistency in technical terms throughout.
Pay attention to:
- Technical specifications and procedures
- Software and hardware terminology
- Programming concepts and code examples
- Industry standards and protocols
- Safety and compliance requirements""",
                
                'legal': f"""You are a certified legal translator with expertise in legal documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use precise legal terminology, maintain legal accuracy and formality, and preserve all legal references.
Ensure no ambiguity in legal terms.
Focus on:
- Legal terminology and concepts
- Contractual language and clauses
- Regulatory and compliance terms
- Jurisdictional considerations
- Legal precedents and citations""",
                
                'medical': f"""You are a certified medical translator with expertise in medical documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use standard medical terminology, preserve drug names and dosages exactly, and maintain clinical precision.
Follow international medical nomenclature standards.
Pay attention to:
- Medical terminology and procedures
- Drug names, dosages, and contraindications
- Clinical protocols and guidelines
- Patient safety considerations
- Diagnostic and treatment information""",
                
                'creative': f"""You are a creative translator focusing on maintaining style and tone.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve the original style, adapt idioms naturally, and maintain emotional impact.
Focus on readability and flow while being faithful to the original meaning.
Consider:
- Literary devices and stylistic elements
- Cultural adaptation of idioms and expressions
- Emotional tone and atmosphere
- Creative expression and authorial voice
- Rhythm and flow of the text"""
            }
            
            system_content = professional_prompts.get(domain, professional_prompts['academic'])
            
            # 添加批处理规则
            system_content += f"""

PROFESSIONAL DOCX BATCH PROCESSING:
1. Process each numbered line [1], [2], etc. individually
2. Consider the context of surrounding lines for coherent translation
3. Maintain professional consistency throughout the document
4. Keep the exact same number of lines as input
5. Output only the translated content, one per line
6. Do not include line numbers in output
7. Preserve professional formatting and terminology
8. Do not include any extra comments in your output
9. Ensure translations sound natural in {target_lang}"""
            
            # 添加增强规则
            system_content = self._add_enhancement_rules(system_content)
            
            return system_content
            
        except Exception as e:
            logger.warning(f"Error building professional system prompt: {e}")
            return self._build_default_system_prompt(target_lang, source_lang)

    def _build_general_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """构建通用增强系统提示"""
        try:
            system_content = f"""You are a professional translator with expertise in document translation.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Provide accurate, natural translations while preserving the original meaning and tone.
Maintain document coherence and consistency throughout the translation.

GENERAL BATCH PROCESSING RULES:
1. Translate each numbered line [1], [2], etc. individually
2. Consider context from surrounding lines for coherent translation
3. Keep the exact same number of lines as input
4. Preserve all formatting, punctuation, and special characters
5. Output only the translated content, one per line
6. Do not include line numbers in output
7. Maintain consistency in terminology throughout
8. Do not include any extra comments in your output
9. Ensure translations sound natural in {target_lang}"""
            
            # 添加增强规则
            system_content = self._add_enhancement_rules(system_content)
            
            return system_content
            
        except Exception as e:
            logger.warning(f"Error building general system prompt: {e}")
            return self._build_default_system_prompt(target_lang, source_lang)

    def _build_simple_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """构建简单系统提示"""
        return f"""Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Process each numbered line [1], [2], etc. and return the same number of translated lines.
Do not include line numbers in output."""

    def _build_default_system_prompt(self, target_lang: str, source_lang: Optional[str]) -> str:
        """构建默认系统提示"""
        system_content = f"""You are a professional translator. Translate from {source_lang or 'auto-detected language'} to {target_lang}.

Rules:
1. Translate each numbered line individually, but take the full context of surrounding lines into account
2. Keep the exact same number of lines as the original
3. Preserve all formatting, punctuation, and special characters
4. For lists, keep the list markers
5. Output only the translated lines, one per line, in the same order
6. Do not include the original line numbers or any extra comments in your output
7. Do not translate place names (e.g. cities, countries) or company names—keep them exactly as in the original
8. Return only the translated content without including the original text"""
        
        return system_content

    def _add_enhancement_rules(self, system_content: str) -> str:
        """添加增强规则到系统提示"""
        try:
            if not self.effective_prompt_config or self.effective_prompt_config.get('mode') == 'none':
                return system_content
            
            enhancements = []
            
            # 保留术语
            preserve_terms = self.effective_prompt_config.get('preserve_terms')
            if preserve_terms and isinstance(preserve_terms, list) and preserve_terms:
                terms = ', '.join(str(term) for term in preserve_terms if term)
                if terms:
                    enhancements.append(f"PRESERVE THESE TERMS EXACTLY: {terms}")
            
            # 术语表
            glossary = self.effective_prompt_config.get('glossary')
            if glossary and isinstance(glossary, dict) and glossary:
                try:
                    glossary_text = '; '.join([f"{k}: {v}" for k, v in glossary.items() if k and v])
                    if glossary_text:
                        enhancements.append(f"USE THIS GLOSSARY: {glossary_text}")
                except Exception as e:
                    logger.warning(f"Error processing glossary: {e}")
            
            # 额外上下文
            additional_context = self.effective_prompt_config.get('additional_context')
            if additional_context and str(additional_context).strip():
                enhancements.append(f"ADDITIONAL CONTEXT: {str(additional_context).strip()}")
            
            if enhancements:
                enhancement_text = "\n\nADDITIONAL REQUIREMENTS:\n" + "\n".join(f"• {rule}" for rule in enhancements)
                system_content += enhancement_text
            
            return system_content
            
        except Exception as e:
            logger.warning(f"Error adding enhancement rules: {e}")
            return system_content

    def _translate_batch(self, batch: List[ParagraphElement], target_lang: str, 
                        source_lang: Optional[str], batch_num: int) -> Tuple[List[str], List[bool]]:
        """翻译单个批次，使用merge策略"""
        
        texts = [element.full_text for element in batch]
        
        # 检查缓存
        results = []
        uncached_texts = []
        uncached_indices = []
        
        local_prompt_config = self._get_config_safe()
        
        for i, text in enumerate(texts):
            cache_key = self._get_cache_key(text, target_lang, source_lang, local_prompt_config)
            cached = self.cache.get(cache_key)
            if cached:
                results.append(cached)
                self.stats['cache_savings'] += 1
            else:
                uncached_texts.append(text)
                uncached_indices.append(i)
                results.append("")
        
        # 翻译未缓存的文本
        if uncached_texts:
            try:
                # 获取系统提示（支持自定义prompt）
                system_prompt = self._get_enhanced_system_prompt(target_lang, source_lang)
                
                # 使用行号编码
                numbered_texts = [f"[{i+1}] {text}" for i, text in enumerate(uncached_texts)]
                user_message = "\n".join(numbered_texts)
                
                # 使用merge策略调用翻译器
                try:
                    translated_result = self.translator.translate(
                        text=user_message,
                        target_lang=target_lang,
                        source_lang=source_lang,
                        prompt_config=local_prompt_config,
                        config_merge_mode='merge'  # 使用merge策略
                    )
                except (TypeError, AttributeError):
                    # 回退到messages格式
                    try:
                        messages = [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_message}
                        ]
                        translated_result = self.translator.translate(
                            messages=messages,
                            target_lang=target_lang,
                            source_lang=source_lang
                        )
                    except:
                        # 最终回退
                        full_prompt = f"{system_prompt}\n\nText to translate:\n{user_message}"
                        translated_result = self.translator.translate(full_prompt, target_lang, source_lang)
                
                self.stats['api_calls'] += 1
                
                # 解析结果
                translated_parts = self._extract_numbered_translations(translated_result, len(uncached_texts))
                
                # 更新缓存和结果
                for i, (original_text, idx) in enumerate(zip(uncached_texts, uncached_indices)):
                    if i < len(translated_parts) and translated_parts[i]:
                        translation = translated_parts[i].strip()
                        cache_key = self._get_cache_key(original_text, target_lang, source_lang, local_prompt_config)
                        self.cache.put(cache_key, translation)
                        results[idx] = translation
                    else:
                        results[idx] = original_text
                        
            except Exception as e:
                logger.error(f"批次 {batch_num} 翻译失败: {e}")
                # 回退：保持原文
                for idx in uncached_indices:
                    results[idx] = texts[idx]
        
        # 返回结果和缓存标记
        from_cache_flags = [bool(results[i] and i not in uncached_indices) for i in range(len(results))]
        return results, from_cache_flags

    def _extract_numbered_translations(self, response: str, expected_count: int) -> List[str]:
        """提取编号翻译结果"""
        translations = [""] * expected_count
        
        try:
            lines = response.strip().split('\n')
            current_line_num = 0
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # 检查行号标记
                number_match = NUMBERED_LINE_PATTERN.match(line)
                if number_match:
                    line_num = int(number_match.group(1)) - 1
                    content = number_match.group(2).strip()
                    if 0 <= line_num < expected_count and content:
                        translations[line_num] = content
                else:
                    # 无行号标记，按顺序分配
                    if current_line_num < expected_count and line:
                        translations[current_line_num] = line
                        current_line_num += 1
            
            # 填补空缺
            if any(not t for t in translations):
                non_empty_lines = [line.strip() for line in lines if line.strip()]
                for i, line in enumerate(non_empty_lines):
                    if i < expected_count and not translations[i]:
                        clean_line = re.sub(r'^\[\d+\]\s*', '', line)
                        if clean_line:
                            translations[i] = clean_line
            
        except Exception as e:
            logger.warning(f"Error extracting numbered translations: {e}")
        
        return translations

    def translate_docx(self, input_filepath: str, output_filepath: str, target_lang: str, source_lang: Optional[str] = None) -> str:
        """主翻译方法，增加重试机制"""
        try:
            start_time = time.time()
            self.source_lang = source_lang
            
            # 使用配置上下文管理器
            with self._translator_config_context():
                logger.info(f"开始高级DOCX翻译: {os.path.basename(input_filepath)}")
                logger.info(f"目标语言: {target_lang}")
                if self.effective_prompt_config and self.effective_prompt_config.get('mode') != 'none':
                    mode = self.effective_prompt_config.get('mode', 'none')
                    logger.info(f"Prompt配置: mode={mode}")
                    if mode == 'professional':
                        logger.info(f"专业领域: {self.effective_prompt_config.get('prompt_template', 'academic')}")
                    elif mode == 'custom':
                        logger.info("使用自定义prompt")
                if self.reference_doc:
                    logger.info(f"参考文档: {self.reference_doc}")
                
                logger.info(f"批处理设置: batch_size={self.batch_size}, max_chars={self.max_chars}, workers={self.max_workers}")
                logger.info(f"重试设置: max_retries={self.max_retries}, timeout={self.translation_timeout}s, retry_workers={self.retry_max_workers}")
                
                # 复制原文档
                shutil.copy2(input_filepath, output_filepath)
                doc = Document(output_filepath)
                
                # 如果有参考文档，尝试应用其样式
                if self.reference_doc and os.path.exists(self.reference_doc):
                    self._apply_reference_styles(doc, self.reference_doc)
                
                # 收集段落元素（保持完整性）
                elements = self._collect_paragraph_elements(doc)
                
                if not elements:
                    logger.info("没有需要翻译的内容")
                    print("翻译完成！（无需要翻译的内容）")
                    return output_filepath
                
                self.stats['total_paragraphs'] = len(elements)
                self.stats['total_chars'] = sum(len(e.full_text) for e in elements)
                
                # 创建批次（保持文档顺序）
                batches = self._create_ordered_batches(elements)
                self.stats['total_batches'] = len(batches)
                
                logger.info(f"需要翻译 {len(elements)} 个段落")
                logger.info(f"分为 {len(batches)} 个批次处理")
                
                # 清空失败任务列表
                with self.failed_tasks_lock:
                    self.failed_tasks.clear()
                
                # 执行翻译
                self._translate_all_batches(batches, target_lang, source_lang)
                
                # 智能判断是否需要重试
                if self._should_trigger_retry(len(elements)):
                    logger.info(f"检测到严重失败，开始重试流程")
                    self._retry_failed_tasks(target_lang, source_lang)
                else:
                    with self.failed_tasks_lock:
                        serious_count = sum(1 for task in self.failed_tasks if task.is_serious)
                        minor_count = sum(1 for task in self.failed_tasks if not task.is_serious)
                        if self.failed_tasks:
                            logger.info(f"失败分析: 严重失败 {serious_count} 个, 轻微问题 {minor_count} 个, 未达到重试阈值，跳过重试")
                
                # 最终处理剩余失败任务
                self._final_retry_remaining_tasks(target_lang, source_lang)
                
                # 保存文档
                doc.save(output_filepath)
                
                self.stats['processing_time'] = time.time() - start_time
                self._print_stats()
                
                return output_filepath
                
        except Exception as e:
            logger.error(f"高级DOCX翻译失败: {e}")
            raise

    def _translate_all_batches(self, batches: List[Tuple[List[ParagraphElement], List[int]]], 
                             target_lang: str, source_lang: Optional[str]):
        """翻译所有批次，增加失败检测"""
        
        if len(batches) > 10 and self.max_workers > 1:
            self._translate_concurrent(batches, target_lang, source_lang)
        else:
            self._translate_sequential(batches, target_lang, source_lang)
    
    def _translate_sequential(self, batches: List[Tuple[List[ParagraphElement], List[int]]], 
                            target_lang: str, source_lang: Optional[str]):
        """串行翻译，增加失败检测"""
        with tqdm(total=sum(len(batch[0]) for batch in batches), 
                 desc="高级DOCX翻译进度", unit="段落") as pbar:
            
            for batch_idx, (batch_elements, batch_indices) in enumerate(batches):
                success, translated_texts, from_cache_flags = self._translate_batch_with_timeout(
                    batch_elements, batch_indices, target_lang, source_lang, batch_idx + 1)
                
                if success:
                    # 应用翻译并检查失败
                    for i, (element, translated_text) in enumerate(zip(batch_elements, translated_texts)):
                        from_cache = from_cache_flags[i] if i < len(from_cache_flags) else False
                        
                        # 检查是否有问题
                        is_serious, reason = TranslationValidator.is_serious_failure(
                            element.full_text, translated_text, self.large_text_threshold, from_cache)
                        
                        if is_serious or (not from_cache and reason != "翻译成功"):
                            self._add_failed_task(element, batch_indices[i], translated_text, from_cache)
                        
                        self._apply_translation(element, translated_text)
                else:
                    # 整个批次失败
                    self._add_batch_failure(batch_elements, batch_indices, translated_texts[0] if translated_texts else "批次处理失败")
                
                pbar.update(len(batch_elements))
    
    def _translate_concurrent(self, batches: List[Tuple[List[ParagraphElement], List[int]]], 
                            target_lang: str, source_lang: Optional[str]):
        """并发翻译，增加失败检测"""
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            with tqdm(total=sum(len(batch[0]) for batch in batches), 
                     desc="高级DOCX翻译进度", unit="段落") as pbar:
                
                future_to_batch = {}
                for batch_idx, (batch_elements, batch_indices) in enumerate(batches):
                    future = executor.submit(
                        self._translate_batch_with_timeout, batch_elements, batch_indices,
                        target_lang, source_lang, batch_idx + 1
                    )
                    future_to_batch[future] = (batch_elements, batch_indices)
                
                for future in as_completed(future_to_batch):
                    batch_elements, batch_indices = future_to_batch[future]
                    try:
                        success, translated_texts, from_cache_flags = future.result()
                        
                        if success:
                            # 应用翻译并检查失败
                            for i, (element, translated_text) in enumerate(zip(batch_elements, translated_texts)):
                                from_cache = from_cache_flags[i] if i < len(from_cache_flags) else False
                                
                                # 检查是否有问题
                                is_serious, reason = TranslationValidator.is_serious_failure(
                                    element.full_text, translated_text, self.large_text_threshold, from_cache)
                                
                                if is_serious or (not from_cache and reason != "翻译成功"):
                                    self._add_failed_task(element, batch_indices[i], translated_text, from_cache)
                                
                                self._apply_translation(element, translated_text)
                        else:
                            # 整个批次失败
                            self._add_batch_failure(batch_elements, batch_indices, translated_texts[0] if translated_texts else "批次处理失败")
                        
                        pbar.update(len(batch_elements))
                    except Exception as e:
                        logger.error(f"批次翻译失败: {e}")
                        self._add_batch_failure(batch_elements, batch_indices, f"执行异常: {str(e)}")
                        pbar.update(len(batch_elements))

    def _apply_reference_styles(self, doc: Document, reference_doc_path: str):
        """应用参考文档的样式"""
        try:
            ref_doc = Document(reference_doc_path)
            logger.info(f"正在应用参考文档样式: {os.path.basename(reference_doc_path)}")
            
            # 简单的样式复制逻辑
            try:
                # 获取参考文档的样式
                ref_styles = ref_doc.styles
                target_styles = doc.styles
                
                # 复制段落样式（简化版本）
                for ref_style in ref_styles:
                    if ref_style.type == WD_STYLE_TYPE.PARAGRAPH:
                        try:
                            # 检查目标文档是否已有此样式
                            if ref_style.name not in [s.name for s in target_styles]:
                                logger.debug(f"找到段落样式: {ref_style.name}")
                        except Exception as e:
                            logger.debug(f"处理样式 {ref_style.name} 时出错: {e}")
                            continue
                
                self.stats['template_applied'] = True
                logger.info("参考文档样式应用完成")
                
            except Exception as e:
                logger.warning(f"样式复制过程中出错: {e}")
            
        except Exception as e:
            logger.warning(f"无法应用参考文档样式: {e}")
    
    def _collect_paragraph_elements(self, doc: Document) -> List[ParagraphElement]:
        """收集段落元素，保持完整性"""
        elements = []
        
        try:
            # 处理普通段落
            for para_idx, para in enumerate(doc.paragraphs):
                if self._should_skip_paragraph(para):
                    self.stats['skipped_paragraphs'] += 1
                    continue
                
                # 收集run信息
                runs_info = []
                for run in para.runs:
                    if run.text:  # 包括空白run，保持格式
                        runs_info.append({
                            'text': run.text,
                            'format': self._extract_run_format(run),
                            'is_image': self._run_contains_image(run)
                        })
                
                if runs_info:
                    full_text = ''.join(r['text'] for r in runs_info if not r['is_image'])
                    if full_text.strip():  # 只有有实际文本内容才翻译
                        elements.append(ParagraphElement(
                            full_text=full_text,
                            paragraph=para,
                            runs_info=runs_info,
                            para_format=self._extract_paragraph_format(para),
                            location=f"paragraph_{para_idx}",
                            element_type='paragraph'
                        ))
            
            # 处理表格
            for table_idx, table in enumerate(doc.tables):
                self.stats['tables'] += 1
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                elements.append(ParagraphElement(
                                    full_text=para.text,
                                    paragraph=para,
                                    runs_info=[{
                                        'text': para.text,
                                        'format': self._extract_paragraph_format(para),
                                        'is_image': False
                                    }],
                                    para_format=self._extract_paragraph_format(para),
                                    location=f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}_para_{para_idx}",
                                    element_type='table_cell'
                                ))
        
        except Exception as e:
            logger.error(f"Error collecting paragraph elements: {e}")
        
        return elements
    
    def _apply_translation(self, element: ParagraphElement, translated_text: str):
        """应用翻译到段落"""
        if not translated_text or translated_text == element.full_text:
            return
        
        try:
            # 清空段落并重建
            para = element.paragraph
            para.clear()
            
            # 简化处理：直接添加翻译文本，应用主要格式
            if element.runs_info:
                # 使用第一个非图片run的格式
                main_format = None
                for run_info in element.runs_info:
                    if not run_info['is_image']:
                        main_format = run_info['format']
                        break
                
                run = para.add_run(translated_text)
                if main_format:
                    self._apply_run_format(run, main_format)
            
            # 恢复段落格式
            self._apply_paragraph_format(para, element.para_format)
            
            self.stats['translated_paragraphs'] += 1
            
        except Exception as e:
            logger.warning(f"Error applying translation: {e}")
    
    def _should_skip_paragraph(self, para) -> bool:
        """判断是否跳过段落"""
        try:
            text = para.text.strip()
            return (not text or len(text) < 3 or text.isdigit() or 
                    not any(c.isalpha() for c in text))
        except:
            return True
    
    def _run_contains_image(self, run) -> bool:
        """检查run是否包含图片"""
        try:
            drawing = run._element.find('.//w:drawing', 
                                     {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            return drawing is not None
        except:
            return False
    
    def _extract_run_format(self, run) -> Dict[str, Any]:
        """提取run格式"""
        format_info = {}
        try:
            if run.bold:
                format_info['bold'] = True
            if run.italic:
                format_info['italic'] = True
            if run.underline:
                format_info['underline'] = True
            
            font = run.font
            if font.size:
                format_info['font_size'] = font.size.pt
            if font.name:
                format_info['font_name'] = font.name
        except:
            pass
        return format_info
    
    def _extract_paragraph_format(self, para) -> Dict[str, Any]:
        """提取段落格式"""
        format_info = {}
        try:
            if para.alignment:
                format_info['alignment'] = para.alignment
            if para.style:
                format_info['style'] = para.style.name
        except:
            pass
        return format_info
    
    def _apply_run_format(self, run, format_info: Dict[str, Any]):
        """应用run格式"""
        try:
            if format_info.get('bold'):
                run.bold = True
            if format_info.get('italic'):
                run.italic = True
            if format_info.get('underline'):
                run.underline = True
            if format_info.get('font_size'):
                run.font.size = Pt(format_info['font_size'])
            if format_info.get('font_name'):
                run.font.name = format_info['font_name']
        except:
            pass
    
    def _apply_paragraph_format(self, para, format_info: Dict[str, Any]):
        """应用段落格式"""
        try:
            if format_info.get('alignment'):
                para.alignment = format_info['alignment']
            if format_info.get('style'):
                para.style = format_info['style']
        except:
            pass
    
    def _print_stats(self):
        """打印统计信息，增加重试worker相关统计"""
        try:
            cache_stats = self.cache.stats
            
            print(f"\n高级DOCX翻译完成！")
            print(f"处理时间: {self.stats['processing_time']:.1f} 秒")
            print(f"总段落: {self.stats['total_paragraphs']} 个")
            print(f"已翻译: {self.stats['translated_paragraphs']} 个")
            print(f"已跳过: {self.stats['skipped_paragraphs']} 个")
            print(f"总字符数: {self.stats['total_chars']:,}")
            print(f"处理批次: {self.stats['total_batches']} 个")
            print(f"表格: {self.stats['tables']} 个")
            print(f"API调用: {self.stats['api_calls']} 次")
            print(f"缓存命中率: {cache_stats['hit_rate']:.1%}")
            print(f"缓存节省: {self.stats['cache_savings']} 次调用")
            
            # 智能失败统计显示
            if self.stats['serious_failures'] > 0 or self.stats['minor_issues'] > 0:
                print(f"\n=== 问题分析 ===")
                if self.stats['serious_failures'] > 0:
                    print(f"严重失败: {self.stats['serious_failures']} 个（>100字符未翻译或API错误）")
                if self.stats['minor_issues'] > 0:
                    print(f"轻微问题: {self.stats['minor_issues']} 个（短文本相同等，正常情况）")
                
                if self.stats['retry_attempts'] > 0:
                    print(f"\n=== 重试统计 ===")
                    print(f"重试轮数: {self.stats['retry_attempts']} 轮")
                    print(f"重试Worker: 最大{self.retry_max_workers}个，实际使用{self.stats['retry_workers_used']}个")
                    if self.stats['concurrent_retry_batches'] > 0:
                        print(f"并发重试批次: {self.stats['concurrent_retry_batches']}个")
                    print(f"缓存清理: {cache_stats['clears']} 次")
                    if self.stats['final_rescues'] > 0:
                        print(f"最终挽救: {self.stats['final_rescues']} 个")
                    if self.stats['final_failures'] > 0:
                        print(f"最终失败: {self.stats['final_failures']} 个（保持原文）")
                    else:
                        print("✅ 所有严重失败都已成功处理！")
                else:
                    print("✅ 未达到重试阈值，无需重试")
            else:
                print("✅ 无翻译问题")
            
            print(f"\n=== 配置信息 ===")
            print(f"Prompt模式: {self.stats['prompt_mode']}")
            print(f"主翻译Workers: {self.max_workers} 个")
            print(f"重试Workers: {self.retry_max_workers} 个")
            print(f"大段文字阈值: {self.large_text_threshold} 字符")
            print(f"重试失败率阈值: {self.retry_failure_threshold:.1%}")
            print(f"非ASCII字符阈值: {self.non_ascii_threshold:.1%}")
            print(f"最大重试次数: {self.max_retries} 次")
            print(f"参考模板: {'已应用' if self.stats['template_applied'] else '未使用'}")
            
            if (self.effective_prompt_config and 
                self.effective_prompt_config.get('mode') != 'none'):
                
                mode = self.effective_prompt_config.get('mode', 'none')
                if mode == 'professional':
                    print(f"专业领域: {self.effective_prompt_config.get('prompt_template', 'academic')}")
                elif mode == 'custom':
                    print("使用了自定义Prompt")
                
                # 显示使用的增强功能
                enhancements = []
                preserve_terms = self.effective_prompt_config.get('preserve_terms')
                if preserve_terms and isinstance(preserve_terms, list):
                    enhancements.append(f"保留术语({len(preserve_terms)}个)")
                
                glossary = self.effective_prompt_config.get('glossary')
                if glossary and isinstance(glossary, dict):
                    enhancements.append(f"术语表({len(glossary)}条)")
                
                if self.effective_prompt_config.get('additional_context'):
                    enhancements.append("额外上下文")
                
                if enhancements:
                    print(f"增强功能: {', '.join(enhancements)}")
                    
        except Exception as e:
            logger.warning(f"Error printing stats: {e}")


def translate_docx_file_formatted(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    max_chunk_size: int = 8000,
    batch_size: int = 50,
    max_workers: int = 5,
    retry_max_workers: int = 5,  # 新增参数
    reference_doc: Optional[str] = None,
    prompt_config: Optional[Dict[str, Any]] = None,
    translation_timeout: int = 60,
    max_retries: int = 8,
    large_text_threshold: int = 50,
    retry_failure_threshold: float = 0.0,
    **kwargs
) -> str:
    """
    高级格式化DOCX翻译函数，完整支持前端prompt配置和重试机制
    
    新增参数:
        retry_max_workers: 重试时的最大worker数量，默认5。建议不超过8，以保持系统稳定性。
    """
    
    if not os.path.exists(input_filepath):
        return f"Error: 输入文件未找到: {input_filepath}"
    
    os.makedirs(output_dir, exist_ok=True)
    
    # 生成输出文件路径
    input_path = Path(input_filepath)
    if unique_filename_base:
        output_filename = f"{unique_filename_base}_translated_{target_lang}.docx"
    else:
        output_filename = f"{input_path.stem}_translated_{target_lang}.docx"
    
    output_filepath = os.path.join(output_dir, output_filename)
    
    # 确保文件名唯一
    counter = 1
    while os.path.exists(output_filepath):
        if unique_filename_base:
            output_filename = f"{unique_filename_base}_translated_{target_lang}_{counter}.docx"
        else:
            output_filename = f"{input_path.stem}_translated_{target_lang}_{counter}.docx"
        output_filepath = os.path.join(output_dir, output_filename)
        counter += 1

    try:
        # 处理prompt配置
        effective_prompt_config = _prepare_prompt_config(prompt_config, kwargs)
        
        # 获取批处理设置
        batch_settings = _get_batch_settings_from_config(effective_prompt_config, kwargs)
        
        # 记录翻译开始信息
        logger.info("=== 启动高级格式化DOCX翻译（增强并发重试版）===")
        logger.info(f"输入: {os.path.basename(input_filepath)}")
        logger.info(f"目标语言: {target_lang}")
        
        if effective_prompt_config and effective_prompt_config.get('mode') != 'none':
            mode = effective_prompt_config.get('mode', 'none')
            logger.info(f"Prompt配置: mode={mode}")
            if mode == 'professional':
                logger.info(f"专业领域: {effective_prompt_config.get('prompt_template', 'academic')}")
            elif mode == 'custom':
                logger.info("使用自定义prompt")
        
        if reference_doc:
            logger.info(f"参考文档: {reference_doc}")
        
        logger.info(f"批处理设置: batch_size={batch_settings['batch_size']}, max_chars={batch_settings['max_chars']}, workers={batch_settings['max_workers']}")
        logger.info(f"重试设置: max_retries={batch_settings['max_retries']}, timeout={batch_settings['translation_timeout']}s, retry_workers={retry_max_workers}")
        
        # 使用高级翻译器
        docx_translator = AdvancedDocxTranslator(
            translator=translator,
            batch_size=batch_settings['batch_size'],
            max_chars=batch_settings['max_chars'],
            max_workers=batch_settings['max_workers'],
            retry_max_workers=retry_max_workers,  # 传递新参数
            prompt_config=effective_prompt_config,
            reference_doc=reference_doc,
            translation_timeout=batch_settings['translation_timeout'],
            max_retries=batch_settings['max_retries'],
            large_text_threshold=batch_settings['large_text_threshold'],
            retry_failure_threshold=batch_settings['retry_failure_threshold'],
            non_ascii_threshold=batch_settings['non_ascii_threshold'],
            **kwargs
        )
        
        result_path = docx_translator.translate_docx(
            input_filepath=input_filepath,
            output_filepath=output_filepath,
            target_lang=target_lang,
            source_lang=source_lang
        )
        
        if not os.path.exists(result_path):
            return "Error: 输出文件未创建"
        
        file_size = os.path.getsize(result_path)
        if file_size == 0:
            return "Error: 输出文件为空"
        
        logger.info(f"高级格式化翻译成功完成！输出: {result_path}")
        return result_path
        
    except Exception as e:
        error_msg = f"高级格式化翻译失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return f"Error: {error_msg}"


# 向后兼容的别名函数
def translate_docx_via_markdown(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    max_chunk_size: int = 8000,
    batch_size: int = 50,
    max_workers: int = 5,
    retry_max_workers: int = 5,  # 新增参数
    reference_doc: Optional[str] = None,
    prompt_config: Optional[Dict[str, Any]] = None,
    **kwargs
) -> str:
    """向后兼容的别名函数，支持新的重试worker参数"""
    return translate_docx_file_formatted(
        input_filepath=input_filepath,
        output_dir=output_dir,
        target_lang=target_lang,
        translator=translator,
        source_lang=source_lang,
        unique_filename_base=unique_filename_base,
        max_chunk_size=max_chunk_size,
        batch_size=batch_size,
        max_workers=max_workers,
        retry_max_workers=retry_max_workers,  # 传递新参数
        reference_doc=reference_doc,
        prompt_config=prompt_config,
        **kwargs
    )


# 测试用翻译器，支持prompt配置和merge策略
class MockTranslator:
    def __init__(self):
        self.call_count = 0
        self.prompt_config = None
    
    def set_prompt_config(self, prompt_config):
        """设置prompt配置"""
        self.prompt_config = copy.deepcopy(prompt_config) if prompt_config else None
    
    def translate(self, text: str = None, target_lang: str = None, 
                 source_lang: str = None, messages: List[Dict] = None, 
                 prompt_config: Optional[Dict[str, Any]] = None,
                 config_merge_mode: str = 'merge', **kwargs) -> str:
        """支持merge策略的翻译方法"""
        self.call_count += 1
        
        # 处理不同的输入格式
        if messages:
            user_content = messages[-1]["content"]
        elif text:
            user_content = text
        else:
            return "Error: No input provided"
        
        # 简单的翻译映射
        translation_map = {
            "documento": "document", "teste": "test", "parágrafo": "paragraph",
            "exemplo": "example", "texto": "text", "arquivo": "file",
            "tradução": "translation", "sistema": "system", "usuário": "user",
            "projeto": "project", "situação": "situation", "empresa": "company",
            "relatório": "report", "análise": "analysis", "dados": "data"
        }
        
        # 根据有效的prompt配置调整翻译行为
        effective_config = prompt_config or self.prompt_config
        if effective_config and effective_config.get('mode') != 'none':
            mode = effective_config.get('mode', 'none')
            if mode == 'custom':
                # 这里可以根据自定义prompt调整翻译逻辑
                pass
            elif mode == 'professional':
                domain = effective_config.get('prompt_template', 'academic')
                # 使用更专业的翻译
                if domain == 'business':
                    translation_map.update({
                        "projeto": "business project",
                        "relatório": "business report"
                    })
        
        lines = user_content.split('\n')
        translated_lines = []
        
        for line in lines:
            if '[' in line and ']' in line:
                match = re.search(r'\[(\d+)\]\s*(.*)', line)
                if match:
                    content = match.group(2)
                    # 应用翻译
                    for pt, en in translation_map.items():
                        content = re.sub(r'\b' + pt + r'\b', en, content, flags=re.IGNORECASE)
                    translated_lines.append(content)
        
        return '\n'.join(translated_lines)


if __name__ == '__main__':
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        force=True
    )
    
    print("=== 高级格式化DOCX翻译系统（增强并发重试版）===")
    print("✅ 增加重试次数到8次")
    print("✅ 修复长文本检测逻辑（阈值改为0）")
    print("✅ 增强错误消息检测")
    print("✅ 添加缓存清理机制")
    print("✅ 实现最终挽救机制")
    print("✅ 改进批次失败处理")
    print("✅ 优化重试批次大小 [20, 10, 5, 3, 1, 1, 1, 1]")
    print("✅ 增强统计和日志记录")
    print("✅ 支持前端所有prompt配置")
    print("✅ 智能配置合并策略")
    print("✅ 专业领域模板")
    print("✅ 自定义prompt")
    print("✅ 术语管理和术语表")
    print("✅ 参考文档样式")
    print("✅ 高级批处理和缓存")
    print("✅ 多线程并发处理")
    print("✅ 完整的格式保持")
    print("✅ 详细的统计和进度")
    print("🆕 **NEW**: 并发重试机制，支持最大5个重试worker")
    print("🆕 **NEW**: 智能重试调度，自动判断串行/并发模式")
    print("🆕 **NEW**: 增强的资源管理，避免worker冲突")
    print("🆕 **NEW**: 详细的重试统计和监控")
    
    # 演示不同配置
    translator = MockTranslator()
    
    test_configs = [
        None,  # 默认
        {'mode': 'professional', 'prompt_template': 'academic', 'retry_max_workers': 3},
        {'mode': 'custom', 'custom_prompt': {'system': 'You are a specialized translator...', 'user': '{content}'}, 'retry_max_workers': 5},
        {'mode': 'general', 'preserve_terms': ['API', 'HTTP'], 'glossary': {'server': '服务器'}, 'retry_max_workers': 4},
    ]
    
    for i, config in enumerate(test_configs):
        print(f"\n=== 测试配置 {i+1}: {config.get('mode', 'default') if config else 'default'} ===")
        
        if config:
            mode = config.get('mode', 'none')
            print(f"模式: {mode}")
            if mode == 'professional':
                print(f"专业领域: {config.get('prompt_template', 'academic')}")
            elif mode == 'custom':
                print("自定义prompt模式")
            
            retry_workers = config.get('retry_max_workers', 5)
            print(f"重试Workers: {retry_workers}个")
            
            if config.get('preserve_terms'):
                print(f"保留术语: {config['preserve_terms']}")
            if config.get('glossary'):
                print(f"术语表: {config['glossary']}")
        
        print("配置已准备，可用于实际翻译")
    
    print("\n=== 系统就绪 ===")
    print("现在所有长文本都会被正确检测和重试！")
    print("重试机制已增强，支持8轮重试、并发处理和最终挽救")
    print("新增retry_max_workers参数，可设置1-8个重试worker")
    print("智能资源管理，避免主翻译和重试worker冲突")
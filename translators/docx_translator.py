# docx_translator.py
import os
import logging
from docx import Document
from typing import Optional, Dict, Any, List, Union, Tuple
import hashlib
from tqdm import tqdm
import re
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import threading
import time
import signal
from enum import Enum
from dataclasses import dataclass

# 简单的日志配置
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 简化的配置常量
DEFAULT_BATCH_SIZE = 50
DEFAULT_MAX_CHARS = 8000
DEFAULT_THREADS = 5
DEFAULT_TRANSLATION_TIMEOUT = 60  # 翻译超时时间（秒）
DEFAULT_MAX_RETRIES = 10           # 最大重试次数
DEFAULT_RETRY_DELAY = 2           # 重试延迟（秒）
DEFAULT_RETRY_BATCH_SIZE = 10     # 重试时的批次大小 - 改为更小的值
DEFAULT_RETRY_THREADS = 5         # 默认重试线程数 - 固定为5

# 全局缓存
_cache = {}
_cache_lock = threading.Lock()

# 失败原因枚举 - 新增
class FailureReason(Enum):
    """失败原因枚举"""
    TIMEOUT = "timeout"
    API_ERROR = "api_error"
    PARSE_ERROR = "parse_error"
    NOT_TRANSLATED = "not_translated"
    EMPTY_RESPONSE = "empty_response"
    CONNECTION_ERROR = "connection_error"
    BATCH_FAILURE = "batch_failure"
    FORMAT_ERROR = "format_error"

# 失败任务数据类 - 新增
@dataclass
class FailedTask:
    """失败任务数据类 - 修复版本，支持结果更新"""
    original_text: str
    paragraph: Any  # 段落对象直接引用
    task_index: int  # 任务索引
    batch_id: int  # 批次ID
    task_id: str  # 任务ID，用于重试时直接更新
    failure_reason: FailureReason
    retry_count: int = 0
    error_message: str = ""
    is_serious: bool = True
    
    def __post_init__(self):
        self.failure_timestamp = time.time()

class TimeoutException(Exception):
    """超时异常"""
    pass

class TranslationValidator:
    """翻译完整性验证器 - 新增"""
    
    # 错误关键词检测
    ERROR_KEYWORDS = [
        # 英文错误关键词
        'timeout', 'readtimeout', 'connecttimeout', 'httptimeout',
        'network error', 'connection error', 'api error', 'service error',
        'translation failed', 'service unavailable', 'request failed',
        'server error', 'bad gateway', 'gateway timeout', 'error:',
        # 中文错误关键词
        '超时', '网络错误', '连接错误', '服务错误', '翻译失败',
        '服务不可用', '请求失败', '服务器错误', '错误:'
    ]
    
    @staticmethod
    def is_error_message(text: str) -> bool:
        """检查是否为错误消息"""
        if not text:
            return True
            
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in TranslationValidator.ERROR_KEYWORDS)
    
    @staticmethod
    def is_serious_failure(original_text: str, translated_text: str, 
                          large_text_threshold: int = 50) -> Tuple[bool, str]:
        """改进的失败检测逻辑"""
        
        # 1. 首先检查明显的错误消息
        if TranslationValidator.is_error_message(translated_text):
            return True, f"API错误消息: {translated_text[:50]}..."
        
        # 2. 检查空值或无效响应
        if not translated_text or translated_text.strip() == "":
            return len(original_text) > 20, f"空翻译结果（原文{len(original_text)}字符）"
        
        # 3. 改进的未翻译检测
        is_unchanged = original_text.strip() == translated_text.strip()
        if is_unchanged:
            # 检查是否可能是正常情况（专有名词、数字、代码等）
            if TranslationValidator._is_likely_untranslatable(original_text):
                return False, "可能是专有名词或代码，保持原文正常"
            
            # 长文本完全未变化才认为是严重失败
            if len(original_text) > large_text_threshold:
                return True, f"长文本未翻译（{len(original_text)}字符）"
            else:
                return False, f"短文本未变化（{len(original_text)}字符，可能正常）"
        
        # 4. 质量检测
        quality_score = TranslationValidator._assess_translation_quality(original_text, translated_text)
        if quality_score < 0.2:
            return True, f"翻译质量过低（分数: {quality_score:.2f}）"
        
        return False, "翻译正常"
    
    @staticmethod
    def _is_likely_untranslatable(text: str) -> bool:
        """判断文本是否可能不需要翻译"""
        text = text.strip()
        
        # 数字、日期、代码
        if re.match(r'^[\d\-\/\.\s\:\(\)]+$', text):
            return True
        
        # URL、邮箱
        if re.match(r'^https?://|.*@.*\..*', text):
            return True
        
        # 专有名词比例
        uppercase_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if uppercase_ratio > 0.3:
            return True
        
        # 常见的不需要翻译的模式
        common_patterns = [
            r'^[A-Z]{2,}$',  # 全大写缩写
            r'^\d+[\.\)]\s*$',  # 列表编号
            r'^[a-zA-Z0-9_\-\.]+$',  # 标识符
        ]
        
        for pattern in common_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    @staticmethod
    def _assess_translation_quality(original: str, translated: str) -> float:
        """评估翻译质量（0-1分数）"""
        if not translated or not original:
            return 0.0
        
        score = 1.0
        
        # 长度检查
        length_ratio = len(translated) / max(len(original), 1)
        if length_ratio < 0.2 or length_ratio > 5.0:
            score -= 0.2
        
        # 字符多样性检查
        if len(set(translated)) / max(len(translated), 1) < 0.05:
            score -= 0.1
        
        # 重复检查
        if len(translated) > 20 and translated.count(translated[:10]) > 5:
            score -= 0.2
        
        return max(0.0, score)

def timeout_handler(signum, frame):
    """超时信号处理器"""
    raise TimeoutException("Translation timeout")

class DocxTranslator:
    """简化的DOCX翻译器类 - 修复版本，支持重试结果更新"""
    
    def __init__(self, translator, target_lang: str, source_lang: Optional[str] = None,
                 translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT,
                 max_retries: int = DEFAULT_MAX_RETRIES,
                 retry_delay: int = DEFAULT_RETRY_DELAY,
                 retry_batch_size: int = DEFAULT_RETRY_BATCH_SIZE,
                 max_retry_workers: int = DEFAULT_RETRY_THREADS,
                 large_text_threshold: int = 50):
        self.translator = translator
        self.target_lang = target_lang
        self.source_lang = source_lang
        self.translation_timeout = translation_timeout
        self.max_retries = max_retries
        self.retry_delay = retry_delay
        self.retry_batch_size = retry_batch_size
        self.max_retry_workers = 5  # 固定为5线程
        self.large_text_threshold = large_text_threshold
        
        self.stats = {
            'success': 0, 
            'failed': 0, 
            'skipped': 0, 
            'timeout': 0, 
            'retried': 0,
            'final_failed': 0,
            'serious_failures': 0,
            'minor_issues': 0,
            'retry_attempts': 0,
            'rescued_tasks': 0,
            'retry_success_rate': 0.0,
            'total_retry_tasks': 0,
            'parallel_retry_used': 0,
            'cache_hits': 0,
            'api_calls': 0
        }
        
        # 时间统计 - 新增
        self.time_stats = {
            'total_start_time': 0,
            'main_translation_time': 0,
            'retry_time': 0,
            'total_time': 0,
            'document_load_time': 0,
            'document_save_time': 0
        }
        
        # 失败任务追踪
        self.failed_tasks = []
        self.failed_tasks_lock = threading.Lock()
        
        # 增强重试策略配置 - 修改为 [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_delays = [1, 2, 4, 8, 12, 16, 20, 25]
    
    def _get_cache_key(self, text: str) -> str:
        """生成简单的缓存键"""
        content = f"{text}_{self.target_lang}_{self.source_lang or 'auto'}"
        return hashlib.md5(content.encode('utf-8')).hexdigest()
    
    def _create_system_prompt(self, prompt_config: Optional[Dict[str, Any]] = None) -> str:
        """创建系统提示 - 简化版"""
        base = f"Translate from {self.source_lang or 'auto-detected language'} to {self.target_lang}."
        
        if not prompt_config:
            return base + "\n\nRules:\n- Translate each [N] line\n- Output only translations\n- Keep formatting"
        
        mode = prompt_config.get('mode', 'none')
        
        # 专业模式
        if mode == 'professional':
            domain = prompt_config.get('prompt_template', 'general')
            domain_prompts = {
                'academic': 'Use academic terminology and formal style.',
                'business': 'Use business terminology and professional tone.',
                'technical': 'Preserve technical accuracy and terminology.',
                'medical': 'Use standard medical terminology.',
                'legal': 'Use precise legal terminology.',
            }
            base = f"You are a {domain} translator. " + base
            if domain in domain_prompts:
                base += f" {domain_prompts[domain]}"
        
        # 自定义模式
        elif mode == 'custom':
            custom_prompt = prompt_config.get('custom_prompt', {})
            if isinstance(custom_prompt, dict) and custom_prompt.get('system'):
                return custom_prompt['system']
        
        # 添加增强功能
        rules = ["\n\nRules:", "- Translate each [N] line", "- Output only translations", "- Keep formatting"]
        
        # 保留术语
        if prompt_config.get('preserve_terms'):
            terms = prompt_config['preserve_terms'][:10]  # 限制数量
            if isinstance(terms, str):
                terms = [t.strip() for t in terms.split(',') if t.strip()]
            if terms:
                rules.append(f"- Keep unchanged: {', '.join(terms)}")
        
        # 术语表
        if prompt_config.get('glossary'):
            glossary = prompt_config['glossary']
            if isinstance(glossary, dict) and len(glossary) <= 5:  # 限制数量
                glossary_text = '; '.join([f"{k}→{v}" for k, v in list(glossary.items())[:5]])
                rules.append(f"- Use glossary: {glossary_text}")
        
        return base + '\n'.join(rules)
    
    def _call_translator_with_timeout(self, text: str, prompt_config: Optional[Dict[str, Any]] = None, 
                                    timeout: Optional[int] = None) -> str:
        """带超时的翻译调用"""
        timeout = timeout or self.translation_timeout
        
        # 使用ThreadPoolExecutor实现超时
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(self._call_translator, text, prompt_config)
            try:
                result = future.result(timeout=timeout)
                return result
            except TimeoutError:
                logger.warning(f"Translation timeout after {timeout} seconds")
                self.stats['timeout'] += 1
                return f"Error: Translation timeout after {timeout} seconds"
            except Exception as e:
                logger.error(f"Translation error: {e}")
                return f"Error: {str(e)}"
    
    def _call_translator(self, text: str, prompt_config: Optional[Dict[str, Any]] = None) -> str:
        """调用翻译器 - 简化版"""
        try:
            # 方式1: 尝试带prompt_config的调用
            if prompt_config:
                try:
                    return self.translator.translate(
                        text=text,
                        target_lang=self.target_lang,
                        source_lang=self.source_lang,
                        prompt_config=prompt_config
                    )
                except TypeError:
                    pass  # 翻译器不支持prompt_config参数
            
            # 方式2: 尝试messages格式
            try:
                system_prompt = self._create_system_prompt(prompt_config)
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ]
                return self.translator.translate(
                    messages=messages,
                    target_lang=self.target_lang,
                    source_lang=self.source_lang
                )
            except TypeError:
                pass
            
            # 方式3: 传统调用
            return self.translator.translate(text, self.target_lang, self.source_lang)
            
        except Exception as e:
            logger.error(f"Translation failed: {e}")
            return f"Error: {str(e)}"
    
    def _extract_translations(self, response: str, expected_count: int) -> List[str]:
        """提取翻译结果"""
        if not response or response.startswith("Error:"):
            return [""] * expected_count
        
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        results = [""] * expected_count
        
        # 尝试匹配编号格式 [1], [2], etc.
        for line in lines:
            match = re.match(r'^\[(\d+)\]\s*(.*)', line)
            if match:
                idx = int(match.group(1)) - 1
                content = match.group(2).strip()
                if 0 <= idx < expected_count and content:
                    results[idx] = content
        
        # 如果没有编号，按顺序分配
        empty_count = sum(1 for r in results if not r)
        if empty_count > 0:
            available_lines = [line for line in lines if not re.match(r'^\[\d+\]', line)]
            j = 0
            for i in range(expected_count):
                if not results[i] and j < len(available_lines):
                    results[i] = available_lines[j].strip()
                    j += 1
        
        return results
    
    def _translate_batch(self, texts: List[str], prompt_config: Optional[Dict[str, Any]] = None,
                        batch_info: Optional[Dict] = None) -> tuple:
        """
        翻译一批文本 - 修复版本，增加缓存标记返回
        返回: (translations, failed_indices, error_message, cache_flags)
        """
        if not texts:
            return [], [], None, []
        
        # 检查缓存
        cached_results = []
        uncached_texts = []
        uncached_indices = []
        cache_flags = [False] * len(texts)  # 跟踪哪些结果来自缓存
        
        with _cache_lock:
            for i, text in enumerate(texts):
                cache_key = self._get_cache_key(text)
                if cache_key in _cache:
                    cached_results.append((i, _cache[cache_key]))
                    cache_flags[i] = True  # 标记为缓存结果
                    self.stats['cache_hits'] += 1
                else:
                    uncached_texts.append(text)
                    uncached_indices.append(i)
        
        # 翻译未缓存的文本
        results = [""] * len(texts)
        failed_indices = []
        error_message = None
        
        # 恢复缓存结果
        for idx, result in cached_results:
            results[idx] = result
        
        if uncached_texts:
            # 创建编号文本
            numbered_texts = [f"[{i+1}] {text}" for i, text in enumerate(uncached_texts)]
            batch_text = "\n".join(numbered_texts)
            
            # 带超时的翻译
            response = self._call_translator_with_timeout(batch_text, prompt_config)
            self.stats['api_calls'] += 1
            
            if response and not response.startswith("Error:"):
                translations = self._extract_translations(response, len(uncached_texts))
                
                # 更新结果和缓存
                with _cache_lock:
                    for i, (orig_idx, translation) in enumerate(zip(uncached_indices, translations)):
                        if translation and translation.strip() and not TranslationValidator.is_error_message(translation):
                            results[orig_idx] = translation.strip()
                            cache_key = self._get_cache_key(uncached_texts[i])
                            _cache[cache_key] = translation.strip()
                        else:
                            # 翻译失败或为空
                            results[orig_idx] = ""
                            failed_indices.append(orig_idx)
            else:
                # 整个批次翻译失败
                error_message = response if response else "Unknown translation error"
                failed_indices = uncached_indices
                for orig_idx in uncached_indices:
                    results[orig_idx] = ""
        
        return results, failed_indices, error_message, cache_flags
    
    def _record_failed_task(self, task: Dict, batch_id: int, failure_reason: FailureReason, 
                          error_msg: str = "", is_serious: bool = True) -> str:
        """记录失败的任务 - 修复版本"""
        task_id = f"task_{batch_id}_{task.get('index', 0)}"
        
        failed_task = FailedTask(
            original_text=task['text'],
            paragraph=task['paragraph'],
            task_index=task.get('index', -1),
            batch_id=batch_id,
            task_id=task_id,
            failure_reason=failure_reason,
            error_message=error_msg,
            is_serious=is_serious,
            retry_count=task.get('retry_count', 0)
        )
        
        with self.failed_tasks_lock:
            self.failed_tasks.append(failed_task)
            
        # 根据严重性更新统计
        if is_serious:
            self.stats['serious_failures'] += 1
            return "serious"
        else:
            self.stats['minor_issues'] += 1
            return "minor"
    
    def _create_batches(self, tasks: List[Dict], batch_size: int = DEFAULT_BATCH_SIZE, 
                       max_chars: int = DEFAULT_MAX_CHARS) -> List[List[Dict]]:
        """创建批次 - 简化版"""
        batches = []
        current_batch = []
        current_chars = 0
        
        for task in tasks:
            text_len = len(task['text'])
            
            # 如果添加当前任务会超出限制，且当前批次不为空，则开始新批次
            if (len(current_batch) >= batch_size or current_chars + text_len > max_chars) and current_batch:
                batches.append(current_batch)
                current_batch = []
                current_chars = 0
            
            current_batch.append(task)
            current_chars += text_len
        
        if current_batch:
            batches.append(current_batch)
        
        return batches
    
    def _process_batch(self, batch_data) -> tuple:
        """处理单个批次 - 修复版本，支持更好的失败检测"""
        batch, batch_id, thread_id, prompt_config = batch_data
        
        try:
            texts = [task['text'] for task in batch]
            logger.info(f"线程-{thread_id}: 处理批次-{batch_id} (共{len(texts)}段落)")
            
            batch_info = {'batch_id': batch_id, 'thread_id': thread_id}
            translations, failed_indices, error_message, cache_flags = self._translate_batch(
                texts, prompt_config, batch_info
            )
            
            # 应用成功的翻译结果
            success_count = 0
            serious_failures = 0
            minor_issues = 0
            
            for i, (task, translation) in enumerate(zip(batch, translations)):
                # 确保任务有唯一ID
                task_index = task.get('index', i)
                
                # 检查翻译结果
                if translation and translation.strip() and not TranslationValidator.is_error_message(translation):
                    # 验证翻译结果
                    is_serious, reason = TranslationValidator.is_serious_failure(
                        task['text'], translation, large_text_threshold=self.large_text_threshold
                    )
                    
                    if not is_serious:
                        try:
                            # 应用翻译到段落
                            para = task['paragraph']
                            para.clear()
                            para.add_run(translation)
                            success_count += 1
                        except Exception as e:
                            logger.warning(f"应用翻译失败: {e}")
                            failure_result = self._record_failed_task(
                                task, batch_id, FailureReason.FORMAT_ERROR, 
                                f"格式应用错误: {str(e)}", True
                            )
                            serious_failures += 1 if failure_result == "serious" else 0
                            minor_issues += 1 if failure_result == "minor" else 0
                    else:
                        # 严重的翻译问题
                        failure_result = self._record_failed_task(
                            task, batch_id, FailureReason.NOT_TRANSLATED, reason, is_serious
                        )
                        serious_failures += 1 if failure_result == "serious" else 0
                        minor_issues += 1 if failure_result == "minor" else 0
                        
                elif i in failed_indices or not translation or TranslationValidator.is_error_message(translation):
                    # 明确的翻译失败
                    if error_message and "timeout" in error_message.lower():
                        failure_reason = FailureReason.TIMEOUT
                        self.stats['timeout'] += 1
                    else:
                        failure_reason = FailureReason.API_ERROR
                    
                    failure_result = self._record_failed_task(
                        task, batch_id, failure_reason, error_message or "翻译失败或空结果", True
                    )
                    serious_failures += 1 if failure_result == "serious" else 0
                    minor_issues += 1 if failure_result == "minor" else 0
            
            logger.info(f"线程-{thread_id}: 批次-{batch_id} 完成，成功: {success_count}/{len(batch)}，严重失败: {serious_failures}，轻微问题: {minor_issues}")
            return len(batch), success_count, serious_failures + minor_issues, error_message, serious_failures
            
        except Exception as e:
            logger.error(f"线程-{thread_id}: 批次-{batch_id} 处理失败: {e}")
            # 记录整个批次失败
            serious_failures = 0
            for task in batch:
                failure_result = self._record_failed_task(
                    task, batch_id, FailureReason.BATCH_FAILURE, f"批次处理异常: {str(e)}", True
                )
                serious_failures += 1 if failure_result == "serious" else 0
            
            return len(batch), 0, len(batch), str(e), serious_failures
    
    def _analyze_failure_patterns(self, failed_tasks: List[FailedTask]) -> Dict[str, Any]:
        """分析失败模式 - 新增"""
        analysis = {
            'total_count': len(failed_tasks),
            'avg_text_length': sum(len(task.original_text) for task in failed_tasks) / max(1, len(failed_tasks)),
            'failure_types': {},
            'has_timeouts': False,
            'has_api_errors': False,
            'consecutive_failures': 0
        }
        
        for task in failed_tasks:
            reason = task.failure_reason.value
            analysis['failure_types'][reason] = analysis['failure_types'].get(reason, 0) + 1
            
            if task.failure_reason == FailureReason.TIMEOUT:
                analysis['has_timeouts'] = True
            elif task.failure_reason == FailureReason.API_ERROR:
                analysis['has_api_errors'] = True
        
        return analysis
    
    def _determine_retry_strategy(self, analysis: Dict[str, Any], retry_count: int) -> Dict[str, Any]:
        """确定重试策略 - 修改为始终使用并行"""
        strategy = {
            'use_concurrent': True,  # 始终使用并行
            'max_workers': 5,  # 固定5线程
            'batch_size_multiplier': 1.0,
            'timeout_multiplier': 1.0,
            'description': '并行重试'
        }
        
        # 根据失败类型调整超时乘数和批次大小，但始终保持并行
        if analysis['has_timeouts']:
            strategy['timeout_multiplier'] = 1.5
            strategy['description'] = '超时优化并行重试'
        
        if analysis['has_api_errors']:
            strategy['batch_size_multiplier'] = 0.5
            strategy['description'] = 'API错误优化并行重试'
        
        if analysis['avg_text_length'] > 1000:
            strategy['timeout_multiplier'] = 2.0
            strategy['description'] = '长文本优化并行重试'
        
        return strategy
    
    def _create_smart_retry_batches(self, failed_tasks: List[FailedTask], retry_count: int) -> List[List[FailedTask]]:
        """创建智能重试批次 - 新增"""
        # 只重试严重失败
        serious_failed_tasks = [task for task in failed_tasks if task.is_serious]
        
        if not serious_failed_tasks:
            return []
        
        # 根据重试次数确定批次大小
        if retry_count < len(self.retry_batch_sizes):
            max_batch_size = self.retry_batch_sizes[retry_count]
        else:
            # 后续重试使用最小批次
            max_batch_size = 1
        
        logger.info(f"第{retry_count + 1}次重试，严重失败任务: {len(serious_failed_tasks)}, 批次大小: {max_batch_size}")
        
        # 按失败原因分组
        tasks_by_reason = {}
        for task in serious_failed_tasks:
            reason = task.failure_reason
            if reason not in tasks_by_reason:
                tasks_by_reason[reason] = []
            tasks_by_reason[reason].append(task)
        
        chunks = []
        
        for reason, tasks in tasks_by_reason.items():
            # 根据失败类型确定批次策略
            if reason == FailureReason.TIMEOUT:
                # 超时失败：减小批次
                batch_size = max(1, max_batch_size // 2)
            elif reason == FailureReason.API_ERROR:
                # API错误：更小批次
                batch_size = max(1, max_batch_size // 3)
            else:
                batch_size = max_batch_size
            
            # 按文本长度排序，长文本优先单独处理
            tasks.sort(key=lambda x: len(x.original_text), reverse=True)
            
            current_chunk = []
            current_chars = 0
            
            for task in tasks:
                text_len = len(task.original_text)
                
                # 超长文本单独处理
                if text_len > DEFAULT_MAX_CHARS // 2:
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = []
                        current_chars = 0
                    chunks.append([task])
                    continue
                
                # 正常批次处理
                if (len(current_chunk) >= batch_size or 
                    current_chars + text_len > DEFAULT_MAX_CHARS) and current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = [task]
                    current_chars = text_len
                else:
                    current_chunk.append(task)
                    current_chars += text_len
            
            if current_chunk:
                chunks.append(current_chunk)
        
        return chunks
    
    def _calculate_adaptive_delay(self, retry_count: int, failure_analysis: Dict[str, Any]) -> float:
        """计算自适应延迟 - 新增"""
        base_delay = self.retry_delays[min(retry_count, len(self.retry_delays) - 1)]
        
        # 根据失败模式调整延迟
        if failure_analysis.get('has_api_errors', False):
            base_delay *= 1.5  # API错误需要更长延迟
        
        if failure_analysis.get('avg_text_length', 0) > 1000:
            base_delay *= 1.2  # 长文本需要更多时间
        
        return base_delay
    
    def _process_retry_batch(self, retry_batch: List[FailedTask], batch_idx: int, retry_count: int,
                           prompt_config: Optional[Dict[str, Any]]) -> int:
        """处理单个重试批次 - 修复超时处理逻辑"""
        logger.info(f"开始重试批次 {batch_idx + 1}，任务数: {len(retry_batch)}")
        
        # 提取文本
        batch_texts = [task.original_text for task in retry_batch]
        
        # 执行重试翻译
        translations, failed_indices, error_message, cache_flags = self._translate_batch(
            batch_texts, prompt_config
        )
        
        logger.info(f"重试批次 {batch_idx + 1} 完成，结果数: {len(translations) if translations else 0}")
        
        # 处理重试结果 - 修复超时处理逻辑
        success_count = 0
        
        if not translations:
            # 整个批次失败
            batch_reason = error_message if error_message else "重试批次失败"
            logger.warning(f"重试批次失败: {batch_reason}")
            
            # 更新重试计数
            for task in retry_batch:
                task.retry_count += 1
                task.error_message = batch_reason
            
            return 0
        
        # 批次成功，检查个别结果并更新原始段落
        for i, task in enumerate(retry_batch):
            if i < len(translations):
                translation = translations[i]
                from_cache = cache_flags[i] if i < len(cache_flags) else False
                
                # 检查翻译结果 - 检查是否超时或错误
                if not translation or TranslationValidator.is_error_message(translation):
                    # 超时或错误 - 不更新段落，保持原文
                    task.retry_count += 1
                    task.error_message = "翻译超时或错误"
                    logger.debug(f"重试失败(超时或错误): {task.original_text[:50]}...")
                    continue
                
                # 检查重试是否成功
                is_serious, reason = TranslationValidator.is_serious_failure(
                    task.original_text, translation, self.large_text_threshold
                )
                
                if not is_serious:
                    # 重试成功 - 核心修复：直接更新原始段落
                    try:
                        paragraph = task.paragraph
                        paragraph.clear()
                        paragraph.add_run(translation)
                        
                        # 更新统计和状态
                        success_count += 1
                        task.retry_count += 1  # 仍然增加重试计数，但标记为成功
                        task.is_serious = False  # 不再是严重问题
                        logger.info(f"✅ 重试救援成功: {task.original_text[:30]}... -> {translation[:30]}...")
                    except Exception as e:
                        logger.warning(f"应用重试翻译失败: {e}")
                        task.retry_count += 1
                        task.error_message = f"格式应用错误: {str(e)}"
                else:
                    # 重试仍然失败
                    logger.debug(f"重试失败: {task.original_text[:50]}... 原因: {reason}")
                    task.retry_count += 1
                    task.error_message = reason or "重试翻译结果仍有问题"
            else:
                # 没有对应的结果
                logger.warning(f"重试结果不足: 任务 {i}, 结果数量 {len(translations)}")
                task.retry_count += 1
                task.error_message = "重试结果缺失"
        
        return success_count
    
    def _execute_concurrent_retry(self, retry_batches: List[List[FailedTask]], 
                                prompt_config: Optional[Dict[str, Any]],
                                strategy: Dict[str, Any]) -> int:
        """并发重试执行 - 修复版本"""
        success_count = 0
        max_workers = 5  # 固定使用5线程
        
        try:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                with tqdm(total=sum(len(batch) for batch in retry_batches), 
                         desc="并行重试", unit="任务") as pbar:
                    
                    # 提交所有重试批次任务
                    future_to_batch = {}
                    for retry_batch_idx, retry_batch in enumerate(retry_batches):
                        future = executor.submit(
                            self._process_retry_batch, retry_batch, retry_batch_idx, 
                            0, prompt_config
                        )
                        future_to_batch[future] = retry_batch
                    
                    # 处理完成的任务
                    for future in as_completed(future_to_batch):
                        retry_batch = future_to_batch[future]
                        try:
                            batch_success = future.result()
                            success_count += batch_success
                            pbar.update(len(retry_batch))
                        except Exception as e:
                            logger.error(f"并发重试批次处理异常: {e}")
                            # 对失败的批次任务增加重试计数
                            for task in retry_batch:
                                task.retry_count += 1
                            pbar.update(len(retry_batch))
            
            self.stats['parallel_retry_used'] += 1
            return success_count
            
        except Exception as e:
            logger.error(f"并发重试执行异常: {e}")
            # 对所有任务增加重试计数
            for batch in retry_batches:
                for task in batch:
                    task.retry_count += 1
            return success_count
    
    def _execute_sequential_retry(self, retry_batches: List[List[FailedTask]], 
                                prompt_config: Optional[Dict[str, Any]]) -> int:
        """串行重试执行 - 保留但不再使用"""
        success_count = 0
        
        with tqdm(total=sum(len(batch) for batch in retry_batches), 
                 desc="串行重试", unit="任务") as pbar:
            
            for retry_batch_idx, retry_batch in enumerate(retry_batches):
                batch_success = self._process_retry_batch(
                    retry_batch, retry_batch_idx, 0, prompt_config
                )
                success_count += batch_success
                pbar.update(len(retry_batch))
        
        return success_count
    
    def _adaptive_retry_strategy(self, prompt_config: Optional[Dict[str, Any]] = None) -> int:
        """自适应重试策略 - 修改为始终使用并行重试"""
        retry_start_time = time.time()  # 开始重试计时
        total_success_count = 0
        
        # 只处理严重失败任务
        serious_tasks = [task for task in self.failed_tasks if task.is_serious]
        if not serious_tasks:
            logger.info("没有严重失败任务需要重试")
            return 0
        
        logger.info(f"开始自适应重试，发现 {len(serious_tasks)} 个严重失败任务")
        self.stats['total_retry_tasks'] = len(serious_tasks)
        
        retry_count = 0
        while retry_count < self.max_retries:
            # 过滤出当前需要重试的任务
            current_retry_tasks = [task for task in self.failed_tasks 
                                 if task.is_serious and task.retry_count <= retry_count]
            
            if not current_retry_tasks:
                logger.info(f"第 {retry_count + 1} 次重试检查：没有(更多)需要重试的任务")
                break
            
            logger.info(f"第 {retry_count + 1} 次重试，处理 {len(current_retry_tasks)} 个严重失败任务")
            
            # 分析失败模式
            failure_analysis = self._analyze_failure_patterns(current_retry_tasks)
            
            # 动态调整策略
            strategy = self._determine_retry_strategy(failure_analysis, retry_count)
            
            # 添加重试延迟
            delay = self._calculate_adaptive_delay(retry_count, failure_analysis)
            if delay > 0:
                logger.info(f"智能延迟: {delay:.1f}秒")
                time.sleep(delay)
            
            # 创建重试批次
            retry_batches = self._create_smart_retry_batches(current_retry_tasks, retry_count)
            
            if not retry_batches:
                logger.info(f"第 {retry_count + 1} 次重试：没有批次需要处理")
                break
            
            # 始终使用并行重试 - 修改点
            logger.info(f"启用并发重试，worker数量: 5")
            success_count = self._execute_concurrent_retry(
                retry_batches, prompt_config, strategy
            )
            
            total_success_count += success_count
            self.stats['rescued_tasks'] += success_count
            self.stats['retry_attempts'] += 1
            
            # 更新失败任务列表 - 移除不再严重的任务
            self.failed_tasks = [task for task in self.failed_tasks if task.is_serious]
            
            logger.info(f"第 {retry_count + 1} 次重试完成，恢复 {success_count} 个任务，当前剩余 {len(self.failed_tasks)} 个严重失败")
            
            retry_count += 1
        
        # 计算重试时间和成功率
        self.time_stats['retry_time'] = time.time() - retry_start_time
        if self.stats['total_retry_tasks'] > 0:
            self.stats['retry_success_rate'] = self.stats['rescued_tasks'] / self.stats['total_retry_tasks']
        
        logger.info(f"重试阶段完成，用时: {self.time_stats['retry_time']:.1f}秒, "
                   f"成功率: {self.stats['retry_success_rate']:.1%}")
        
        return total_success_count
    
    def translate_file(self, input_path: str, output_path: str, 
                      prompt_config: Optional[Dict[str, Any]] = None,
                      batch_size: int = DEFAULT_BATCH_SIZE,
                      max_chars: int = DEFAULT_MAX_CHARS,
                      max_workers: int = DEFAULT_THREADS) -> str:
        """翻译DOCX文件 - 修复版本，支持重试结果更新"""
        try:
            # 开始总计时
            self.time_stats['total_start_time'] = time.time()
            
            logger.info(f"开始翻译: {input_path} -> {output_path}")
            
            # 加载文档
            if not os.path.exists(input_path):
                return f"Error: File not found: {input_path}"
            
            doc_load_start = time.time()
            doc = Document(input_path)
            self.time_stats['document_load_time'] = time.time() - doc_load_start
            
            # 收集翻译任务
            tasks = []
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text and len(text) > 2 and any(c.isalpha() for c in text):
                    tasks.append({'text': text, 'paragraph': para, 'index': idx})
                else:
                    self.stats['skipped'] += 1
            
            if not tasks:
                logger.info("No content to translate")
                doc.save(output_path)
                self.time_stats['total_time'] = time.time() - self.time_stats['total_start_time']
                return output_path
            
            logger.info(f"发现 {len(tasks)} 段落需要翻译")
            
            # 重置统计和失败任务列表
            self.stats.update({
                'success': 0, 'failed': 0, 'timeout': 0, 'retried': 0, 'final_failed': 0,
                'serious_failures': 0, 'minor_issues': 0, 'retry_attempts': 0,
                'rescued_tasks': 0, 'retry_success_rate': 0.0, 'total_retry_tasks': 0,
                'parallel_retry_used': 0, 'cache_hits': 0, 'api_calls': 0
            })
            self.failed_tasks = []
            
            # 创建批次
            batches = self._create_batches(tasks, batch_size, max_chars)
            logger.info(f"创建了 {len(batches)} 个批次")
            
            # 清理缓存
            with _cache_lock:
                _cache.clear()
            
            # === 第一阶段：多线程批量翻译 ===
            logger.info("=== 阶段1：初始批量翻译 ===")
            main_translation_start = time.time()
            
            with ThreadPoolExecutor(max_workers=min(max_workers, len(batches))) as executor:
                # 准备批次数据
                batch_data_list = [
                    (batch, i+1, i+1, prompt_config) 
                    for i, batch in enumerate(batches)
                ]
                
                # 提交任务
                future_to_batch = {
                    executor.submit(self._process_batch, batch_data): batch_data[0]
                    for batch_data in batch_data_list
                }
                
                # 处理结果
                with tqdm(total=len(tasks), desc="翻译进度", unit="段落") as pbar:
                    for future in as_completed(future_to_batch):
                        total, success, failed, error, serious_failures = future.result()
                        self.stats['success'] += success
                        self.stats['failed'] += failed
                        pbar.update(total)
                        
                        if error:
                            logger.warning(f"批次错误: {error}")
            
            self.time_stats['main_translation_time'] = time.time() - main_translation_start
            logger.info(f"主翻译完成，耗时: {self.time_stats['main_translation_time']:.1f}秒")
            
            # === 第二阶段：重试失败的任务 ===
            serious_failures = len([task for task in self.failed_tasks if task.is_serious])
            if serious_failures > 0:
                logger.info(f"=== 阶段2：重试 {serious_failures} 个严重失败任务 ===")
                retry_start_time = time.time()
                
                # 执行自适应重试策略
                retry_success_count = self._adaptive_retry_strategy(prompt_config)
                
                self.time_stats['retry_time'] = time.time() - retry_start_time
                logger.info(f"重试阶段完成，恢复 {retry_success_count} 个任务，耗时: {self.time_stats['retry_time']:.1f}秒")
                
                # 更新统计
                self.stats['success'] += retry_success_count
                self.stats['failed'] -= retry_success_count
                self.stats['retried'] += retry_success_count
                
                # 统计最终失败任务
                final_failures = len([task for task in self.failed_tasks if task.is_serious])
                self.stats['final_failed'] = final_failures
                
                if final_failures > 0:
                    logger.warning(f"重试后仍有 {final_failures} 个严重失败任务保持原文")
                    # 对最终失败的任务保持原文
                    for task in self.failed_tasks:
                        if task.is_serious:
                            try:
                                paragraph = task.paragraph
                                paragraph.clear()
                                paragraph.add_run(task.original_text)  # 保持原文
                            except Exception as e:
                                logger.error(f"Failed to restore original text: {e}")
            else:
                logger.info("=== 阶段2：没有需要重试的严重失败任务 ===")
            
            # 保存文档
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            save_start = time.time()
            doc.save(output_path)
            self.time_stats['document_save_time'] = time.time() - save_start
            
            # 计算总时间
            self.time_stats['total_time'] = time.time() - self.time_stats['total_start_time']
            
            # 输出详细统计
            total = len(tasks)
            success_rate = (self.stats['success'] / total * 100) if total > 0 else 0
            
            logger.info(f"翻译完成: 成功率 {success_rate:.1f}%")
            logger.info(f"总耗时: {self.time_stats['total_time']:.1f}秒 (主翻译: {self.time_stats['main_translation_time']:.1f}s, "
                       f"重试: {self.time_stats['retry_time']:.1f}s, 文档加载: {self.time_stats['document_load_time']:.1f}s, "
                       f"文档保存: {self.time_stats['document_save_time']:.1f}s)")
            
            print(f"\n✅ 翻译完成!")
            print(f"📄 输出: {output_path}")
            print(f"📊 翻译统计:")
            print(f"  ✨ 成功翻译: {self.stats['success']} 段落")
            print(f"  🔄 重试恢复: {self.stats['retried']} 段落")
            print(f"  ⏰ 超时失败: {self.stats['timeout']} 段落")
            print(f"  ❌ 最终失败: {self.stats['final_failed']} 段落")
            print(f"  ⏭️ 跳过: {self.stats['skipped']} 段落")
            print(f"  🎯 成功率: {success_rate:.1f}%")
            print(f"  ⏱️ 总用时: {self.time_stats['total_time']:.1f}秒")
            print(f"     - 主翻译: {self.time_stats['main_translation_time']:.1f}秒")
            print(f"     - 重试: {self.time_stats['retry_time']:.1f}秒")
            print(f"     - 文档加载: {self.time_stats['document_load_time']:.1f}秒")
            print(f"     - 文档保存: {self.time_stats['document_save_time']:.1f}秒")
            
            if self.stats['retry_attempts'] > 0:
                print(f"  🔁 重试统计: {self.stats['retry_attempts']}轮，成功率{self.stats['retry_success_rate']:.1%}")
            
            if self.stats['final_failed'] > 0:
                print(f"  ⚠️ 注意: {self.stats['final_failed']} 段落保持原文")
            
            return output_path
            
        except Exception as e:
            logger.exception(f"翻译错误: {e}")
            return f"Error: {str(e)}"

# 便捷函数
def translate_docx_file(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    chunk_size: Optional[int] = None,
    prompt_config: Optional[Dict[str, Any]] = None,
    translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT,
    max_retries: int = DEFAULT_MAX_RETRIES,
    retry_delay: int = DEFAULT_RETRY_DELAY,
    retry_batch_size: int = DEFAULT_RETRY_BATCH_SIZE,
    max_retry_workers: int = 5,  # 固定为5
    large_text_threshold: int = 50,
    **kwargs
) -> str:
    """
    翻译DOCX文件 - 修复版本，支持重试结果更新
    
    Args:
        input_filepath: 输入文件路径
        output_dir: 输出目录
        target_lang: 目标语言
        translator: 翻译器对象
        source_lang: 源语言
        unique_filename_base: 自定义文件名基础
        chunk_size: 批次大小
        prompt_config: prompt配置
        translation_timeout: 翻译超时时间（秒）
        max_retries: 最大重试次数
        retry_delay: 重试延迟（秒）
        retry_batch_size: 重试批次大小
        max_retry_workers: 重试线程数 (固定为5)
        large_text_threshold: 大文本阈值
    
    Returns:
        输出文件路径或错误信息
    """
    
    # 准备输出路径
    os.makedirs(output_dir, exist_ok=True)
    filename = os.path.basename(input_filepath)
    name, ext = os.path.splitext(filename)
    
    if unique_filename_base:
        output_filename = f"{unique_filename_base}_{target_lang}{ext}"
    else:
        output_filename = f"{name}_translated_{target_lang}{ext}"
    
    output_path = os.path.join(output_dir, output_filename)
    
    # 创建翻译器实例
    docx_translator = DocxTranslator(
        translator, 
        target_lang, 
        source_lang,
        translation_timeout=translation_timeout,
        max_retries=max_retries,
        retry_delay=retry_delay,
        retry_batch_size=retry_batch_size,
        max_retry_workers=5,  # 固定为5
        large_text_threshold=large_text_threshold
    )
    
    # 设置批次大小
    batch_size = chunk_size or DEFAULT_BATCH_SIZE
    
    # 执行翻译
    return docx_translator.translate_file(
        input_path=input_filepath,
        output_path=output_path,
        prompt_config=prompt_config,
        batch_size=batch_size
    )

def translate_docx_file_advanced(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    chunk_size: Optional[int] = None,
    prompt_config: Optional[Dict[str, Any]] = None,
    preserve_formatting: bool = True,
    quality_check: bool = False,
    translation_timeout: int = DEFAULT_TRANSLATION_TIMEOUT,
    max_retries: int = DEFAULT_MAX_RETRIES,
    retry_delay: int = DEFAULT_RETRY_DELAY,
    retry_batch_size: int = DEFAULT_RETRY_BATCH_SIZE,
    max_retry_workers: int = 5,  # 固定为5
    large_text_threshold: int = 50,
    **kwargs
) -> str:
    """高级DOCX翻译 - 修复版本，支持重试结果更新"""
    
    # 如果启用质量检查，调整配置
    if quality_check and prompt_config:
        enhanced_config = prompt_config.copy()
        enhanced_config['advanced_mode'] = True
        prompt_config = enhanced_config
    
    return translate_docx_file(
        input_filepath=input_filepath,
        output_dir=output_dir,
        target_lang=target_lang,
        translator=translator,
        source_lang=source_lang,
        unique_filename_base=unique_filename_base,
        chunk_size=chunk_size,
        prompt_config=prompt_config,
        translation_timeout=translation_timeout,
        max_retries=max_retries,
        retry_delay=retry_delay,
        retry_batch_size=retry_batch_size,
        max_retry_workers=5,  # 固定为5
        large_text_threshold=large_text_threshold,
        **kwargs
    )

# 测试用模拟翻译器 - 增强版
class MockTranslator:
    """模拟翻译器 - 支持超时和随机失败测试"""
    
    def __init__(self, failure_rate: float = 0.1, timeout_rate: float = 0.05):
        self.failure_rate = failure_rate
        self.timeout_rate = timeout_rate
        self.call_count = 0
    
    def translate(self, text=None, target_lang=None, source_lang=None, 
                 messages=None, prompt_config=None, **kwargs):
        self.call_count += 1
        
        # 模拟超时
        import random
        if random.random() < self.timeout_rate:
            time.sleep(70)  # 超过默认超时时间
        
        # 模拟随机失败
        if random.random() < self.failure_rate:
            raise Exception(f"Mock translation failure #{self.call_count}")
        
        # 模拟正常延迟
        time.sleep(random.uniform(0.1, 0.5))
        
        # 处理输入
        if messages and isinstance(messages, list):
            content = messages[-1]["content"]
        elif isinstance(text, str):
            content = text
        else:
            return "Error: Invalid input"
        
        # 简单的翻译模拟
        lines = content.split('\n')
        results = []
        
        for line in lines:
            if '[' in line and ']' in line:
                # 提取编号内容
                match = re.match(r'^\[(\d+)\]\s*(.*)', line.strip())
                if match:
                    content_part = match.group(2)
                    results.append(f"[{match.group(1)}] [{target_lang}] {content_part}")
                else:
                    results.append(f"[{target_lang}] {line}")
            else:
                results.append(f"[{target_lang}] {line}")
        
        return '\n'.join(results) if results else f"[{target_lang}] {content}"

if __name__ == '__main__':
    print("=== 增强高效DOCX翻译器（修复重试结果更新版）===")
    print("🚀 新特性：")
    print("✅ 直接更新重试结果 - 不再存在'很多没有翻译'问题")
    print("✅ 始终使用并行重试(5线程) - 大幅提高重试效率")
    print("✅ 增强失败检测 - 减少误判，提高成功率")
    print("✅ 完整时间统计 - 主翻译/重试/加载/保存各阶段耗时")
    print("✅ 智能批次管理 - 根据失败原因调整批次大小 [10, 5, 2, 1, 1, 1, 1, 1]")
    print("✅ 自适应延迟 - 智能调整重试间隔")
    print("✅ 安全超时处理 - 确保超时和错误情况下保持原文")
    
    # 测试翻译功能
    translator = MockTranslator(failure_rate=0.2, timeout_rate=0.1)
    
    # 测试翻译功能
    docx_translator = DocxTranslator(
        translator, 
        "中文", 
        translation_timeout=5,  # 短超时用于测试
        max_retries=2
    )
    
    test_texts = ["Hello world", "How are you?", "Good morning", "This is a test"]
    
    print(f"\n=== 功能测试 ===")
    results, failed_indices, error, _ = docx_translator._translate_batch(test_texts)
    for i, (original, translated) in enumerate(zip(test_texts, results)):
        status = "❌ FAILED" if i in failed_indices else "✅ SUCCESS"
        print(f"{original} -> {translated} {status}")
    
    print(f"\n✅ 增强版翻译器就绪！支持超时检测和并行重试！")
import os
import shutil
import logging
import subprocess
import tempfile
from typing import Optional, List, Dict, Any, Tuple, Union
from pathlib import Path
from tqdm import tqdm
import re
import time
import hashlib
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
from enum import Enum
from copy import deepcopy
from dataclasses import dataclass
from collections import OrderedDict
from contextlib import contextmanager

logger = logging.getLogger(__name__)

# 配置合并策略枚举
class ConfigMergeStrategy(Enum):
    MERGE = "merge"              # 默认：智能合并前端和实例配置
    OVERRIDE = "override"        # 前端配置完全覆盖
    INSTANCE_ONLY = "instance_only"  # 只使用实例配置

# 失败原因枚举
class FailureReason(Enum):
    """失败原因枚举"""
    TIMEOUT = "timeout"
    API_ERROR = "api_error"
    PARSE_ERROR = "parse_error"
    NOT_TRANSLATED = "not_translated"
    EMPTY_RESPONSE = "empty_response"
    CONNECTION_ERROR = "connection_error"
    BATCH_FAILURE = "batch_failure"

@dataclass
class FailedTask:
    """统一的失败任务数据类 - 修复版本，支持结果更新"""
    original_text: str
    original_index: int  # 在原始列表中的索引
    original_task_id: str  # 新增：原始任务ID，用于更新翻译结果
    element: Any  # 任务元素对象
    failure_reason: FailureReason
    retry_count: int = 0
    error_message: str = ""
    is_serious: bool = True
    
    def __post_init__(self):
        self.failure_timestamp = time.time()

# 全局配置合并策略
_global_merge_strategy = ConfigMergeStrategy.MERGE
_instance_prompt_config = None

def set_config_merge_strategy(strategy: Union[str, ConfigMergeStrategy]):
    """设置全局配置合并策略"""
    global _global_merge_strategy
    if isinstance(strategy, str):
        try:
            _global_merge_strategy = ConfigMergeStrategy(strategy)
        except ValueError:
            logger.warning(f"Invalid merge strategy: {strategy}, using default 'merge'")
            _global_merge_strategy = ConfigMergeStrategy.MERGE
    else:
        _global_merge_strategy = strategy
    
    logger.info(f"Markdown-based translator config merge strategy set to: {_global_merge_strategy.value}")

def set_instance_prompt_config(config: Optional[Dict[str, Any]]):
    """设置实例级别的prompt配置"""
    global _instance_prompt_config
    _instance_prompt_config = deepcopy(config) if config else None
    if config:
        logger.info(f"Markdown-based translator instance prompt config set: mode={config.get('mode', 'none')}")

def get_effective_config(frontend_config: Optional[Dict[str, Any]] = None, 
                        merge_strategy: Optional[Union[str, ConfigMergeStrategy]] = None) -> Optional[Dict[str, Any]]:
    """获取当前有效配置，基于合并策略"""
    global _instance_prompt_config, _global_merge_strategy
    
    # 确定使用的合并策略
    if merge_strategy:
        if isinstance(merge_strategy, str):
            try:
                strategy = ConfigMergeStrategy(merge_strategy)
            except ValueError:
                strategy = _global_merge_strategy
        else:
            strategy = merge_strategy
    else:
        strategy = _global_merge_strategy
    
    logger.debug(f"Markdown translator using merge strategy: {strategy.value}")
    
    # 根据策略合并配置
    if strategy == ConfigMergeStrategy.INSTANCE_ONLY:
        effective_config = deepcopy(_instance_prompt_config) if _instance_prompt_config else None
        logger.debug("Using instance-only configuration for markdown translator")
    
    elif strategy == ConfigMergeStrategy.OVERRIDE:
        effective_config = deepcopy(frontend_config) if frontend_config else deepcopy(_instance_prompt_config)
        logger.debug("Using override strategy for markdown translator (frontend takes precedence)")
    
    else:  # ConfigMergeStrategy.MERGE
        effective_config = _merge_prompt_configs(_instance_prompt_config, frontend_config)
        logger.debug("Using intelligent merge strategy for markdown translator")
    
    # 标准化最终配置
    if effective_config:
        effective_config = _normalize_prompt_config(effective_config)
        logger.debug(f"Effective markdown config: mode={effective_config.get('mode', 'none')}")
    
    return effective_config

def _merge_prompt_configs(instance_config: Optional[Dict[str, Any]], 
                         frontend_config: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """核心配置合并逻辑 - 智能合并前端和实例配置"""
    
    if not instance_config and not frontend_config:
        return None
    
    if not instance_config:
        return deepcopy(frontend_config)
    
    if not frontend_config:
        return deepcopy(instance_config)
    
    # 开始智能合并
    merged = deepcopy(instance_config)
    
    logger.debug(f"Merging markdown configs - Instance: {instance_config.get('mode', 'none')}, Frontend: {frontend_config.get('mode', 'none')}")
    
    # 1. 模式合并：前端优先
    if 'mode' in frontend_config:
        merged['mode'] = frontend_config['mode']
    
    # 2. 基础字段合并：前端优先
    basic_fields = ['prompt_template', 'professional_domain', 'custom_prompt', 
                   'custom_system_prompt', 'custom_user_prompt', 'system', 'user']
    
    for field in basic_fields:
        if field in frontend_config:
            merged[field] = frontend_config[field]
    
    # 3. 智能字段合并
    
    # preserve_terms：术语列表合并去重
    merged_terms = set()
    for config in [instance_config, frontend_config]:
        terms = config.get('preserve_terms', [])
        if isinstance(terms, str):
            terms = [term.strip() for term in terms.split(',') if term.strip()]
        elif isinstance(terms, list):
            terms = [str(term).strip() for term in terms if str(term).strip()]
        merged_terms.update(terms)
    
    if merged_terms:
        merged['preserve_terms'] = list(merged_terms)
        logger.debug(f"Merged preserve_terms for markdown: {len(merged_terms)} unique terms")
    
    # glossary：术语表字典合并（前端优先）
    merged_glossary = {}
    if instance_config.get('glossary') and isinstance(instance_config['glossary'], dict):
        merged_glossary.update(instance_config['glossary'])
    if frontend_config.get('glossary') and isinstance(frontend_config['glossary'], dict):
        merged_glossary.update(frontend_config['glossary'])  # 前端优先覆盖
    
    if merged_glossary:
        merged['glossary'] = merged_glossary
        logger.debug(f"Merged glossary for markdown: {len(merged_glossary)} entries")
    
    # additional_context：上下文信息拼接
    contexts = []
    for config in [instance_config, frontend_config]:
        context = config.get('additional_context', '').strip()
        if context:
            contexts.append(context)
    
    if contexts:
        merged['additional_context'] = ' | '.join(contexts)
        logger.debug(f"Merged additional_context for markdown: {len(merged['additional_context'])} chars")
    
    # 4. 批处理设置合并：取更优的值
    batch_fields = {
        'max_units_per_chunk': max,
        'max_chars_per_chunk': max,
        'max_chunk_size': max,
        'batch_size': max,
        'max_workers': max,
        'thread_count': max
    }
    
    for field, merge_func in batch_fields.items():
        values = []
        for config in [instance_config, frontend_config]:
            if field in config and isinstance(config[field], (int, float)):
                values.append(config[field])
        if values:
            merged[field] = merge_func(values)
    
    # 5. 质量和功能设置：逻辑OR合并
    quality_fields = ['ensure_consistency', 'quality_level', 'preserve_formatting', 
                     'advanced_mode', 'detailed_logging', 'use_cache', 'use_template']
    
    for field in quality_fields:
        if (instance_config.get(field) or frontend_config.get(field)):
            merged[field] = frontend_config.get(field, instance_config.get(field))
    
    logger.debug(f"Markdown config merge completed: {len(merged)} fields in final config")
    return merged

def _normalize_prompt_config(config: Dict[str, Any]) -> Dict[str, Any]:
    """标准化prompt配置格式，确保与前端格式兼容"""
    normalized = config.copy()
    
    # 确保mode字段存在
    if 'mode' not in normalized:
        if 'custom_prompt' in normalized:
            normalized['mode'] = 'custom'
        elif 'prompt_template' in normalized or 'professional_domain' in normalized:
            normalized['mode'] = 'professional'
        elif any(k in normalized for k in ['preserve_terms', 'glossary', 'additional_context']):
            normalized['mode'] = 'general'
        else:
            normalized['mode'] = 'none'
    
    # 处理保留术语 - 支持逗号分隔的字符串（前端格式）
    preserve_terms = normalized.get('preserve_terms')
    if preserve_terms:
        if isinstance(preserve_terms, str):
            # 前端格式：逗号分隔的字符串
            terms_list = [term.strip() for term in preserve_terms.split(',') if term.strip()]
            normalized['preserve_terms'] = terms_list
        elif isinstance(preserve_terms, list):
            # 确保列表中的字符串都是清理过的
            normalized['preserve_terms'] = [str(term).strip() for term in preserve_terms if str(term).strip()]
    
    # 处理术语表 - 确保是字典格式
    glossary = normalized.get('glossary')
    if glossary and not isinstance(glossary, dict):
        logger.warning(f"Glossary should be a dictionary, got {type(glossary)}, ignoring")
        normalized.pop('glossary', None)
    
    # 处理自定义prompt
    if normalized.get('mode') == 'custom':
        custom_prompt = normalized.get('custom_prompt', {})
        if not custom_prompt or not isinstance(custom_prompt, dict):
            # 检查是否有分离的system和user字段
            system_prompt = normalized.get('custom_system_prompt', normalized.get('system'))
            user_prompt = normalized.get('custom_user_prompt', normalized.get('user'))
            
            if system_prompt:
                normalized['custom_prompt'] = {
                    'system': system_prompt,
                    'user': user_prompt or 'Please translate the following content to {target_lang}:\n\n{content}'
                }
            else:
                logger.warning("Custom mode selected but no valid custom prompt provided, falling back to general mode")
                normalized['mode'] = 'general'
    
    # 处理专业模板 - 前端使用 'professional_domain' 字段
    if normalized.get('mode') == 'professional':
        domain = normalized.get('professional_domain', normalized.get('prompt_template', 'academic'))
        normalized['prompt_template'] = domain
    
    return normalized

class TranslationValidator:
    """翻译完整性验证器 - 修复版本，减少误判"""
    
    # 扩展的错误关键词检测
    ERROR_KEYWORDS = [
        # 英文错误关键词
        'timeout', 'readtimeout', 'connecttimeout', 'httptimeout',
        'network error', 'connection error', 'api error', 'service error',
        'translation failed', 'service unavailable', 'request failed',
        'server error', 'bad gateway', 'gateway timeout',
        # 中文错误关键词
        '超时', '网络错误', '连接错误', '服务错误', '翻译失败',
        '服务不可用', '请求失败', '服务器错误'
    ]
    
    @staticmethod
    def is_error_message(text: str) -> bool:
        """检查是否为错误消息"""
        if not text:
            return True
            
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in TranslationValidator.ERROR_KEYWORDS)
    
    @staticmethod
    def is_serious_failure(original_text: str, translated_text: str, 
                          large_text_threshold: int = 50, from_cache: bool = False) -> Tuple[bool, str]:
        """改进的失败检测逻辑 - 减少误判"""
        
        # 1. 首先检查明显的错误消息
        if TranslationValidator.is_error_message(translated_text):
            return True, f"API错误消息: {translated_text[:50]}..."
        
        # 2. 检查空值或无效响应
        if not translated_text or translated_text.strip() == "":
            return len(original_text) > 20, f"空翻译结果（原文{len(original_text)}字符）"
        
        # 3. 改进的未翻译检测 - 更宽松的判断
        is_unchanged = original_text.strip() == translated_text.strip()
        if is_unchanged:
            # 检查是否可能是正常情况（专有名词、数字、代码等）
            if TranslationValidator._is_likely_untranslatable(original_text):
                return False, "可能是专有名词或代码，保持原文正常"
            
            # 长文本完全未变化才认为是严重失败 - 阈值调整为100
            if len(original_text) > 100:  # 更宽松的阈值
                return True, f"长文本未翻译（{len(original_text)}字符）"
            else:
                return False, f"短文本未变化（{len(original_text)}字符，可能正常）"
        
        # 4. 质量检测：检查翻译是否合理 - 降低阈值
        quality_score = TranslationValidator._assess_translation_quality(original_text, translated_text)
        if quality_score < 0.2:  # 降低阈值，减少误判
            return True, f"翻译质量过低（分数: {quality_score:.2f}）"
        
        # 5. 缓存结果也需要验证（修复原有问题）
        if from_cache and len(original_text) > 100:  # 更宽松的阈值
            # 即使来自缓存，长文本也要检查
            if is_unchanged or quality_score < 0.3:
                return True, f"缓存中的问题翻译（{len(original_text)}字符）"
        
        return False, "翻译正常"
    
    @staticmethod
    def _is_likely_untranslatable(text: str) -> bool:
        """判断文本是否可能不需要翻译 - 扩展识别范围"""
        text = text.strip()
        
        # 数字、日期、代码
        if re.match(r'^[\d\-\/\.\s\:\(\)]+$', text):
            return True
        
        # URL、邮箱
        if re.match(r'^https?://|.*@.*\..*', text):
            return True
        
        # 代码块标识
        if text.startswith('```') or text.startswith('```'):
            return True
        
        # 专有名词比例 - 降低阈值
        uppercase_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if uppercase_ratio > 0.3:  # 更宽松
            return True
        
        # 常见的不需要翻译的模式
        common_patterns = [
            r'^[A-Z]{2,}$',  # 全大写缩写
            r'^\d+[\.\)]\s*$',  # 列表编号
            r'^[a-zA-Z0-9_\-\.]+$',  # 标识符
        ]
        
        for pattern in common_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    @staticmethod
    def _assess_translation_quality(original: str, translated: str) -> float:
        """评估翻译质量（0-1分数）- 更宽松的评估"""
        if not translated or not original:
            return 0.0
        
        score = 1.0
        
        # 长度检查 - 更宽松
        length_ratio = len(translated) / max(len(original), 1)
        if length_ratio < 0.2 or length_ratio > 5.0:  # 更宽松的范围
            score -= 0.2  # 减少扣分
        
        # 字符多样性检查 - 更宽松
        if len(set(translated)) / max(len(translated), 1) < 0.05:  # 更宽松的阈值
            score -= 0.1  # 减少扣分
        
        # 重复检查 - 更宽松
        if len(translated) > 20 and translated.count(translated[:10]) > 5:  # 更严格的条件
            score -= 0.2
        
        return max(0.0, score)

class SmartCache:
    """智能LRU缓存，支持prompt配置差异化和失败缓存清理"""
    
    def __init__(self, max_size: int = 1000):
        self._cache = OrderedDict()
        self.max_size = max_size
        self._lock = threading.RLock()
        self._hits = 0
        self._misses = 0
        self._clears = 0  # 新增清理统计
    
    def get(self, key: str) -> Optional[str]:
        if not key or not isinstance(key, str):
            return None
            
        with self._lock:
            if key in self._cache:
                value = self._cache.pop(key)
                self._cache[key] = value
                self._hits += 1
                return value
            self._misses += 1
            return None
    
    def put(self, key: str, value: str):
        if not key or not isinstance(key, str) or not isinstance(value, str):
            return
            
        with self._lock:
            if key in self._cache:
                self._cache.pop(key)
            elif len(self._cache) >= self.max_size:
                self._cache.popitem(last=False)
            self._cache[key] = value
    
    def remove(self, key: str) -> bool:
        """移除指定缓存项"""
        if not key or not isinstance(key, str):
            return False
            
        with self._lock:
            if key in self._cache:
                self._cache.pop(key)
                self._clears += 1
                return True
            return False
    
    def clear_pattern(self, pattern: str) -> int:
        """清理匹配模式的缓存项"""
        if not pattern:
            return 0
            
        with self._lock:
            keys_to_remove = [key for key in self._cache.keys() if pattern in key]
            for key in keys_to_remove:
                self._cache.pop(key)
            self._clears += len(keys_to_remove)
            return len(keys_to_remove)
    
    def clear(self):
        with self._lock:
            cleared_count = len(self._cache)
            self._cache.clear()
            self._hits = 0
            self._misses = 0
            self._clears += cleared_count
            
    @property
    def stats(self) -> Dict[str, Any]:
        with self._lock:
            total = self._hits + self._misses
            return {
                'size': len(self._cache),
                'hits': self._hits,
                'misses': self._misses,
                'clears': self._clears,
                'hit_rate': self._hits / total if total > 0 else 0
            }

class MarkdownBasedDocxTranslator:
    """基于Markdown转换的DOCX翻译器 - 修复重试结果更新版本"""
    
    def __init__(self, translator, batch_size: int = 50, max_chunk_size: int = 8000, 
                 max_workers: int = 5, retry_max_workers: int = 5, use_cache: bool = True, 
                 template_path: Optional[str] = None, mode: str = 'optimized', 
                 max_units_per_chunk: int = 100, prompt_config=None, 
                 config_merge_mode: Optional[Union[str, ConfigMergeStrategy]] = None,
                 translation_timeout: int = 60, max_retries: int = 8,
                 large_text_threshold: int = 50, retry_failure_threshold: float = 0.0,
                 **kwargs):
        self.translator = translator
        self.batch_size = batch_size
        self.max_chunk_size = max_chunk_size
        self.max_workers = max_workers
        self.retry_max_workers = retry_max_workers  # 重试专用worker数量
        self.use_cache = use_cache
        self.translation_cache = SmartCache(1000) if use_cache else None
        self.template_path = template_path
        self.mode = mode
        self.max_units_per_chunk = max_units_per_chunk
        self.translation_timeout = translation_timeout
        self.max_retries = max_retries
        self.large_text_threshold = large_text_threshold  # 修改为50
        self.retry_failure_threshold = retry_failure_threshold
        
        # 验证worker配置
        if self.retry_max_workers < 1:
            self.retry_max_workers = 1
            logger.warning("retry_max_workers不能小于1，已重置为1")
        
        if self.retry_max_workers > 8:
            logger.warning(f"retry_max_workers={self.retry_max_workers}可能过高，建议不超过8")
        
        # 失败任务追踪
        self.failed_tasks: List[FailedTask] = []
        self.failed_tasks_lock = threading.Lock()
        
        # 增强重试策略配置 - 改为 [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]
        self.retry_delays = [1, 2, 4, 8, 12, 16, 20, 25]  # 更合理的延迟
        
        # 时间统计 - 新增完整时间追踪
        self.time_stats = {
            'total_start_time': 0,
            'main_translation_time': 0,
            'retry_time': 0,
            'total_time': 0,
            'conversion_time': 0,
            'formatting_time': 0
        }
        
        # 配置合并策略支持
        self.config_merge_mode = config_merge_mode
        self.original_translator_config = None
        
        # 使用配置合并策略处理prompt配置
        self.effective_prompt_config = get_effective_config(prompt_config, config_merge_mode)
        
        if self.effective_prompt_config:
            self.prompt_template = self.effective_prompt_config.get('prompt_template')
            self.custom_prompt = self.effective_prompt_config.get('custom_prompt')
            self.preserve_terms = self.effective_prompt_config.get('preserve_terms')
            self.glossary = self.effective_prompt_config.get('glossary')
            self.additional_context = self.effective_prompt_config.get('additional_context')
            
            # 更新批处理参数
            if 'max_chunk_size' in self.effective_prompt_config:
                self.max_chunk_size = self.effective_prompt_config['max_chunk_size']
            if 'batch_size' in self.effective_prompt_config:
                self.batch_size = self.effective_prompt_config['batch_size']
            if 'max_workers' in self.effective_prompt_config:
                self.max_workers = self.effective_prompt_config['max_workers']
            
            logger.info(f"MarkdownBasedDocxTranslator initialized with merged config: mode={self.effective_prompt_config.get('mode')}")
        else:
            self.prompt_template = None
            self.custom_prompt = None
            self.preserve_terms = None
            self.glossary = None
            self.additional_context = None
        
        self.stats = {
            'total_elements': 0,
            'translated_elements': 0,
            'skipped_tables': 0,
            'api_calls': 0,
            'cache_hits': 0,
            'total_texts_translated': 0,
            'batch_translations': 0,
            'single_translations': 0,
            'headers_translated': 0,
            'paragraphs_translated': 0,
            'list_items_translated': 0,
            'quotes_translated': 0,
            'portuguese_simple_strategy_used': 0,
            'docx_marks_cleaned': 0,
            'formatting_applied': False,
            'custom_prompt_used': False,
            'config_merge_strategy_used': config_merge_mode or _global_merge_strategy.value,
            'order_validation_passed': False,
            'batches_reordered': 0,
            'sequence_preserved': True,
            # 重试相关统计 - 增强版
            'serious_failures': 0,
            'minor_issues': 0,
            'retry_attempts': 0,
            'final_failures': 0,
            'final_rescues': 0,
            'retry_workers_used': 0,
            'concurrent_retry_batches': 0,
            'cache_clears': 0,
            'retry_strategy_used': 'none',
            'retry_success_rate': 0.0,  # 新增：重试成功率
            'total_retry_tasks': 0,     # 新增：总重试任务数
            'rescued_tasks': 0,         # 新增：成功救援任务数
            'parallel_retry_used': 0,   # 新增：并行重试使用次数
        }

    def set_config_merge_strategy(self, strategy: Union[str, ConfigMergeStrategy]):
        """设置当前实例的配置合并策略"""
        self.config_merge_mode = strategy
        logger.info(f"Updated config merge strategy for markdown translator instance: {strategy}")

    def get_effective_config(self) -> Optional[Dict[str, Any]]:
        """获取当前实例的有效配置"""
        return self.effective_prompt_config

    @contextmanager
    def _translator_config_context(self):
        """线程安全的配置上下文管理器"""
        try:
            self._apply_prompt_config_to_translator()
            yield
        except Exception as e:
            logger.error(f"Error in translator config context: {e}")
            raise
        finally:
            self._restore_translator_config()

    def _apply_prompt_config_to_translator(self):
        """应用prompt配置到翻译器 - 使用合并后的配置"""
        if self.effective_prompt_config and hasattr(self.translator, 'set_prompt_config'):
            try:
                # 保存翻译器的原始配置
                self.original_translator_config = getattr(self.translator, 'prompt_config', None)
                self.translator.set_prompt_config(self.effective_prompt_config)
                logger.info("Applied merged prompt config to translator in MarkdownBasedDocxTranslator")
                return True
            except Exception as e:
                logger.warning(f"Failed to apply merged prompt config to translator: {e}")
        return False

    def _restore_translator_config(self):
        """恢复翻译器的原始配置"""
        if self.effective_prompt_config and hasattr(self.translator, 'set_prompt_config'):
            try:
                if self.original_translator_config is not None:
                    self.translator.set_prompt_config(self.original_translator_config)
                else:
                    # 如果原来没有配置，清除当前配置
                    if hasattr(self.translator, 'prompt_config'):
                        self.translator.prompt_config = None
                logger.debug("Restored translator config in MarkdownBasedDocxTranslator")
            except Exception as e:
                logger.warning(f"Failed to restore translator config: {e}")

    def _is_portuguese(self, source_lang: Optional[str]) -> bool:
        """判断是否为葡萄牙语"""
        if not source_lang:
            return False
        source_lang_lower = source_lang.lower().strip()
        portuguese_variants = [
            'pt', 'portuguese', 'portugues', 'português',
            'pt-br', 'pt-pt', 'portugal', 'brasil', 'brazil',
            'por', 'portugu',
        ]
        
        if source_lang_lower in portuguese_variants:
            return True
        
        portuguese_keywords = ['portugu', 'brasil', 'brazil']
        for keyword in portuguese_keywords:
            if keyword in source_lang_lower:
                return True
        
        return False

    def _clear_failed_task_cache(self, task: FailedTask, target_lang: str, source_lang: Optional[str]):
        """清理失败任务的缓存"""
        try:
            cache_key = self._get_text_hash(task.original_text, target_lang, source_lang or 'auto')
            if self.translation_cache and self.translation_cache.remove(cache_key):
                self.stats['cache_clears'] += 1
                logger.debug(f"清理失败任务缓存: {task.original_text[:50]}...")
        except Exception as e:
            logger.warning(f"Error clearing failed task cache: {e}")

    def _should_trigger_retry(self, total_processed: int) -> bool:
        """判断是否应该触发重试"""
        with self.failed_tasks_lock:
            serious_failures = [task for task in self.failed_tasks if task.is_serious]
            
            if not serious_failures:
                return False
            
            # 严重失败率超过阈值才重试
            failure_rate = len(serious_failures) / max(1, total_processed)
            should_retry = failure_rate >= self.retry_failure_threshold
            
            logger.info(f"严重失败: {len(serious_failures)}/{total_processed} ({failure_rate:.1%}), "
                       f"阈值: {self.retry_failure_threshold:.1%}, 是否重试: {should_retry}")
            
            return should_retry

    def _add_failed_task(self, original_text: str, original_task_id: str, original_index: int, 
                        reason: str, from_cache: bool = False):
        """添加失败任务，智能判断是否为严重失败 - 修复版本，添加task_id"""
        is_serious, detailed_reason = TranslationValidator.is_serious_failure(
            original_text, reason, self.large_text_threshold, from_cache)
        
        with self.failed_tasks_lock:
            failed_task = FailedTask(
                original_text=original_text,
                original_index=original_index,
                original_task_id=original_task_id,  # 关键：保存原始任务ID
                element={'text': original_text, 'index': original_index, 'task_id': original_task_id},
                failure_reason=FailureReason.API_ERROR if is_serious else FailureReason.NOT_TRANSLATED,
                error_message=detailed_reason,
                is_serious=is_serious
            )
            self.failed_tasks.append(failed_task)
            
            # 更新统计
            if is_serious:
                self.stats['serious_failures'] += 1
                logger.debug(f"严重失败: {detailed_reason[:50]}...")
            else:
                self.stats['minor_issues'] += 1
                logger.debug(f"轻微问题: {detailed_reason[:50]}...")

    def _add_batch_failure(self, batch_data: List[Tuple[str, str]], batch_indices: List[int], reason: str):
        """添加批次级失败，所有任务都标记为严重失败 - 修复版本"""
        with self.failed_tasks_lock:
            for (task_id, text), original_index in zip(batch_data, batch_indices):
                failed_task = FailedTask(
                    original_text=text,
                    original_index=original_index,
                    original_task_id=task_id,  # 关键：保存原始任务ID
                    element={'text': text, 'index': original_index, 'task_id': task_id},
                    failure_reason=FailureReason.BATCH_FAILURE,
                    error_message=reason,
                    is_serious=True
                )
                self.failed_tasks.append(failed_task)
                self.stats['serious_failures'] += 1
            logger.warning(f"批次失败: {len(batch_data)} 个任务, 原因: {reason}")

    def _analyze_failure_patterns(self, tasks: List[FailedTask]) -> Dict[str, Any]:
        """分析失败模式"""
        analysis = {
            'total_count': len(tasks),
            'avg_text_length': sum(len(task.original_text) for task in tasks) / len(tasks),
            'failure_types': {},
            'has_timeouts': False,
            'has_api_errors': False,
            'consecutive_failures': 0
        }
        
        for task in tasks:
            reason = task.failure_reason.value
            analysis['failure_types'][reason] = analysis['failure_types'].get(reason, 0) + 1
            
            if task.failure_reason == FailureReason.TIMEOUT:
                analysis['has_timeouts'] = True
            elif task.failure_reason == FailureReason.API_ERROR:
                analysis['has_api_errors'] = True
        
        return analysis

    def _determine_retry_strategy(self, analysis: Dict[str, Any], retry_count: int) -> Dict[str, Any]:
        """确定重试策略 - 简化并行判断"""
        strategy = {
            'use_concurrent': True,  # 默认使用并行
            'concurrent_threshold': 2,  # 简化阈值
            'max_workers': self.retry_max_workers,
            'batch_size_multiplier': 1.0,
            'timeout_multiplier': 1.0,
            'description': '并行重试'
        }
        
        # 如果有超时，降低并发度但仍使用并行
        if analysis['has_timeouts']:
            strategy['max_workers'] = max(2, strategy['max_workers'] // 2)
            strategy['timeout_multiplier'] = 1.5
            strategy['description'] = '超时优化并行重试'
        
        # 如果有API错误，使用更小批次但仍并行
        if analysis['has_api_errors']:
            strategy['batch_size_multiplier'] = 0.5
            strategy['description'] = 'API错误优化并行重试'
        
        # 大文本处理 - 仍使用并行
        if analysis['avg_text_length'] > 1000:
            strategy['concurrent_threshold'] = 1  # 任何数量都并行
            strategy['timeout_multiplier'] = 2.0
            strategy['description'] = '长文本优化并行重试'
        
        # 只有在worker数量为1时才串行
        if strategy['max_workers'] == 1:
            strategy['use_concurrent'] = False
            strategy['description'] = '串行重试'
        
        return strategy

    def _should_use_concurrent_retry(self, retry_batches):
        """判断是否应该使用并发重试 - 简化逻辑，默认并行"""
        # 简化判断：只要有2个以上批次就并行
        return len(retry_batches) >= 2

    def _create_smart_retry_batches(self, failed_tasks: List[FailedTask], retry_count: int) -> List[List[FailedTask]]:
        """创建智能重试批次，只处理严重失败"""
        # 只重试严重失败
        serious_failed_tasks = [task for task in failed_tasks if task.is_serious]
        
        if not serious_failed_tasks:
            return []
        
        # 根据重试次数确定批次大小
        if retry_count < len(self.retry_batch_sizes):
            max_batch_size = self.retry_batch_sizes[retry_count]
        else:
            # 后续重试使用最小批次
            max_batch_size = 1
        
        logger.info(f"第{retry_count + 1}次重试，严重失败任务: {len(serious_failed_tasks)}, 批次大小: {max_batch_size}")
        
        # 按失败原因分组
        grouped_tasks = self._group_tasks_by_failure_type(serious_failed_tasks)
        
        chunks = []
        
        for failure_type, tasks in grouped_tasks.items():
            # 根据失败类型确定批次策略
            if failure_type == FailureReason.TIMEOUT:
                # 超时失败：减小批次
                batch_size = max(1, max_batch_size // 2)
            elif failure_type == FailureReason.API_ERROR:
                # API错误：更小批次
                batch_size = max(1, max_batch_size // 3)
            else:
                batch_size = max_batch_size
            
            # 按文本长度排序，长文本优先单独处理
            tasks.sort(key=lambda x: len(x.original_text), reverse=True)
            
            current_chunk = []
            current_chars = 0
            
            for task in tasks:
                text_len = len(task.original_text)
                
                # 超长文本单独处理
                if text_len > self.max_chunk_size // 2:
                    if current_chunk:
                        chunks.append(current_chunk)
                        current_chunk = []
                        current_chars = 0
                    chunks.append([task])
                    continue
                
                # 正常批次处理
                if (len(current_chunk) >= batch_size or 
                    current_chars + text_len > self.max_chunk_size) and current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = [task]
                    current_chars = text_len
                else:
                    current_chunk.append(task)
                    current_chars += text_len
            
            if current_chunk:
                chunks.append(current_chunk)
        
        return chunks

    def _group_tasks_by_failure_type(self, tasks: List[FailedTask]) -> Dict[FailureReason, List[FailedTask]]:
        """按失败类型分组任务"""
        groups = {}
        for task in tasks:
            reason = task.failure_reason
            if reason not in groups:
                groups[reason] = []
            groups[reason].append(task)
        return groups

    def _calculate_adaptive_delay(self, retry_count: int, failure_analysis: Dict[str, Any]) -> float:
        """计算自适应延迟"""
        base_delay = self.retry_delays[min(retry_count, len(self.retry_delays) - 1)]
        
        # 根据失败模式调整延迟
        if failure_analysis.get('has_api_errors', False):
            base_delay *= 1.5  # API错误需要更长延迟
        
        if failure_analysis.get('avg_text_length', 0) > 1000:
            base_delay *= 1.2  # 长文本需要更多时间
        
        return base_delay

    def _selective_cache_cleanup(self, tasks: List[FailedTask], target_lang: str, source_lang: Optional[str]):
        """选择性缓存清理"""
        if not self.translation_cache:
            return
        
        cleanup_count = 0
        for task in tasks:
            cache_key = self._get_text_hash(task.original_text, target_lang, source_lang or 'auto')
            if self.translation_cache.remove(cache_key):
                cleanup_count += 1
        
        if cleanup_count > 0:
            self.stats['cache_clears'] += cleanup_count
            logger.info(f"清理了 {cleanup_count} 个失败任务的缓存")

    def _adaptive_retry_strategy(self, target_lang: str, source_lang: Optional[str], 
                               translated_results: Dict[str, str]):
        """自适应重试策略主控制器 - 修复版本，支持结果更新"""
        retry_count = 0
        retry_start_time = time.time()  # 开始重试计时
        
        logger.info(f"开始自适应重试，目标更新 {len(translated_results)} 个翻译结果")
        
        while retry_count < self.max_retries:
            with self.failed_tasks_lock:
                serious_tasks = [task for task in self.failed_tasks 
                               if task.is_serious and task.retry_count <= retry_count]
            
            if not serious_tasks:
                logger.info(f"第 {retry_count + 1} 次重试检查：没有严重失败任务需要重试")
                break
            
            logger.info(f"第 {retry_count + 1} 次重试，处理 {len(serious_tasks)} 个严重失败任务")
            self.stats['total_retry_tasks'] += len(serious_tasks)
            
            # 分析失败模式
            failure_analysis = self._analyze_failure_patterns(serious_tasks)
            
            # 动态调整策略
            strategy = self._determine_retry_strategy(failure_analysis, retry_count)
            self.stats['retry_strategy_used'] = strategy['description']
            
            logger.info(f"重试策略: {strategy['description']}")
            
            # 清理相关缓存
            self._selective_cache_cleanup(serious_tasks, target_lang, source_lang)
            
            # 清空当前重试轮次的失败任务
            with self.failed_tasks_lock:
                self.failed_tasks = [task for task in self.failed_tasks 
                                   if not task.is_serious or task.retry_count > retry_count]
            
            # 添加重试延迟
            delay = self._calculate_adaptive_delay(retry_count, failure_analysis)
            if delay > 0:
                logger.info(f"智能延迟: {delay:.1f}秒")
                time.sleep(delay)
            
            # 创建重试批次
            retry_batches = self._create_smart_retry_batches(serious_tasks, retry_count)
            
            if not retry_batches:
                logger.info(f"第 {retry_count + 1} 次重试：没有批次需要处理")
                break
            
            # 简化并行重试判断 - 优先使用并行
            use_concurrent = self._should_use_concurrent_retry(retry_batches)
            actual_workers = min(strategy['max_workers'], len(retry_batches)) if use_concurrent else 1
            self.stats['retry_workers_used'] = actual_workers
            
            # 执行重试 - 传递翻译结果引用
            if use_concurrent:
                logger.info(f"启用并发重试，worker数量: {actual_workers}")
                self.stats['concurrent_retry_batches'] = len(retry_batches)
                self.stats['parallel_retry_used'] += 1
                self._execute_concurrent_retry(retry_batches, target_lang, source_lang, 
                                             retry_count, strategy, translated_results)
            else:
                logger.info("使用串行重试")
                self._execute_sequential_retry(retry_batches, target_lang, source_lang, 
                                             retry_count, strategy, translated_results)
            
            self.stats['retry_attempts'] += 1
            retry_count += 1
        
        # 统计最终失败的任务
        with self.failed_tasks_lock:
            final_serious_failures = [task for task in self.failed_tasks if task.is_serious]
            self.stats['final_failures'] = len(final_serious_failures)
            if final_serious_failures:
                logger.warning(f"自适应重试后仍有 {len(final_serious_failures)} 个严重失败任务")
        
        # 计算重试时间和成功率
        self.time_stats['retry_time'] = time.time() - retry_start_time
        if self.stats['total_retry_tasks'] > 0:
            self.stats['retry_success_rate'] = self.stats['rescued_tasks'] / self.stats['total_retry_tasks']
        
        logger.info(f"重试阶段完成，用时: {self.time_stats['retry_time']:.1f}秒, "
                   f"成功率: {self.stats['retry_success_rate']:.1%}")

    def _execute_concurrent_retry(self, retry_batches: List[List[FailedTask]], target_lang: str, 
                                source_lang: Optional[str], retry_count: int, strategy: Dict[str, Any],
                                translated_results: Dict[str, str]):
        """并发重试执行 - 修复版本，支持结果更新"""
        try:
            max_workers = strategy['max_workers']
            timeout_multiplier = strategy.get('timeout_multiplier', 1.0)
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                with tqdm(total=sum(len(batch) for batch in retry_batches), 
                         desc=f"第{retry_count + 1}次重试(并发)", unit="任务") as pbar:
                    
                    # 提交所有重试批次任务
                    future_to_batch = {}
                    for retry_batch_idx, retry_batch in enumerate(retry_batches):
                        future = executor.submit(
                            self._process_retry_batch, retry_batch, target_lang, source_lang, 
                            retry_batch_idx, retry_count, timeout_multiplier, translated_results
                        )
                        future_to_batch[future] = retry_batch
                    
                    # 处理完成的任务
                    for future in as_completed(future_to_batch):
                        retry_batch = future_to_batch[future]
                        try:
                            future.result(timeout=self.translation_timeout * timeout_multiplier + 30)
                            pbar.update(len(retry_batch))
                        except Exception as e:
                            logger.error(f"并发重试批次处理异常: {e}")
                            # 处理异常情况
                            with self.failed_tasks_lock:
                                for task in retry_batch:
                                    new_task = FailedTask(
                                        original_text=task.original_text,
                                        original_index=task.original_index,
                                        original_task_id=task.original_task_id,
                                        element=task.element,
                                        failure_reason=FailureReason.BATCH_FAILURE,
                                        error_message=f"并发执行异常: {str(e)}",
                                        is_serious=True
                                    )
                                    new_task.retry_count = retry_count + 1
                                    self.failed_tasks.append(new_task)
                            pbar.update(len(retry_batch))
                            
        except Exception as e:
            logger.error(f"并发重试执行异常: {e}")
            # 回退到串行重试
            logger.info("回退到串行重试")
            self._execute_sequential_retry(retry_batches, target_lang, source_lang, 
                                         retry_count, strategy, translated_results)

    def _execute_sequential_retry(self, retry_batches: List[List[FailedTask]], target_lang: str, 
                                source_lang: Optional[str], retry_count: int, strategy: Dict[str, Any],
                                translated_results: Dict[str, str]):
        """串行重试执行 - 修复版本，支持结果更新"""
        timeout_multiplier = strategy.get('timeout_multiplier', 1.0)
        
        with tqdm(total=sum(len(batch) for batch in retry_batches), 
                 desc=f"第{retry_count + 1}次重试", unit="任务") as pbar:
            
            for retry_batch_idx, retry_batch in enumerate(retry_batches):
                self._process_retry_batch(retry_batch, target_lang, source_lang, 
                                        retry_batch_idx, retry_count, timeout_multiplier, translated_results)
                pbar.update(len(retry_batch))

    def _process_retry_batch(self, retry_batch: List[FailedTask], target_lang: str, 
                           source_lang: Optional[str], batch_idx: int, retry_count: int, 
                           timeout_multiplier: float, translated_results: Dict[str, str]):
        """处理单个重试批次 - 核心修复：直接更新翻译结果"""
        logger.info(f"开始重试批次 {batch_idx + 1}，任务数: {len(retry_batch)}")
        
        # 提取文本和索引
        batch_texts = [task.original_text for task in retry_batch]
        batch_indices = [task.original_index for task in retry_batch]
        
        # 执行重试翻译
        success, retry_results, cache_flags = self._translate_batch_with_timeout(
            batch_texts, batch_indices, target_lang, source_lang, -1)
        
        logger.info(f"重试批次 {batch_idx + 1} 完成，成功: {success}, 结果数: {len(retry_results) if retry_results else 0}")
        
        # 处理重试结果 - 关键修复：直接更新原始翻译结果
        if not success:
            # 整个批次失败
            batch_reason = retry_results[0] if retry_results else "重试批次失败"
            logger.warning(f"重试批次失败: {batch_reason}")
            
            with self.failed_tasks_lock:
                for task in retry_batch:
                    new_task = FailedTask(
                        original_text=task.original_text,
                        original_index=task.original_index,
                        original_task_id=task.original_task_id,
                        element=task.element,
                        failure_reason=FailureReason.BATCH_FAILURE,
                        error_message=batch_reason,
                        is_serious=True
                    )
                    new_task.retry_count = retry_count + 1
                    self.failed_tasks.append(new_task)
        else:
            # 批次成功，检查个别结果并更新翻译结果
            for i, task in enumerate(retry_batch):
                if i < len(retry_results):
                    result = retry_results[i]
                    from_cache = cache_flags[i] if i < len(cache_flags) else False
                    
                    # 检查重试是否成功
                    is_valid, reason = self._enhanced_result_validation(task.original_text, result, from_cache)
                    
                    if is_valid and not TranslationValidator.is_error_message(result):
                        # 重试成功 - 核心修复：直接更新原始翻译结果
                        original_task_id = task.original_task_id
                        if original_task_id in translated_results:
                            translated_results[original_task_id] = result
                            self.stats['rescued_tasks'] += 1
                            logger.info(f"✅ 重试救援成功: {task.original_text[:30]}... -> {result[:30]}...")
                        else:
                            logger.warning(f"原始任务ID {original_task_id} 不在翻译结果中")
                    else:
                        # 重试仍然失败
                        logger.debug(f"重试失败，保持原文: {task.original_text[:50]}... 原因: {reason}")
                        
                        with self.failed_tasks_lock:
                            new_task = FailedTask(
                                original_text=task.original_text,
                                original_index=task.original_index,
                                original_task_id=task.original_task_id,
                                element=task.element,
                                failure_reason=FailureReason.API_ERROR,
                                error_message=reason,
                                is_serious=True
                            )
                            new_task.retry_count = retry_count + 1
                            self.failed_tasks.append(new_task)
                else:
                    # 没有对应的结果
                    logger.warning(f"重试结果不足: 任务 {i}, 结果数量 {len(retry_results)}")
                    with self.failed_tasks_lock:
                        new_task = FailedTask(
                            original_text=task.original_text,
                            original_index=task.original_index,
                            original_task_id=task.original_task_id,
                            element=task.element,
                            failure_reason=FailureReason.PARSE_ERROR,
                            error_message="重试结果缺失",
                            is_serious=True
                        )
                        new_task.retry_count = retry_count + 1
                        self.failed_tasks.append(new_task)

    def _final_retry_remaining_tasks(self, target_lang: str, source_lang: Optional[str],
                                   translated_results: Dict[str, str]):
        """最终处理剩余失败任务 - 修复版本，支持结果更新"""
        with self.failed_tasks_lock:
            remaining_tasks = [task for task in self.failed_tasks if task.is_serious]
        
        if not remaining_tasks:
            logger.info("没有剩余失败任务需要最终处理")
            return
        
        logger.info(f"最终处理 {len(remaining_tasks)} 个剩余失败任务")
        
        success_count = 0
        
        with tqdm(total=len(remaining_tasks), desc="最终挽救", unit="任务") as pbar:
            for task in remaining_tasks:
                try:
                    # 清理缓存
                    self._clear_failed_task_cache(task, target_lang, source_lang)
                    
                    # 单独翻译，增加延迟避免频率限制
                    time.sleep(1)
                    
                    logger.debug(f"最终处理任务: {task.original_text[:50]}...")
                    
                    result = self._translate_single_text(task.original_text, target_lang, source_lang)
                    
                    self.stats['api_calls'] += 1
                    
                    if (result and not TranslationValidator.is_error_message(result) and 
                        result.strip() != task.original_text.strip()):
                        
                        # 进一步验证翻译质量
                        is_valid, reason = self._enhanced_result_validation(task.original_text, result, False)
                        
                        if is_valid:
                            # 核心修复：直接更新原始翻译结果
                            original_task_id = task.original_task_id
                            if original_task_id in translated_results:
                                translated_results[original_task_id] = result
                                success_count += 1
                                self.stats['final_rescues'] += 1
                                self.stats['rescued_tasks'] += 1
                                logger.info(f"✅ 最终挽救成功: {task.original_text[:30]}... -> {result[:30]}...")
                            else:
                                logger.warning(f"原始任务ID {original_task_id} 不在翻译结果中")
                        else:
                            logger.debug(f"最终处理结果仍不合格: {reason}")
                            
                    else:
                        logger.debug(f"最终处理仍失败: {task.original_text[:50]}...")
                        
                except Exception as e:
                    logger.error(f"最终处理异常: {e}")
                
                pbar.update(1)
        
        if success_count > 0:
            logger.info(f"✅ 最终挽救成功 {success_count}/{len(remaining_tasks)} 个任务")
        else:
            logger.warning(f"❌ 最终挽救失败，{len(remaining_tasks)} 个任务保持原文")

    def _translate_batch_with_timeout(self, batch_texts: List[str], batch_indices: List[int],
                                    target_lang: str, source_lang: Optional[str], 
                                    batch_num: int) -> Tuple[bool, List[str], List[bool]]:
        """带超时的批次翻译，返回结果和缓存标记"""
        try:
            # 使用Future实现超时控制
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._translate_batch_core, batch_texts, target_lang, source_lang, batch_num)
                
                try:
                    batch_results, from_cache_flags = future.result(timeout=self.translation_timeout)
                    return True, batch_results, from_cache_flags
                    
                except TimeoutError:
                    logger.warning(f"批次 {batch_num} 翻译超时 ({self.translation_timeout}秒)")
                    failure_reasons = [f"翻译超时 ({self.translation_timeout}秒)"] * len(batch_texts)
                    cache_flags = [False] * len(batch_texts)
                    return False, failure_reasons, cache_flags
                    
        except Exception as e:
            logger.error(f"批次 {batch_num} 翻译异常: {e}")
            failure_reasons = [f"翻译异常: {str(e)}"] * len(batch_texts)
            cache_flags = [False] * len(batch_texts)
            return False, failure_reasons, cache_flags

    def _translate_batch_core(self, batch_texts: List[str], target_lang: str, 
                            source_lang: Optional[str], batch_num: int) -> Tuple[List[str], List[bool]]:
        """核心批次翻译逻辑"""
        results = []
        from_cache_flags = []
        
        # 检查缓存
        uncached_texts = []
        uncached_indices = []
        
        for i, text in enumerate(batch_texts):
            cache_key = self._get_text_hash(text, target_lang, source_lang or 'auto')
            if self.translation_cache:
                cached = self.translation_cache.get(cache_key)
                if cached:
                    results.append(cached)
                    from_cache_flags.append(True)
                    continue
            
            uncached_texts.append(text)
            uncached_indices.append(i)
            results.append("")
            from_cache_flags.append(False)
        
        # 翻译未缓存的文本
        if uncached_texts:
            try:
                if len(uncached_texts) == 1:
                    # 单文本翻译
                    translated = self._translate_single_text(uncached_texts[0], target_lang, source_lang)
                    translated_results = [translated]
                else:
                    # 批量翻译
                    batch_data = [(f"text_{i}", text) for i, text in enumerate(uncached_texts)]
                    translated_results_dict = self._translate_text_batch(batch_data, target_lang, source_lang)
                    translated_results = [translated_results_dict.get(f"text_{i}", text) 
                                        for i, text in enumerate(uncached_texts)]
                
                # 更新结果和缓存
                for i, (text, translated) in enumerate(zip(uncached_texts, translated_results)):
                    idx = uncached_indices[i]
                    results[idx] = translated
                    
                    # 更新缓存
                    if self.translation_cache:
                        cache_key = self._get_text_hash(text, target_lang, source_lang or 'auto')
                        self.translation_cache.put(cache_key, translated)
                
                self.stats['api_calls'] += 1
                
            except Exception as e:
                logger.error(f"批次翻译失败: {e}")
                # 失败时返回原文
                for i, idx in enumerate(uncached_indices):
                    results[idx] = uncached_texts[i]
        
        return results, from_cache_flags

    def _enhanced_result_validation(self, original_text: str, result: str, from_cache: bool) -> Tuple[bool, str]:
        """增强的结果验证"""
        
        # 基础验证
        is_serious, reason = TranslationValidator.is_serious_failure(
            original_text, result, self.large_text_threshold, from_cache)
        
        if is_serious:
            return False, reason
        
        # 检查是否为循环翻译（翻译回原文）
        if self._is_circular_translation(original_text, result):
            return False, "检测到循环翻译"
        
        return True, "验证通过"

    def _is_circular_translation(self, original: str, translated: str) -> bool:
        """检查是否为循环翻译 - 更宽松的检测"""
        # 简单的循环检测
        if len(original) < 10 or len(translated) < 10:
            return False  # 短文本不检测循环
            
        similarity = len(set(original.lower().split()) & set(translated.lower().split()))
        total_words = len(set(original.lower().split()) | set(translated.lower().split()))
        
        if total_words > 0 and similarity / total_words > 0.9:  # 更宽松的阈值
            return True
        
        return False

    def _get_enhanced_system_prompt(self, target_lang: str, source_lang: str) -> str:
        """获取增强的系统提示 - 支持配置合并后的自定义prompt配置"""
        
        # 根据合并后的prompt配置构建系统提示
        if self.effective_prompt_config:
            mode = self.effective_prompt_config.get('mode', 'none')
            
            if mode == 'custom' and self.custom_prompt:
                # 使用完全自定义的prompt
                if isinstance(self.custom_prompt, dict):
                    system_content = self.custom_prompt.get('system', '')
                else:
                    system_content = str(self.custom_prompt)
                
                self.stats['custom_prompt_used'] = True
                logger.info("Using merged custom prompt for single text translation")
                return system_content
            
            elif mode == 'professional' and self.prompt_template:
                # 使用专业模板 - 扩展的专业领域
                logger.info(f"Using professional template from merged config: {self.prompt_template}")
                
                professional_prompts = {
                    'academic': f"""You are an expert academic translator specializing in scholarly documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain academic tone, preserve citations and references, and use appropriate academic terminology.
Ensure consistency in technical terms throughout the translation.""",
                
                    'business': f"""You are a professional business translator with expertise in corporate documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use appropriate business terminology, maintain formal tone, and keep company names/brands unchanged.
Ensure clarity and professionalism in the translation.""",
                
                    'technical': f"""You are a technical translator specializing in technical documentation.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve technical accuracy, keep code snippets and commands unchanged, and use industry-standard terminology.
Maintain consistency in technical terms throughout.""",
                
                    'legal': f"""You are a certified legal translator with expertise in legal documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use precise legal terminology, maintain legal accuracy and formality, and preserve all legal references.
Ensure no ambiguity in legal terms.""",
                
                    'medical': f"""You are a certified medical translator with expertise in medical documents.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use standard medical terminology, preserve drug names and dosages exactly, and maintain clinical precision.
Follow international medical nomenclature standards.""",
                
                    'creative': f"""You are a creative translator focusing on maintaining style and tone.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Preserve the original style, adapt idioms naturally, and maintain emotional impact.
Focus on readability and flow while being faithful to the original meaning.""",
                
                    # 新增专业领域
                    'scientific': f"""You are a scientific translator with expertise in research papers and scientific publications.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain scientific accuracy, preserve equations and formulas, use standard scientific terminology.
Ensure precision in methodology descriptions and results interpretation.""",
                
                    'financial': f"""You are a financial translator specializing in financial documents and reports.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use precise financial terminology, maintain accuracy in numerical data, preserve financial statements format.
Ensure compliance with international financial reporting standards.""",
                
                    'marketing': f"""You are a marketing translator focusing on promotional and marketing content.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Maintain persuasive tone, adapt cultural references appropriately, preserve brand messaging.
Focus on engagement and cultural sensitivity while maintaining marketing effectiveness.""",
                
                    'educational': f"""You are an educational content translator specializing in learning materials.
Translate from {source_lang or 'auto-detected language'} to {target_lang}.
Use clear and accessible language, maintain pedagogical structure, preserve learning objectives.
Ensure content remains suitable for the intended educational level."""
                }
                
                system_content = professional_prompts.get(self.prompt_template, professional_prompts['academic'])
                self.stats['custom_prompt_used'] = True
                
        if not hasattr(self, 'system_content') or 'system_content' not in locals():
            # 默认增强prompt逻辑
            base_format_rules = """
CRITICAL MARKDOWN FORMAT PRESERVATION RULES:
1. PRESERVE ALL MARKDOWN SYNTAX - Keep headers (#, ##, ###), lists (-, *, 1.), quotes (>), etc.
2. MAINTAIN EXACT STRUCTURE - Do not alter indentation, spacing, or line breaks
3. KEEP SPECIAL CHARACTERS - Preserve *, _, `, [], (), and other markdown symbols
4. NO FORMAT CHANGES - Do not add or remove markdown formatting
5. PRESERVE LINE STRUCTURE - Each input line should correspond to one output line
6. MAINTAIN LIST FORMATTING - Keep bullet points, numbering, and indentation exactly as original
7. PRESERVE EMPHASIS - Keep **bold**, *italic*, `code` formatting intact
8. KEEP LINKS - Preserve [text](url) and reference-style links exactly
"""

            # 构建增强规则 - 使用合并后的配置
            enhanced_rules = []
            if self.effective_prompt_config:
                if self.preserve_terms:
                    terms_text = ', '.join(self.preserve_terms)
                    enhanced_rules.append(f"PRESERVE THESE TERMS EXACTLY: {terms_text}")
                
                if self.glossary:
                    glossary_text = '; '.join([f"{k}: {v}" for k, v in self.glossary.items()])
                    enhanced_rules.append(f"USE THIS GLOSSARY: {glossary_text}")
                
                if self.additional_context:
                    enhanced_rules.append(f"ADDITIONAL CONTEXT: {self.additional_context}")

            enhanced_rules_text = "\n".join([f"- {rule}" for rule in enhanced_rules])
            if enhanced_rules_text:
                enhanced_rules_text = f"\nADDITIONAL REQUIREMENTS:\n{enhanced_rules_text}"

            if self._is_portuguese(source_lang):
                system_content = f"""You are a professional translator specializing in Portuguese to {target_lang} translation.

{base_format_rules}

PORTUGUESE-SPECIFIC TRANSLATION RULES:
1. TRANSLATE ALL CONTENT - Never leave any Portuguese text untranslated
2. HANDLE ACCENTS PROPERLY - Process ã, ê, ô, ç, á, é, í, ó, ú correctly
3. COMPLETE SENTENCES - Translate entire sentences, not fragments
4. TECHNICAL PRECISION - Maintain accuracy in technical/legal terminology
5. NO LANGUAGE MIXING - Output only in {target_lang}, no Portuguese remnants
6. PRESERVE NAMES - Keep place names, company names, and proper nouns unchanged
7. MAINTAIN CONTEXT - Consider surrounding content for accurate translation
8. Do not include any extra comments in your output
9. Return only the translated content without including the original text

TECHNICAL TERMINOLOGY:
- CONTRATADA → Contractor
- CONTRATANTE → Client/Contracting Party  
- LTDA → Ltd.
- SERVIÇOS → Services
- RESPONSABILIDADES → Responsibilities
- EXECUÇÃO → Execution
- ÂMBITO → Scope
- MUNICÍPIO → Municipality
{enhanced_rules_text}

OUTPUT FORMAT: Return only the translated content with identical markdown structure."""

            else:
                system_content = f"""You are a professional translator from {source_lang or 'auto-detected language'} to {target_lang}.

{base_format_rules}

TRANSLATION REQUIREMENTS:
1. TRANSLATE ALL TEXT - Convert all readable content to {target_lang}
2. PRESERVE NAMES - Keep place names, company names, and proper nouns unchanged
3. MAINTAIN CONTEXT - Consider surrounding content for accurate translation
4. NATURAL LANGUAGE - Ensure translations sound natural in {target_lang}
5. CONSISTENCY - Use consistent terminology throughout
6. Do not include any extra comments in your output
7. Return only the translated content without including the original text
{enhanced_rules_text}

OUTPUT FORMAT: Return only the translated content with identical markdown structure."""

        return system_content

    def _get_enhanced_batch_prompt(self, target_lang: str, source_lang: str) -> str:
        """获取增强的批量翻译提示 - 支持自定义prompt配置"""
        
        # 如果有自定义prompt配置，优先使用
        if self.effective_prompt_config:
            mode = self.effective_prompt_config.get('mode', 'default')
            
            if mode == 'custom' and self.custom_prompt:
                # 对于批量翻译，需要调整自定义prompt
                if isinstance(self.custom_prompt, dict):
                    batch_custom_prompt = self.custom_prompt.get('system', '')
                else:
                    batch_custom_prompt = str(self.custom_prompt)
                
                if "numbered line" not in batch_custom_prompt.lower():
                    batch_custom_prompt += "\n\nFor batch translation: Process each numbered line and return the same number of translated lines."
                self.stats['custom_prompt_used'] = True
                logger.info("Using custom prompt for batch translation")
                return batch_custom_prompt
        
        # 默认批量prompt逻辑
        format_preservation = """
MANDATORY FORMAT PRESERVATION:
- Keep ALL markdown syntax: #, ##, *, -, 1., >, **, __, `, etc.
- Preserve exact indentation and spacing
- Maintain line-by-line correspondence
- Do not add or remove formatting elements
- Keep list markers and numbering exactly as shown
"""

        # 构建增强规则
        enhanced_rules = []
        if self.effective_prompt_config:
            if self.preserve_terms:
                terms_text = ', '.join(self.preserve_terms)
                enhanced_rules.append(f"PRESERVE THESE TERMS EXACTLY: {terms_text}")
            
            if self.glossary:
                glossary_text = '; '.join([f"{k}: {v}" for k, v in self.glossary.items()])
                enhanced_rules.append(f"USE THIS GLOSSARY: {glossary_text}")
            
            if self.additional_context:
                enhanced_rules.append(f"ADDITIONAL CONTEXT: {self.additional_context}")

        enhanced_rules_text = "\n".join([f"{i+11}. {rule}" for i, rule in enumerate(enhanced_rules)])

        if self._is_portuguese(source_lang):
            return f"""Professional Portuguese to {target_lang} batch translator.

{format_preservation}

PORTUGUESE BATCH TRANSLATION PROTOCOL:
1. Translate each numbered line completely to {target_lang}
2. Never leave Portuguese text untranslated
3. Handle Portuguese accents (ã, ê, ô, ç) properly
4. Maintain technical term consistency
5. Preserve all markdown formatting exactly
6. Return same number of lines as input
7. No line numbers in output - just translated content
8. Do not include any extra comments in your output
9. Return only the translated content without including the original text
10. Maintain context across all lines for coherence
{enhanced_rules_text}

TECHNICAL TERMS GUIDE:
CONTRATADA=Contractor, CONTRATANTE=Client, LTDA=Ltd., SERVIÇOS=Services

Output format: Translated lines only, preserving exact markdown structure."""

        else:
            return f"""Professional {source_lang or 'source language'} to {target_lang} batch translator.

{format_preservation}

BATCH TRANSLATION PROTOCOL:
1. Translate each numbered line to natural {target_lang}
2. Maintain context across all lines for coherence
3. Preserve all markdown formatting exactly
4. Keep proper nouns and place names unchanged
5. Return same number of lines as input
6. No line numbers in output - just translated content
7. Do not include any extra comments in your output
8. Return only the translated content without including the original text
{enhanced_rules_text}

Output format: Translated lines only, preserving exact markdown structure."""

    def process_document(
        self,
        input_filepath: str,
        output_filepath: str,
        target_lang: str,
        source_lang: Optional[str] = None
    ) -> str:
        """处理文档翻译 - 修复版本，集成完整重试机制和时间统计"""
        try:
            # 开始总计时
            self.time_stats['total_start_time'] = time.time()
            
            # 应用prompt配置到翻译器
            with self._translator_config_context():
                logger.info(f"开始处理: {os.path.basename(input_filepath)}")
                logger.info(f"源语言: {source_lang}, 目标语言: {target_lang}")
                logger.info(f"配置合并策略: {self.stats['config_merge_strategy_used']}")
                logger.info(f"模板路径: {self.template_path if self.template_path else '不使用模板'}")
                logger.info(f"重试设置: max_retries={self.max_retries}, timeout={self.translation_timeout}s, retry_workers={self.retry_max_workers}")
                if self.effective_prompt_config:
                    logger.info(f"有效Prompt配置: mode={self.effective_prompt_config.get('mode')}, template={self.prompt_template}")
                
                # 1. DOCX → Markdown
                conversion_start = time.time()
                markdown_content = self._convert_docx_to_markdown(input_filepath)
                self.time_stats['conversion_time'] += time.time() - conversion_start
                
                # 2. 解析Markdown结构 (表格作为原子级元素)
                markdown_structure = self._parse_markdown_structure(markdown_content)
                
                if not markdown_structure:
                    logger.info("没有需要翻译的内容")
                    shutil.copy2(input_filepath, output_filepath)
                    self._finalize_timing()
                    return output_filepath
                
                # 3. 提取翻译任务 - 保持原始顺序
                translation_tasks = self._extract_translation_tasks_from_markdown(
                    markdown_structure, source_lang
                )
                
                if not translation_tasks:
                    logger.info("没有需要翻译的文本")
                    shutil.copy2(input_filepath, output_filepath)
                    self._finalize_timing()
                    return output_filepath
                
                # 4. 验证任务顺序
                self._validate_task_order(translation_tasks)
                
                # 清空失败任务列表
                with self.failed_tasks_lock:
                    self.failed_tasks.clear()
                
                # 5. 主翻译阶段 - 计时
                main_translation_start = time.time()
                translated_results = self._execute_batch_translation(
                    translation_tasks, target_lang, source_lang
                )
                self.time_stats['main_translation_time'] = time.time() - main_translation_start
                
                logger.info(f"主翻译完成，用时: {self.time_stats['main_translation_time']:.1f}秒, "
                           f"结果数: {len(translated_results)}")
                
                # 6. 验证翻译结果顺序
                self._validate_translation_order(translation_tasks, translated_results)
                
                # 7. 智能判断是否需要重试 - 传递翻译结果引用
                if self._should_trigger_retry(len(translation_tasks)):
                    logger.info(f"检测到严重失败，开始自适应重试流程")
                    self._adaptive_retry_strategy(target_lang, source_lang, translated_results)
                else:
                    with self.failed_tasks_lock:
                        serious_count = sum(1 for task in self.failed_tasks if task.is_serious)
                        minor_count = sum(1 for task in self.failed_tasks if not task.is_serious)
                        if self.failed_tasks:
                            logger.info(f"失败分析: 严重失败 {serious_count} 个, 轻微问题 {minor_count} 个, 未达到重试阈值，跳过重试")
                
                # 8. 最终处理剩余失败任务 - 传递翻译结果引用
                self._final_retry_remaining_tasks(target_lang, source_lang, translated_results)
                
                logger.info(f"所有重试完成，最终翻译结果数: {len(translated_results)}")
                
                # 9. 重建翻译后的Markdown - 严格按原始结构顺序
                translated_markdown = self._rebuild_markdown_with_translations(
                    markdown_structure, translated_results
                )
                
                # 10. Markdown → DOCX
                conversion_start = time.time()
                self._convert_markdown_to_docx(translated_markdown, output_filepath)
                self.time_stats['conversion_time'] += time.time() - conversion_start
                
                # 11. 清理DOCX文件中的Pandoc标记
                self._clean_docx_pandoc_marks(output_filepath)
                
                # 12. 应用文档格式设置
                formatting_start = time.time()
                self._apply_document_formatting(output_filepath)
                self.time_stats['formatting_time'] = time.time() - formatting_start
                
                # 13. 完成时间统计和日志
                self._finalize_timing()
                self._log_results_enhanced(source_lang)
                
                return output_filepath
                
        except Exception as e:
            logger.error(f"翻译失败: {e}")
            raise

    def _finalize_timing(self):
        """完成时间统计"""
        self.time_stats['total_time'] = time.time() - self.time_stats['total_start_time']
        logger.info(f"总用时: {self.time_stats['total_time']:.1f}秒 (主翻译: {self.time_stats['main_translation_time']:.1f}s, "
                   f"重试: {self.time_stats['retry_time']:.1f}s, 转换: {self.time_stats['conversion_time']:.1f}s, "
                   f"格式: {self.time_stats['formatting_time']:.1f}s)")

    def _convert_docx_to_markdown(self, docx_path: str) -> str:
        """使用pandoc将DOCX转换为Markdown - 保持基本参数"""
        try:
            result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True)
            if result.returncode != 0:
                raise RuntimeError("pandoc未安装或不可用")
            
            # 基本转换参数，保持格式完整
            result = subprocess.run([
                'pandoc',
                '-f', 'docx',
                '-t', 'markdown',
                '--wrap=none',
                '--extract-media=.',
                docx_path
            ], capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode != 0:
                raise RuntimeError(f"pandoc转换失败: {result.stderr}")
            
            logger.info(f"DOCX转Markdown成功，内容长度: {len(result.stdout)}")
            return result.stdout
            
        except Exception as e:
            logger.error(f"DOCX转Markdown失败: {e}")
            raise

    def _parse_markdown_structure(self, markdown_content: str) -> List[Dict[str, Any]]:
        """解析Markdown结构 - 表格作为原子级元素"""
        structure = []
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 检测表格开始 - 简单判断：包含 | 就是表格
            if '|' in line and line.strip():
                # 收集整个表格作为一个元素
                table_lines = []
                table_start = i
                
                # 连续收集所有表格行
                while i < len(lines) and '|' in lines[i] and lines[i].strip():
                    table_lines.append(lines[i])
                    i += 1
                
                # 表格作为原子级元素
                table_content = '\n'.join(table_lines)
                structure.append({
                    'line_number': table_start,
                    'original_line': table_content,
                    'type': 'table_block',
                    'content': table_content,
                    'skip_translation': True,
                    'metadata': {'line_count': len(table_lines)},
                    'original_index': table_start
                })
                
                self.stats['skipped_tables'] += 1
                logger.debug(f"识别表格块，行数: {len(table_lines)}")
                continue
            
            # 处理非表格行
            line_info = self._parse_single_line(line, i)
            structure.append(line_info)
            i += 1
        
        logger.info(f"解析Markdown结构完成，共 {len(structure)} 个元素 (包含 {self.stats['skipped_tables']} 个表格)")
        return structure

    def _parse_single_line(self, line: str, line_number: int) -> Dict[str, Any]:
        """解析单行为元素"""
        element = {
            'line_number': line_number,
            'original_line': line,
            'type': 'paragraph',
            'content': line.strip(),
            'skip_translation': False,
            'original_index': line_number
        }
        
        # 检测标题
        if re.match(r'^#+\s+', line):
            level = len(re.match(r'^#+', line).group(0))
            content = re.sub(r'^#+\s+', '', line).strip()
            element.update({
                'type': 'header',
                'content': content,
                'metadata': {'level': level}
            })
        
        # 检测列表项
        elif re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            indent_match = re.match(r'^(\s*)', line)
            indent = len(indent_match.group(1)) if indent_match else 0
            
            if re.match(r'^\s*\d+\.\s+', line):
                marker_match = re.match(r'^\s*(\d+\.)\s+', line)
                content = re.sub(r'^\s*\d+\.\s+', '', line).strip()
                list_type = 'ordered'
            else:
                marker_match = re.match(r'^\s*([-*+])\s+', line)
                content = re.sub(r'^\s*[-*+]\s+', '', line).strip()
                list_type = 'unordered'
            
            marker = marker_match.group(1) if marker_match else ''
            element.update({
                'type': 'list_item',
                'content': content,
                'metadata': {
                    'list_type': list_type,
                    'marker': marker,
                    'indent': indent
                }
            })
        
        # 检测引用
        elif line.strip().startswith('>'):
            content = re.sub(r'^\s*>\s*', '', line).strip()
            element.update({
                'type': 'quote',
                'content': content
            })
        
        # 检测代码块、空行等 - 跳过翻译
        elif (line.strip().startswith('```') or 
              not line.strip() or 
              line.strip().startswith('![') or  # 图片
              line.strip().startswith('<!--')):  # 注释
            
            if line.strip().startswith('```'):
                element_type = 'code_block'
            elif line.strip().startswith('!['):
                element_type = 'image'
            elif line.strip().startswith('<!--'):
                element_type = 'comment'
            else:
                element_type = 'empty_line'
            
            element.update({
                'type': element_type,
                'skip_translation': True
            })
        
        return element

    def _extract_translation_tasks_from_markdown(
        self, 
        markdown_structure: List[Dict[str, Any]], 
        source_lang: Optional[str]
    ) -> List[Dict[str, Any]]:
        """从Markdown结构中提取翻译任务 - 保持原始文档顺序"""
        tasks = []
        use_simple_strategy = self._is_portuguese(source_lang)
        
        for elem_idx, element in enumerate(markdown_structure):
            # 跳过所有标记为不翻译的元素
            if element.get('skip_translation', False):
                continue
            
            # 只处理需要翻译的文本元素
            if element['type'] in ['header', 'paragraph', 'list_item', 'quote']:
                content = element['content'].strip()
                if (content and 
                    len(content) > 1 and 
                    any(c.isalpha() for c in content)):
                    
                    if use_simple_strategy:
                        self.stats['portuguese_simple_strategy_used'] += 1
                    
                    tasks.append({
                        'type': element['type'],
                        'id': f"{element['type']}_{elem_idx}",
                        'text': content,
                        'element_ref': element,
                        'original_index': elem_idx,
                        'sequence_order': len(tasks),
                        'char_count': len(content),
                        'use_simple_strategy': use_simple_strategy,
                        'source_lang': source_lang
                    })
            
            self.stats['total_elements'] += 1
        
        logger.info(f"提取到 {len(tasks)} 个翻译任务（严格保持原始顺序）")
        if use_simple_strategy:
            logger.info(f"使用葡萄牙语简化策略: {self.stats['portuguese_simple_strategy_used']} 个")
        
        return tasks

    def _validate_task_order(self, tasks: List[Dict[str, Any]]) -> bool:
        """验证翻译任务的顺序正确性"""
        
        # 检查任务是否按原始顺序
        for i in range(1, len(tasks)):
            current_index = tasks[i].get('original_index', 0)
            previous_index = tasks[i-1].get('original_index', 0)
            if current_index < previous_index:
                logger.warning(f"任务顺序异常：任务 {i} 的原始索引({current_index}) 小于任务 {i-1} 的原始索引({previous_index})")
                self.stats['sequence_preserved'] = False
                return False
        
        logger.info("翻译任务顺序验证通过 - 保持原始文档顺序")
        self.stats['sequence_preserved'] = True
        return True

    def _validate_translation_order(self, tasks: List[Dict[str, Any]], 
                                  translations: Dict[str, str]) -> bool:
        """验证翻译结果的顺序正确性"""
        
        # 检查翻译结果完整性
        missing_translations = []
        for task in tasks:
            if task['id'] not in translations:
                missing_translations.append(task['id'])
        
        if missing_translations:
            logger.warning(f"缺失翻译结果：{missing_translations}")
            self.stats['order_validation_passed'] = False
            return False
        
        # 检查翻译结果是否按原始顺序
        translation_order_correct = True
        for i in range(1, len(tasks)):
            current_task = tasks[i]
            previous_task = tasks[i-1]
            
            current_index = current_task.get('original_index', 0)
            previous_index = previous_task.get('original_index', 0)
            
            if current_index < previous_index:
                logger.warning(f"翻译结果顺序异常：任务 {current_task['id']} 的原始位置({current_index}) 在 {previous_task['id']} 的原始位置({previous_index}) 之前")
                translation_order_correct = False
        
        if translation_order_correct:
            logger.info("翻译结果顺序验证通过")
            self.stats['order_validation_passed'] = True
        else:
            self.stats['order_validation_passed'] = False
        
        return translation_order_correct

    def _execute_batch_translation(self, tasks: List[Dict[str, Any]], 
                                 target_lang: str, source_lang: Optional[str]) -> Dict[str, str]:
        """执行批量翻译 - 集成失败检测机制"""
        
        # 1. 文本去重和缓存检查
        unique_texts = {}
        cache_hits = 0
        
        for task in tasks:
            text = task['text']
            text_hash = self._get_text_hash(text, target_lang, source_lang or 'auto')
            
            if self.translation_cache:
                cached = self.translation_cache.get(text_hash)
                if cached:
                    unique_texts[task['id']] = cached
                    cache_hits += 1
                    continue
            
            unique_texts[task['id']] = text
        
        self.stats['cache_hits'] = cache_hits
        
        # 2. 智能分批 - 保持顺序
        translation_results = {}
        texts_to_translate = [(k, v) for k, v in unique_texts.items() 
                             if not (self.translation_cache and 
                                   self._get_text_hash(v, target_lang, source_lang or 'auto') in self.translation_cache._cache)]
        
        if texts_to_translate:
            batches = self._create_smart_batches(texts_to_translate, tasks)
            
            if self.max_workers > 1 and len(batches) > 3:
                translation_results = self._translate_concurrent_batches(batches, target_lang, source_lang)
            else:
                translation_results = self._translate_sequential_batches(batches, target_lang, source_lang)
        
        # 3. 更新缓存和处理重复文本
        if self.translation_cache:
            for task_id, translated_text in translation_results.items():
                original_text = unique_texts[task_id]
                text_hash = self._get_text_hash(original_text, target_lang, source_lang or 'auto')
                self.translation_cache.put(text_hash, translated_text)
        
        # 4. 按原始任务顺序构建最终结果
        final_results = {}
        for task in tasks:  # 按原始任务顺序遍历
            task_id = task['id']
            text = task['text']
            
            if task_id in translation_results:
                final_results[task_id] = translation_results[task_id]
            elif task_id in unique_texts:
                final_results[task_id] = unique_texts[task_id]
            else:
                text_hash = self._get_text_hash(text, target_lang, source_lang or 'auto')
                if self.translation_cache:
                    cached = self.translation_cache.get(text_hash)
                    if cached:
                        final_results[task_id] = cached
                    else:
                        final_results[task_id] = text
                else:
                    final_results[task_id] = text
        
        logger.info(f"批量翻译完成，共处理 {len(final_results)} 个任务，保持原始顺序")
        return final_results

    def _create_smart_batches(self, texts_with_ids: List[Tuple[str, str]], 
                            tasks: List[Dict[str, Any]]) -> List[List[Tuple[str, str]]]:
        """创建智能批次 - 保持文档顺序和连续性"""
        batches = []
        current_batch = []
        current_chars = 0
        
        # 按原始顺序处理
        task_order = {task['id']: task.get('sequence_order', 0) for task in tasks}
        
        # 按翻译任务的原始顺序排序
        sorted_texts = sorted(texts_with_ids, key=lambda x: task_order.get(x[0], 999999))
        
        effective_batch_size = self.batch_size
        effective_chunk_size = self.max_chunk_size
        
        for task_id, text in sorted_texts:
            text_len = len(text)
            
            if (len(current_batch) >= effective_batch_size or 
                current_chars + text_len > effective_chunk_size) and current_batch:
                batches.append(current_batch)
                current_batch = [(task_id, text)]
                current_chars = text_len
            else:
                current_batch.append((task_id, text))
                current_chars += text_len
        
        if current_batch:
            batches.append(current_batch)
        
        logger.info(f"创建了 {len(batches)} 个批次，严格保持文档顺序")
        return batches

    def _translate_concurrent_batches(self, batches: List[List[Tuple[str, str]]], 
                                    target_lang: str, source_lang: Optional[str]) -> Dict[str, str]:
        """并发批量翻译 - 保证结果顺序并集成失败检测"""
        results = {}
        batch_futures = {}
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            with tqdm(total=sum(len(batch) for batch in batches), 
                     desc="翻译进度", unit="文本") as pbar:
                
                # 提交所有批次，保存批次索引
                for batch_idx, batch in enumerate(batches):
                    future = executor.submit(
                        self._translate_single_batch, batch, target_lang, source_lang, batch_idx
                    )
                    batch_futures[future] = (batch_idx, batch)
                
                # 收集结果，但要按批次顺序重新排列
                batch_results = {}
                
                for future in as_completed(batch_futures):
                    batch_idx, batch = batch_futures[future]
                    try:
                        batch_result = future.result()
                        batch_results[batch_idx] = batch_result
                        pbar.update(len(batch))
                    except Exception as e:
                        logger.error(f"批次 {batch_idx} 翻译失败: {e}")
                        # 失败时保持原文，并添加到失败任务列表
                        batch_result = {task_id: text for task_id, text in batch}
                        batch_results[batch_idx] = batch_result
                        
                        # 添加批次失败
                        self._add_batch_failure(batch, [batch_idx] * len(batch), f"批次执行异常: {str(e)}")
                        
                        pbar.update(len(batch))
                
                # 按批次顺序合并结果
                for batch_idx in sorted(batch_results.keys()):
                    results.update(batch_results[batch_idx])
                    self.stats['batches_reordered'] += 1
        
        logger.info(f"并发翻译完成，重新排序了 {self.stats['batches_reordered']} 个批次")
        return results

    def _translate_sequential_batches(self, batches: List[List[Tuple[str, str]]], 
                                    target_lang: str, source_lang: Optional[str]) -> Dict[str, str]:
        """顺序批量翻译 - 天然保持顺序，集成失败检测"""
        results = {}
        
        with tqdm(total=sum(len(batch) for batch in batches), 
                 desc="翻译进度", unit="文本") as pbar:
            
            for batch_idx, batch in enumerate(batches):
                try:
                    batch_results = self._translate_single_batch(batch, target_lang, source_lang, batch_idx)
                    results.update(batch_results)
                except Exception as e:
                    logger.error(f"顺序批次 {batch_idx} 翻译失败: {e}")
                    # 失败时保持原文，并添加到失败任务列表
                    batch_results = {task_id: text for task_id, text in batch}
                    results.update(batch_results)
                    
                    # 添加批次失败
                    self._add_batch_failure(batch, [batch_idx] * len(batch), f"顺序执行异常: {str(e)}")
                
                pbar.update(len(batch))
        
        logger.info("顺序翻译完成，天然保持原始顺序")
        return results

    def _translate_single_batch(self, batch: List[Tuple[str, str]], 
                              target_lang: str, source_lang: Optional[str], 
                              batch_idx: int) -> Dict[str, str]:
        """翻译单个批次 - 集成失败检测"""
        
        if len(batch) == 1:
            task_id, text = batch[0]
            try:
                translated = self._translate_single_text(text, target_lang, source_lang)
                
                # 检查翻译结果
                is_valid, reason = self._enhanced_result_validation(text, translated, False)
                if not is_valid:
                    self._add_failed_task(text, task_id, batch_idx, translated, False)
                
                self.stats['single_translations'] += 1
                self.stats['api_calls'] += 1
                return {task_id: translated}
            except Exception as e:
                logger.error(f"单文本翻译失败: {e}")
                self._add_failed_task(text, task_id, batch_idx, f"单文本翻译异常: {str(e)}", False)
                return {task_id: text}
        else:
            try:
                results = self._translate_text_batch(batch, target_lang, source_lang)
                
                # 检查批次翻译结果
                for task_id, text in batch:
                    if task_id in results:
                        translated = results[task_id]
                        is_valid, reason = self._enhanced_result_validation(text, translated, False)
                        if not is_valid:
                            self._add_failed_task(text, task_id, batch_idx, translated, False)
                
                self.stats['batch_translations'] += 1
                self.stats['api_calls'] += 1
                return results
            except Exception as e:
                logger.error(f"批量翻译失败: {e}")
                # 添加整批失败
                self._add_batch_failure(batch, [batch_idx] * len(batch), f"批量翻译异常: {str(e)}")
                return {task_id: text for task_id, text in batch}

    def _translate_text_batch(self, batch: List[Tuple[str, str]], 
                            target_lang: str, source_lang: Optional[str]) -> Dict[str, str]:
        """批量翻译一组文本 - 使用增强的批量prompt"""
        
        numbered_texts = []
        task_ids = []
        
        for i, (task_id, text) in enumerate(batch):
            numbered_texts.append(f"[{i+1}] {text}")
            task_ids.append(task_id)
        
        user_content = "\n".join(numbered_texts)
        system_prompt = self._get_enhanced_batch_prompt(target_lang, source_lang or 'auto-detected')
        
        try:
            # 支持配置合并策略的翻译器调用
            if hasattr(self.translator, 'translate') and hasattr(self.translator, 'set_prompt_config'):
                # 使用翻译器的prompt配置能力
                response = self.translator.translate(
                    text=user_content,
                    target_lang=target_lang,
                    source_lang=source_lang,
                    prompt_config=self.effective_prompt_config,
                    config_merge_mode=self.config_merge_mode
                )
            else:
                # 回退到messages格式
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ]
                response = self.translator.translate(
                    messages=messages,
                    target_lang=target_lang,
                    source_lang=source_lang
                )
            
            translated_lines = self._parse_batch_response(response, len(batch))
            
            results = {}
            for i, task_id in enumerate(task_ids):
                if i < len(translated_lines) and translated_lines[i]:
                    results[task_id] = translated_lines[i]
                else:
                    results[task_id] = batch[i][1]
            
            self.stats['total_texts_translated'] += len(batch)
            return results
            
        except Exception as e:
            logger.error(f"批量翻译失败: {e}")
            return {task_id: text for task_id, text in batch}

    def _parse_batch_response(self, response: str, expected_count: int) -> List[str]:
        """解析批量翻译响应"""
        lines = [line.strip() for line in response.strip().split('\n') if line.strip()]
        
        cleaned_lines = []
        for line in lines:
            # 清理行号前缀
            cleaned = re.sub(r'^\[\d+\]\s*', '', line)
            cleaned = re.sub(r'^\d+\.\s*', '', cleaned)
            cleaned = re.sub(r'^\d+\)\s*', '', cleaned)
            
            if cleaned:
                cleaned_lines.append(cleaned)
        
        # 确保返回正确数量的行
        while len(cleaned_lines) < expected_count:
            cleaned_lines.append("")
        
        return cleaned_lines[:expected_count]

    def _translate_single_text(self, text: str, target_lang: str, source_lang: Optional[str]) -> str:
        """单个文本翻译 - 使用增强的系统prompt"""
        
        system_prompt = self._get_enhanced_system_prompt(target_lang, source_lang or 'auto-detected')
        
        try:
            # 支持配置合并策略的翻译器调用
            if hasattr(self.translator, 'translate') and hasattr(self.translator, 'set_prompt_config'):
                # 使用翻译器的prompt配置能力
                response = self.translator.translate(
                    text=text,
                    target_lang=target_lang,
                    source_lang=source_lang,
                    prompt_config=self.effective_prompt_config,
                    config_merge_mode=self.config_merge_mode
                )
            else:
                # 回退到messages格式
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ]
                response = self.translator.translate(
                    messages=messages,
                    target_lang=target_lang,
                    source_lang=source_lang
                )
            
            return response.strip()
        except Exception as e:
            logger.error(f"单个翻译失败: {e}")
            return text

    def _rebuild_markdown_with_translations(
        self, 
        markdown_structure: List[Dict[str, Any]], 
        translations: Dict[str, str]
    ) -> str:
        """重建翻译后的Markdown - 严格按原始结构顺序"""
        
        translated_lines = []
        
        with tqdm(total=len(markdown_structure), desc="重建Markdown", unit="元素") as pbar:
            for elem_idx, element in enumerate(markdown_structure):
                
                # 所有标记为跳过翻译的元素直接输出原始内容
                if element.get('skip_translation', False):
                    translated_lines.append(element['original_line'])
                    if element['type'] == 'table_block':
                        logger.debug(f"保留表格块原样 (行 {elem_idx})")
                
                elif element['type'] == 'header':
                    task_id = f"header_{elem_idx}"
                    translated_content = translations.get(task_id, element['content'])
                    level = element['metadata'].get('level', 1)
                    header_line = '#' * level + ' ' + translated_content
                    translated_lines.append(header_line)
                    self.stats['headers_translated'] += 1
                
                elif element['type'] == 'list_item':
                    task_id = f"list_item_{elem_idx}"
                    translated_content = translations.get(task_id, element['content'])
                    metadata = element['metadata']
                    
                    indent = ' ' * metadata.get('indent', 0)
                    marker = metadata.get('marker', '-')
                    list_line = indent + marker + ' ' + translated_content
                    translated_lines.append(list_line)
                    self.stats['list_items_translated'] += 1
                
                elif element['type'] == 'paragraph':
                    task_id = f"paragraph_{elem_idx}"
                    translated_content = translations.get(task_id, element['content'])
                    translated_lines.append(translated_content)
                    self.stats['paragraphs_translated'] += 1
                
                elif element['type'] == 'quote':
                    task_id = f"quote_{elem_idx}"
                    translated_content = translations.get(task_id, element['content'])
                    quote_line = '> ' + translated_content
                    translated_lines.append(quote_line)
                    self.stats['quotes_translated'] += 1
                
                else:
                    # 其他类型保持原样
                    translated_lines.append(element['original_line'])
                
                self.stats['translated_elements'] += 1
                pbar.update(1)
        
        logger.info(f"重建完成，共 {len(translated_lines)} 行，严格保持原始顺序")
        return '\n'.join(translated_lines)

    def _convert_markdown_to_docx(self, markdown_content: str, output_path: str):
        """使用pandoc将Markdown转换为DOCX"""
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
                temp_md.write(markdown_content)
                temp_md_path = temp_md.name
            
            try:
                result = subprocess.run([
                    'pandoc',
                    '-f', 'markdown',
                    '-t', 'docx',
                    '--wrap=none',
                    '-o', output_path,
                    temp_md_path
                ], capture_output=True, text=True, encoding='utf-8')
                
                if result.returncode != 0:
                    raise RuntimeError(f"pandoc转换失败: {result.stderr}")
                
                logger.info(f"Markdown转DOCX成功: {output_path}")
                
            finally:
                if os.path.exists(temp_md_path):
                    os.unlink(temp_md_path)
                    
        except Exception as e:
            logger.error(f"Markdown转DOCX失败: {e}")
            raise

    def _clean_docx_pandoc_marks(self, docx_path: str):
        """清理DOCX文件中的Pandoc标记"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx库未安装，跳过DOCX标记清理。请安装: pip install python-docx")
            return
        
        try:
            logger.info(f"开始清理DOCX文件中的Pandoc标记: {docx_path}")
            
            doc = Document(docx_path)
            total_cleaned = 0
            
            # 定义需要清理的Pandoc属性模式
            patterns_to_clean = [
                r'\{\.mark\}',
                r'\{\.highlight\}',
                r'\{\.underline\}',
                r'\{\.strikethrough\}',
                r'\{\.small-caps\}',
                r'\{\.inserted\}',
                r'\{\.deleted\}',
                r'\{\#[\w\-]+\}',              # {#id}
                r'\{\.[\w\-]+\}',              # {.class}
                r'\{[\w\-]+=[\w\-"\']+\}',     # {attr=value}
                r'\{[\.#\w\-="\'\s]+\}',       # 组合属性
            ]
            
            # 清理所有段落中的标记
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                if '{' in original_text and '}' in original_text:
                    cleaned_text = original_text
                    for pattern in patterns_to_clean:
                        matches = re.findall(pattern, cleaned_text)
                        if matches:
                            cleaned_text = re.sub(pattern, '', cleaned_text)
                            total_cleaned += len(matches)
                    
                    # 清理多余空格
                    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
                    
                    # 只有文本发生变化时才更新
                    if cleaned_text != original_text:
                        paragraph.text = cleaned_text
            
            # 清理所有表格中的标记
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text
                        if '{' in original_text and '}' in original_text:
                            cleaned_text = original_text
                            for pattern in patterns_to_clean:
                                matches = re.findall(pattern, cleaned_text)
                                if matches:
                                    cleaned_text = re.sub(pattern, '', cleaned_text)
                                    total_cleaned += len(matches)
                            
                            # 清理多余空格
                            cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
                            
                            # 只有文本发生变化时才更新
                            if cleaned_text != original_text:
                                cell.text = cleaned_text
            
            # 保存清理后的文档
            if total_cleaned > 0:
                doc.save(docx_path)
                logger.info(f"DOCX标记清理完成，共清理 {total_cleaned} 个标记")
            else:
                logger.info("DOCX文件中未发现需要清理的Pandoc标记")
            
            self.stats['docx_marks_cleaned'] = total_cleaned
            
        except Exception as e:
            logger.error(f"DOCX标记清理失败: {e}")

    def _apply_document_formatting(self, docx_path: str):
        """应用文档格式设置"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_LINE_SPACING
        except ImportError:
            logger.warning("python-docx库未安装，跳过文档格式设置。请安装: pip install python-docx")
            return
        
        try:
            # 如果指定了模板，使用模板创建新文档
            if self.template_path and os.path.exists(self.template_path):
                logger.info(f"使用模板创建文档: {self.template_path}")
                
                # 读取翻译后的内容
                doc = Document(docx_path)
                content_paragraphs = []
                for paragraph in doc.paragraphs:
                    content_paragraphs.append(paragraph.text)
                
                # 读取模板
                template_doc = Document(self.template_path)
                
                # 清空模板内容并添加翻译后的内容
                for paragraph in template_doc.paragraphs:
                    paragraph.clear()
                
                # 添加翻译后的内容到模板
                for content in content_paragraphs:
                    if content.strip():
                        template_doc.add_paragraph(content)
                
                # 保存使用模板的文档
                template_doc.save(docx_path)
                logger.info("模板应用完成")
                self.stats['formatting_applied'] = True
                
            else:
                logger.info("开始应用默认字体和格式设置")
                
                doc = Document(docx_path)
                
                # 设置默认字体样式
                for paragraph in doc.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            # 设置字体
                            run.font.name = 'Segoe UI Semilight'  # 西文字体
                            run.font.size = None  # 默认字体大小
                            
                            # 设置中文字体
                            run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
                    
                    # 设置段落行距为1.25倍
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    paragraph.paragraph_format.line_spacing = 1.25
                
                # 处理表格格式 - 只设置字体为10号，不设置行间距
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs:
                                    for run in paragraph.runs:
                                        # 表格字体设置为10号
                                        run.font.name = 'Segoe UI Semilight'  # 西文字体
                                        run.font.size = Pt(10)  # 表格字体大小
                                        
                                        # 设置中文字体
                                        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
                                
                                # 表格段落不设置行间距，保持默认
                                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # 保存格式设置后的文档
                doc.save(docx_path)
                logger.info("默认格式设置完成 - 中文宋体，西文Segoe UI Semilight，1.25倍行距，表格10号字体")
                self.stats['formatting_applied'] = True
                
        except Exception as e:
            logger.error(f"文档格式设置失败: {e}")

    def _get_text_hash(self, text: str, target_lang: str, source_lang: str) -> str:
        """生成文本哈希用于缓存 - 包含配置信息"""
        # 包含配置信息在哈希中，确保不同配置的翻译分开缓存
        config_info = ""
        if self.effective_prompt_config:
            mode = self.effective_prompt_config.get('mode', '')
            template = self.effective_prompt_config.get('prompt_template', '')
            config_info = f"_{mode}_{template}"
        
        content = f"{text}_{target_lang}_{source_lang}{config_info}"
        return hashlib.md5(content.encode('utf-8')).hexdigest()

    def _log_results_enhanced(self, source_lang: Optional[str]):
        """记录结果 - 增强版，包含完整时间统计和重试信息"""
        logger.info(f"=== 翻译完成 ===")
        logger.info(f"总用时: {self.time_stats['total_time']:.1f}秒")
        logger.info(f"  - 主翻译: {self.time_stats['main_translation_time']:.1f}秒")
        logger.info(f"  - 重试阶段: {self.time_stats['retry_time']:.1f}秒")
        logger.info(f"  - 文档转换: {self.time_stats['conversion_time']:.1f}秒")
        logger.info(f"  - 格式设置: {self.time_stats['formatting_time']:.1f}秒")
        
        logger.info(f"源语言: {source_lang}, 葡萄牙语判断: {self._is_portuguese(source_lang)}")
        logger.info(f"配置合并策略: {self.stats['config_merge_strategy_used']}")
        logger.info(f"重试设置: max_retries={self.max_retries}, timeout={self.translation_timeout}s, retry_workers={self.retry_max_workers}")
        
        # 重试机制统计 - 增强版
        if self.stats['serious_failures'] > 0 or self.stats['minor_issues'] > 0:
            logger.info(f"\n=== 问题分析 ===")
            if self.stats['serious_failures'] > 0:
                logger.info(f"严重失败: {self.stats['serious_failures']} 个（>50字符未翻译或API错误）")
            if self.stats['minor_issues'] > 0:
                logger.info(f"轻微问题: {self.stats['minor_issues']} 个（短文本相同等，正常情况）")
            
            if self.stats['retry_attempts'] > 0:
                logger.info(f"\n=== 重试统计 ===")
                logger.info(f"重试轮数: {self.stats['retry_attempts']} 轮")
                logger.info(f"重试策略: {self.stats['retry_strategy_used']}")
                logger.info(f"重试Worker: 最大{self.retry_max_workers}个，实际使用{self.stats['retry_workers_used']}个")
                logger.info(f"并行重试使用: {self.stats['parallel_retry_used']} 次")
                if self.stats['concurrent_retry_batches'] > 0:
                    logger.info(f"并发重试批次: {self.stats['concurrent_retry_batches']}个")
                if self.translation_cache:
                    logger.info(f"缓存清理: {self.stats['cache_clears']} 次")
                logger.info(f"总重试任务: {self.stats['total_retry_tasks']} 个")
                logger.info(f"成功救援: {self.stats['rescued_tasks']} 个")
                logger.info(f"重试成功率: {self.stats['retry_success_rate']:.1%}")
                if self.stats['final_rescues'] > 0:
                    logger.info(f"最终挽救: {self.stats['final_rescues']} 个")
                if self.stats['final_failures'] > 0:
                    logger.info(f"最终失败: {self.stats['final_failures']} 个（保持原文）")
                else:
                    logger.info("✅ 所有严重失败都已成功处理！")
            else:
                logger.info("✅ 未达到重试阈值，无需重试")
        else:
            logger.info("✅ 无翻译问题")
        
        logger.info(f"\n=== 基础统计 ===")
        logger.info(f"模板使用: {'是' if self.template_path else '否'}")
        logger.info(f"格式设置: {'成功' if self.stats['formatting_applied'] else '跳过'}")
        logger.info(f"合并后Prompt配置: {'是' if self.stats['custom_prompt_used'] else '否'}")
        logger.info(f"顺序保持状态: {'成功' if self.stats['sequence_preserved'] else '异常'}")
        logger.info(f"顺序验证结果: {'通过' if self.stats['order_validation_passed'] else '失败'}")
        logger.info(f"批次重排序: {self.stats['batches_reordered']} 个")
        logger.info(f"元素: {self.stats['translated_elements']}/{self.stats['total_elements']}")
        logger.info(f"跳过表格: {self.stats['skipped_tables']} 个 (作为原子级元素)")
        logger.info(f"清理DOCX标记: {self.stats['docx_marks_cleaned']} 个")
        logger.info(f"API调用: {self.stats['api_calls']} 次")
        
        if self.effective_prompt_config:
            mode = self.effective_prompt_config.get('mode', 'none')
            logger.info(f"有效配置模式: {mode}")
            if mode == 'professional':
                logger.info(f"专业领域: {self.effective_prompt_config.get('prompt_template', 'academic')}")
            elif mode == 'custom':
                logger.info("使用了自定义Prompt")
            
            # 显示合并的增强功能
            enhancements = []
            if self.preserve_terms:
                enhancements.append(f"保留术语({len(self.preserve_terms)}个)")
            if self.glossary:
                enhancements.append(f"术语表({len(self.glossary)}条)")
            if self.additional_context:
                enhancements.append("额外上下文")
            
            if enhancements:
                logger.info(f"合并的增强功能: {', '.join(enhancements)}")
        
        if self.stats['portuguese_simple_strategy_used'] > 0:
            logger.info(f"葡萄牙语简化策略: {self.stats['portuguese_simple_strategy_used']} 个")
        
        if self.stats['api_calls'] > 0:
            avg_texts_per_call = self.stats['total_texts_translated'] / self.stats['api_calls']
            logger.info(f"平均每次API调用翻译: {avg_texts_per_call:.1f} 个文本")
        
        logger.info(f"\n=== 翻译统计 ===")
        logger.info(f"  - 标题: {self.stats['headers_translated']}")
        logger.info(f"  - 段落: {self.stats['paragraphs_translated']}")
        logger.info(f"  - 列表项: {self.stats['list_items_translated']}")
        logger.info(f"  - 引用: {self.stats['quotes_translated']}")
        if self.translation_cache:
            cache_stats = self.translation_cache.stats
            logger.info(f"  - 缓存命中: {self.stats['cache_hits']}, 命中率: {cache_stats['hit_rate']:.1%}")


def translate_docx_via_markdown(
    input_filepath: str,
    output_dir: str,
    target_lang: str,
    translator,
    source_lang: Optional[str] = None,
    unique_filename_base: Optional[str] = None,
    max_chunk_size: int = 8000,
    batch_size: int = 50,
    max_workers: int = 5,
    retry_max_workers: int = 5,  # 重试worker参数
    template_path: Optional[str] = None,
    prompt_config=None,
    config_merge_mode: Optional[Union[str, ConfigMergeStrategy]] = None,
    translation_timeout: int = 60,  # 超时参数
    max_retries: int = 8,  # 最大重试次数
    large_text_threshold: int = 50,  # 大文本阈值
    retry_failure_threshold: float = 0.0,  # 重试阈值
    **kwargs
) -> str:
    """
    基于Markdown转换的DOCX翻译函数 - 修复版本，解决翻译结果更新问题
    
    主要修复:
        - 重试成功的结果现在会正确更新到最终文档
        - 优化并行重试策略，默认使用并行处理
        - 增加完整的时间统计（总用时、主翻译时间、重试时间等）
        - 改进失败检测逻辑，减少误判
        - 增强重试成功率统计
    """
    
    if not os.path.exists(input_filepath):
        return f"Error: 输入文件未找到: {input_filepath}"
    
    # 检查模板文件
    if template_path and not os.path.exists(template_path):
        logger.warning(f"模板文件未找到: {template_path}，将使用默认格式设置")
        template_path = None
    
    os.makedirs(output_dir, exist_ok=True)
    
    input_path = Path(input_filepath)
    if unique_filename_base:
        output_filename = f"{unique_filename_base}_translated_{target_lang}.docx"
    else:
        output_filename = f"{input_path.stem}_translated_{target_lang}.docx"
    
    output_filepath = os.path.join(output_dir, output_filename)
    
    counter = 1
    while os.path.exists(output_filepath):
        if unique_filename_base:
            output_filename = f"{unique_filename_base}_translated_{target_lang}_{counter}.docx"
        else:
            output_filename = f"{input_path.stem}_translated_{target_lang}_{counter}.docx"
        output_filepath = os.path.join(output_dir, output_filename)
        counter += 1
    
    # 配置合并和翻译器配置管理
    original_prompt_config = None
    try:
        # 获取有效配置（应用配置合并策略）
        effective_config = get_effective_config(prompt_config, config_merge_mode)
        
        # 如果提供了有效配置，应用到翻译器
        if effective_config and hasattr(translator, 'set_prompt_config'):
            # 保存翻译器的原始prompt配置
            original_prompt_config = getattr(translator, 'prompt_config', None)
            translator.set_prompt_config(effective_config)
            logger.info(f"Applied effective config to translator for markdown-based DOCX translation: mode={effective_config.get('mode')}, strategy={config_merge_mode or _global_merge_strategy.value}")
        
        # 从有效配置中提取批处理参数
        if effective_config:
            if 'max_chunk_size' in effective_config:
                max_chunk_size = effective_config['max_chunk_size']
            if 'batch_size' in effective_config:
                batch_size = effective_config['batch_size']
            if 'max_workers' in effective_config:
                max_workers = effective_config['max_workers']
        
        # 从kwargs中提取其他配置参数，保持向后兼容
        mode = kwargs.get('mode', 'optimized_with_retry_fix' if effective_config else 'optimized')
        max_units_per_chunk = kwargs.get('max_units_per_chunk', 100)
        
        # 处理prompt相关的kwargs参数
        processor_kwargs = {}
        if effective_config:
            # 传递有效配置到处理器
            processor_kwargs['prompt_config'] = prompt_config  # 保持原始前端配置
            processor_kwargs['config_merge_mode'] = config_merge_mode
            
            # 处理其他prompt相关参数
            for key in ['prompt_template', 'custom_prompt', 'preserve_terms', 'glossary', 'additional_context']:
                if key in kwargs:
                    processor_kwargs[key] = kwargs[key]
        
        processor = MarkdownBasedDocxTranslator(
            translator=translator,
            batch_size=batch_size,
            max_chunk_size=max_chunk_size,
            max_workers=max_workers,
            retry_max_workers=retry_max_workers,
            use_cache=True,
            template_path=template_path,
            mode=mode,
            max_units_per_chunk=max_units_per_chunk,
            translation_timeout=translation_timeout,
            max_retries=max_retries,
            large_text_threshold=large_text_threshold,
            retry_failure_threshold=retry_failure_threshold,
            **processor_kwargs
        )
        
        result_path = processor.process_document(
            input_filepath=input_filepath,
            output_filepath=output_filepath,
            target_lang=target_lang,
            source_lang=source_lang
        )
        
        if not os.path.exists(result_path):
            return "Error: 输出文件未创建"
        
        file_size = os.path.getsize(result_path)
        if file_size == 0:
            return "Error: 输出文件为空"
        
        logger.info(f"✅ Successfully translated DOCX via markdown with retry mechanism fix. Output: {result_path}")
        logger.info(f"⏱️  Total time: {processor.time_stats['total_time']:.1f}s, Retry success rate: {processor.stats['retry_success_rate']:.1%}")
        return result_path
        
    except Exception as e:
        error_msg = f"翻译失败: {str(e)}"
        logger.error(error_msg, exc_info=True)
        return f"Error: {error_msg}"
    
    finally:
        # 恢复翻译器的原始prompt配置
        if effective_config and hasattr(translator, 'set_prompt_config'):
            if original_prompt_config is not None:
                translator.set_prompt_config(original_prompt_config)
                logger.debug("Restored original prompt config to translator")
            else:
                # 如果原来没有配置，清除当前配置
                if hasattr(translator, 'prompt_config'):
                    translator.prompt_config = None
                    logger.debug("Cleared prompt config from translator")


if __name__ == '__main__':
    print("=== 智能Markdown-based DOCX翻译器（修复重试结果更新版）===")
    print("✅ 支持前端所有prompt配置")
    print("✅ 智能配置合并策略 (merge/override/instance_only)") 
    print("✅ 扩展专业领域模板 (10+ 领域)")
    print("✅ 自定义prompt完全支持")
    print("✅ 术语管理和术语表智能合并")
    print("✅ 基于Markdown的精确转换")
    print("✅ 表格作为原子级元素保护")
    print("✅ Pandoc标记清理")
    print("✅ 实例级配置管理")
    print("✅ 完整的错误处理和日志")
    print("✅ 严格文档顺序保持")
    print("✅ 顺序验证和监控机制")
    print("✅ 并发处理结果重排序")
    print("🔥 **FIXED**: 重试结果更新问题 - 重试成功的翻译现在会正确应用到最终文档")
    print("🔥 **ENHANCED**: 优化并行重试 - 默认使用并行处理，显著提高重试效率")
    print("🔥 **ENHANCED**: 完整时间统计 - 总用时、主翻译、重试、转换、格式化分项计时")
    print("🔥 **ENHANCED**: 减少误判 - 改进失败检测逻辑，降低正常翻译被误判的概率")
    print("🔥 **NEW**: large_text_threshold = 50")
    print("🔥 **NEW**: retry_batch_sizes = [10, 5, 2, 1, 1, 1, 1, 1]")
    print("🔥 **NEW**: 重试成功率统计")
    
    print("\n=== 核心修复说明 ===")
    print("🎯 翻译结果更新: 重试成功的结果通过original_task_id直接更新translated_results")
    print("🎯 并行重试优化: 简化并行判断条件，优先使用并发处理")
    print("🎯 失败检测改进: 提高TranslationValidator的容错性，减少误判")
    print("🎯 时间统计完善: 分项记录各阶段用时，便于性能分析")
    print("🎯 成功率监控: 实时跟踪重试成功率和救援效果")
    
    print("\n智能Markdown-based DOCX翻译器（修复版）已就绪！🚀")
    print("现在重试成功的翻译会正确更新到最终文档，不再出现'很多没有翻译'的问题！")
    print("优化的并行重试机制大幅提升翻译效率和成功率！")